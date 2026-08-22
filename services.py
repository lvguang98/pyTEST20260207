"""services.py - merged backend services.

Originally data_service.py / file_service.py / template_service.py.
Public class names unchanged: DataService / FileService / TemplateVariableManager / TemplateService.
"""
import os
import subprocess
import json
import datetime
import tempfile
import shutil
import logging
import win32com.client
import pandas as pd
from docx import Document
from typing import Dict, List, Any, Optional, Tuple
from path_utils import path_utils


class DataService:
    """数据处理服务 - 负责所有数据计算、验证、转换操作"""

    def __init__(self):
        """初始化DataService"""
        pass

    def validate_case_data(self, case_data: Dict[str, Any]) -> List[str]:
        """
        统一验证案件数据
        :param case_data: 包含所有案件数据的字典
        :return: 错误信息列表
        """
        errors = []

        # 1. 验证人员信息
        person_data = {
            '姓名': case_data.get('姓名', ''),
            '身份证号': case_data.get('身份证号', ''),
            '年龄': case_data.get('年龄', '')
        }
        person_errors = self.validate_person_data(person_data)
        if person_errors:
            errors.extend([f"个人信息: {err}" for err in person_errors])

        # 2. 验证公司信息
        company_name = case_data.get('用工单位', '')
        if not company_name:
            errors.append("用工单位不能为空")

        # 3. 验证案件信息
        case_number = case_data.get('案本号', '')
        if not case_number:
            errors.append("案本号不能为空")

        return errors

    @staticmethod
    def validate_required_fields(data: Dict[str, Any], field_names: List[str]) -> List[str]:
        """
        验证必填字段
        :param data: 数据字典
        :param field_names: 字段名列表
        :return: 错误信息列表
        """
        errors = []
        for field in field_names:
            value = data.get(field)
            if value is None or (isinstance(value, str) and not value.strip()):
                errors.append(f"{field}不能为空")
        return errors

    # ============ 身份证相关方法 ============

    @staticmethod
    def calculate_age_from_idcard(idcard: str) -> Optional[int]:
        """
        根据身份证号计算年龄
        :param idcard: 身份证号码（15位或18位）
        :return: 年龄（整数），如果身份证无效返回None
        """
        try:
            if not idcard or len(idcard) not in (15, 18):
                return None

            # 提取出生日期
            if len(idcard) == 18:
                birth_date_str = idcard[6:14]  # YYYYMMDD
            else:  # 15位身份证
                birth_date_str = f"19{idcard[6:12]}"  # 19YYMMDD

            # 转换为日期
            birth_date = datetime.datetime.strptime(birth_date_str, "%Y%m%d")

            # 计算年龄
            today = datetime.datetime.now()
            age = today.year - birth_date.year

            # 如果今年生日还没过，减1岁
            if (today.month, today.day) < (birth_date.month, birth_date.day):
                age -= 1

            return age

        except Exception as e:
            print(f"[DataService] 计算年龄失败: {e}")
            return None

    @staticmethod
    def extract_gender_from_idcard(idcard: str) -> Optional[str]:
        """
        根据身份证号提取性别
        :param idcard: 身份证号码（15位或18位）
        :return: "男" 或 "女"，如果身份证无效返回None
        """
        try:
            if not idcard:
                return None

            if len(idcard) == 18:
                gender_digit = int(idcard[16])  # 第17位
            elif len(idcard) == 15:
                gender_digit = int(idcard[14])  # 第15位
            else:
                return None

            # 奇数男性，偶数女性
            return "男" if gender_digit % 2 == 1 else "女"

        except Exception as e:
            print(f"[DataService] 提取性别失败: {e}")
            return None

    @staticmethod
    def validate_idcard(idcard: str) -> Tuple[bool, str]:
        """
        验证身份证号格式
        :param idcard: 身份证号码
        :return: (是否有效, 错误信息)
        """
        if not idcard:
            return False, "身份证号不能为空"

        # 移除空格
        idcard = idcard.strip()

        # 检查长度
        if len(idcard) not in (15, 18):
            return False, "身份证号必须是15位或18位"

        # 检查字符
        if len(idcard) == 18:
            # 18位：前17位必须是数字，最后一位可以是数字或X
            if not idcard[:17].isdigit():
                return False, "身份证号前17位必须是数字"
            if not (idcard[17].isdigit() or idcard[17].upper() == 'X'):
                return False, "身份证号最后一位必须是数字或X"
        else:  # 15位
            if not idcard.isdigit():
                return False, "15位身份证号必须全部是数字"

        return True, "身份证号格式正确"

    # ============ 数据验证方法 ============

    @staticmethod
    def validate_person_data(person_data: Dict[str, Any]) -> List[str]:
        """
        验证人员数据
        :param person_data: 人员数据字典
        :return: 错误信息列表
        """
        errors = []

        # 验证姓名
        name = person_data.get('姓名', '')
        if not name:
            errors.append("姓名不能为空")
        elif len(name) > 15:
            errors.append("姓名长度不能超过15个字符")

        # 验证身份证号
        idcard = person_data.get('身份证号', '')
        if idcard:
            is_valid, msg = DataService.validate_idcard(idcard)
            if not is_valid:
                errors.append(f"身份证号: {msg}")

        # 验证年龄
        age = person_data.get('年龄', '')
        if age:
            if isinstance(age, str):
                if not age.isdigit():
                    errors.append("年龄必须是数字")
                else:
                    age_int = int(age)
                    if not (16 <= age_int <= 100):
                        errors.append("年龄必须在16-100岁之间")
            elif isinstance(age, int):
                if not (16 <= age <= 100):
                    errors.append("年龄必须在16-100岁之间")

        return errors

    @staticmethod
    def get_current_datetime(format_str: str = "%Y年%m月%d日%H时%M分") -> str:
        """
        获取当前日期时间并格式化
        :param format_str: 格式字符串
        :return: 格式化后的日期时间字符串
        """
        return datetime.datetime.now().strftime(format_str)

    # ============ 案件相关数据方法 ============

    @staticmethod
    def generate_case_number(base_number: int, is_death_case: bool = False,
                             date_str: Optional[str] = None) -> str:
        """
        生成案本号
        :param base_number: 基础编号
        :param is_death_case: 是否为工亡案件
        :param date_str: 日期字符串（YYYYMMDD），如果不提供则使用当前日期
        :return: 案本号字符串
        """
        if not date_str:
            date_str = datetime.datetime.now().strftime("%Y%m%d")

        case_type = "GW" if is_death_case else "GS"
        return f"{date_str}-{case_type}-{base_number:03d}"




class FileService:
    def __init__(self, base_path, logger=None):
        self.BASE_PATH = base_path
        self.logger = logger or self._create_default_logger()
        self.WPS_PATH = self.find_wps_path()
        self.logger.setLevel(logging.WARNING)

    def create_enhanced_case_folder(self, base_path: str, case_info: Dict[str, Any]) -> str:
        """创建案件文件夹 —— 直接用案本号命名"""
        import os

        case_number = case_info['case_number']

        folder_name = case_number
        folder_path = os.path.join(base_path, folder_name)

        # 处理重复
        counter = 1
        while os.path.exists(folder_path):
            folder_name = f"{case_number}-{counter:02d}"
            folder_path = os.path.join(base_path, folder_name)
            counter += 1

        os.makedirs(folder_path, exist_ok=True)
        return folder_path

    def _create_default_logger(self):
        """创建默认日志器"""
        import logging
        logger = logging.getLogger('FileService')
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            logger.setLevel(logging.INFO)
        return logger

    def open_document(self, file_path):
        """打开文档（统一使用这个方法）"""
        self.logger.info(f"打开文档: {file_path}")
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                self.logger.error(f"文件不存在: {file_path}")
                return False, "文件不存在"

            # 尝试用WPS打开
            if self.WPS_PATH and os.path.exists(self.WPS_PATH):
                try:
                    self.logger.info(f"尝试用WPS打开: {self.WPS_PATH}")
                    subprocess.Popen([self.WPS_PATH, file_path])
                    self.logger.info("WPS打开成功")
                    return True, "用WPS打开成功"
                except Exception as e:
                    self.logger.error(f"WPS打开失败: {e}")

            # 尝试用Word打开
            try:
                self.logger.info("尝试用Word打开")
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = True
                word_app.Documents.Open(file_path)
                self.logger.info("Word打开成功")
                return True, "用Word打开成功"
            except Exception as e:
                self.logger.error(f"Word打开失败: {e}")

            # 尝试用系统默认程序打开
            try:
                self.logger.info("尝试用默认程序打开")
                os.startfile(file_path)
                self.logger.info("默认程序打开成功")
                return True, "用默认程序打开成功"
            except Exception as e:
                self.logger.error(f"默认程序打开失败: {e}")
                return False, "所有打开方式都失败"

        except Exception as e:
            self.logger.error(f"打开文档异常: {e}")
            return False, f"打开文件时发生错误: {str(e)}"

    def ensure_folder_exists(self, folder_path):
        """确保文件夹存在"""
        os.makedirs(folder_path, exist_ok=True)
        return folder_path

    def find_wps_path(self):
        """从原类复制，完全一样"""
        possible_paths = [
            r"C:\Program Files (x86)\Kingsoft\WPS Office\ksolaunch.exe",
            r"C:\Program Files\Kingsoft\WPS Office\ksolaunch.exe",
            os.path.join(os.environ.get('LOCALAPPDATA', ''), "Kingsoft", "WPS Office", "ksolaunch.exe"),
        ]

        for path in possible_paths:
            if os.path.exists(path):
                return path
        return None

    def get_case_folder_path(self, person_name, case_versions=None):
        """
        获取案件文件夹路径
        person_name: 人员姓名
        case_versions: 版本字典（可选）
        返回：完整的文件夹路径
        """
        if case_versions and person_name in case_versions:
            version = case_versions[person_name]
            folder_name = f"{person_name}{version:02d}"
        else:
            # 扫描现有文件夹
            existing_folders = []
            if os.path.exists(self.BASE_PATH):
                existing_folders = [
                    f for f in os.listdir(self.BASE_PATH)
                    if f.startswith(person_name) and
                       os.path.isdir(os.path.join(self.BASE_PATH, f))
                ]

            if existing_folders:
                max_version = 0
                for folder in existing_folders:
                    version_str = folder.replace(person_name, "").strip()
                    if version_str.isdigit():
                        version = int(version_str)
                        if version > max_version:
                            max_version = version
                version = max_version + 1
            else:
                version = 1

            if case_versions:
                case_versions[person_name] = version

            folder_name = f"{person_name}{version:02d}"

        return os.path.join(self.BASE_PATH, folder_name)

    def create_case_folder(self, person_name, case_versions=None):
        """
        创建案件文件夹
        返回：创建的文件夹路径
        """
        folder_path = self.get_case_folder_path(person_name, case_versions)
        os.makedirs(folder_path, exist_ok=True)
        return folder_path

    def save_to_excel(self, template_path, excel_filename, column_name, new_item, existing_items):
        """保存到Excel - 使用文书模板目录"""
        from path_utils import path_utils
        excel_path = path_utils.get_document_template_path(excel_filename)

        print(f"💾 保存Excel到: {excel_path}")

        try:
            if os.path.exists(excel_path):
                existing_df = pd.read_excel(excel_path)
                existing_list = existing_df[column_name].tolist()
                all_items = list(set(existing_list + [new_item] + existing_items))
                df = pd.DataFrame(all_items, columns=[column_name])
            else:
                all_items = list(set([new_item] + existing_items))
                df = pd.DataFrame(all_items, columns=[column_name])

            df.to_excel(excel_path, index=False)
            return df[column_name].tolist()

        except Exception as e:
            print(f"[FileService ERROR] 保存到Excel失败: {str(e)}")
            return existing_items

    def search_cases_fuzzy(self, base_path: str, keyword: str) -> List[Dict[str, Any]]:
        """根据关键词模糊搜索案件（匹配文件夹名中任意部分）"""
        import os
        import re
        from datetime import datetime

        if not keyword or not keyword.strip():
            return []

        kw = keyword.strip()
        matched_cases = []

        for folder_name in os.listdir(base_path):
            folder_path = os.path.join(base_path, folder_name)
            if not os.path.isdir(folder_path):
                continue

            # 模糊匹配：关键词出现在文件夹名的任意位置
            if kw not in folder_name:
                continue

            # 解析文件夹名：张三-案本202608111234 或 张三-案本202608111234-01
            pattern = r'^(.+?)-(案本|工亡)(\d{8})(\d{4})(-\d{2})?$'
            match = re.match(pattern, folder_name)
            if match:
                name_in_folder = match.group(1)
                case_number = folder_name.rstrip(match.group(5) or '')
                id_last4 = match.group(4)
            else:
                case_number = folder_name
                name_in_folder = ''
                id_last4 = ''

            created_time = os.path.getctime(folder_path)
            created_date = datetime.fromtimestamp(created_time).strftime('%Y-%m-%d')
            file_count = len([f for f in os.listdir(folder_path) if f.endswith('.docx')])

            matched_cases.append({
                'folder_name': folder_name,
                'folder_path': folder_path,
                'case_number': case_number,
                'person_name': name_in_folder,
                'id_last4': id_last4,
                'created_date': created_date,
                'file_count': file_count,
                'has_person': self._check_has_person_record(folder_path, name_in_folder)
            })

        matched_cases.sort(key=lambda x: x['folder_name'], reverse=True)
        return matched_cases

    def search_cases_by_person_name(self, base_path: str, person_name: str) -> List[Dict[str, Any]]:
        """
        根据姓名搜索案件文件夹

        Args:
            base_path: 基础路径
            person_name: 要搜索的姓名

        Returns:
            匹配的案件列表
        """
        import os
        import re
        from datetime import datetime

        if not person_name or not person_name.strip():
            return []

        search_name = person_name.strip()
        matched_cases = []

        for folder_name in os.listdir(base_path):
            folder_path = os.path.join(base_path, folder_name)

            if not os.path.isdir(folder_path):
                continue

            # 解析文件夹名：张三-案本202608111234
            pattern = r'^(.+?)-(案本|工亡)(\d{8})(\d{4})(-\d{2})?$'
            match = re.match(pattern, folder_name)

            if match:
                name_in_folder = match.group(1)

                if name_in_folder == search_name:
                    case_number = folder_name.rstrip(match.group(5) or '')
                    id_last4 = match.group(4)
                    created_time = os.path.getctime(folder_path)
                    created_date = datetime.fromtimestamp(created_time).strftime('%Y-%m-%d')
                    file_count = len([f for f in os.listdir(folder_path)
                                      if f.endswith('.docx')])

                    matched_cases.append({
                        'folder_name': folder_name,
                        'folder_path': folder_path,
                        'case_number': case_number,
                        'person_name': name_in_folder,
                        'id_last4': id_last4,
                        'created_date': created_date,
                        'file_count': file_count,
                        'has_person': self._check_has_person_record(folder_path, name_in_folder)
                    })

        # 按创建时间倒序排列
        matched_cases.sort(key=lambda x: x['folder_name'], reverse=True)
        return matched_cases

    def _check_has_person_record(self, folder_path: str, person_name: str) -> bool:
        """检查文件夹中是否已有本人笔录"""
        import os
        import re

        for filename in os.listdir(folder_path):
            if filename.endswith('.docx'):
                # 匹配：本人谈话笔录（...）（本人）姓名.docx
                if re.search(r'本人谈话笔录.*本人.*' + re.escape(person_name), filename):
                    return True
        return False


class TemplateVariableManager:
    """模板变量管理器 - 从原文件迁移过来"""

    def __init__(self, data_model):
        self.data = data_model
        self.variables_cache: Dict[str, Any] = {}
        self.introduction_cache: Dict[str, str] = {}

    def clear_cache(self):
        """清空所有缓存（切换证人/角色等数据变化时调用，避免命中旧数据）"""
        self.variables_cache.clear()
        self.introduction_cache.clear()

    def get_self_introduction(self, role: str,
                              company: str, employer: str, site: str,
                              name: str, position: str, id_address: str = "") -> str:
        """生成自我介绍（带缓存）"""
        cache_key = f"{role}_{name}_{company}_{employer}_{site}"

        if cache_key in self.introduction_cache:
            return self.introduction_cache[cache_key]

        # 根据角色生成自我介绍
        if role == "证人":
            intro = self._generate_witness_intro(name, position, id_address, company, employer, site)
        elif role == "法人":
            intro = self._generate_legal_intro(name, position, id_address, company)
        else:  # 本人
            intro = self._generate_person_intro(name, position, id_address, company, employer, site)

        self.introduction_cache[cache_key] = intro
        return intro

    def _generate_person_intro(self, name: str, position: str, id_address: str,
                               company: str, employer: str, site: str) -> str:
        """生成本人自我介绍"""
        parts = []

        if id_address:
            parts.append(f"身份证地址为{id_address}")

        if company:
            parts.append(f"系{company}的职工")
            if employer:
                parts.append(f"被指派到{employer}")
                if site:
                    parts.append(f"承建的{site}工地")
            elif site:
                parts.append(f"被指派到{site}工地")
        else:
            parts.append("")

        # 构建结果
        base = f"答：我是{name}"
        if parts:
            base += "，" + "，".join(filter(None, parts))

        if position:
            base += f"，从事{position}工作。"
        else:
            base += "。"

        return base

    def _generate_witness_intro(self, name: str, position: str, id_address: str,
                                company: str, employer: str, site: str) -> str:
        """生成证人自我介绍"""
        parts = []

        if id_address:
            parts.append(f"身份证地址为{id_address}")

        if company:
            parts.append(f"系{company}的职工")
            if employer:
                parts.append(f"被指派到{employer}")
                if site:
                    parts.append(f"承建的{site}工地")
            elif site:
                parts.append(f"被指派到{site}工地")
        else:
            parts.append("")

        # 构建结果
        base = f"答：我是{name}"
        if parts:
            base += "，" + "，".join(filter(None, parts))

        if position:
            base += f"，从事{position}工作。"
        else:
            base += "。"

        return base

    def _generate_legal_intro(self, name: str, position: str,
                              id_address: str, company: str) -> str:
        """生成法人自我介绍"""
        parts = []

        if id_address:
            parts.append(f"身份证地址为{id_address}")

        if company:
            if position:
                parts.append(f"系{company}的{position}")
            else:
                parts.append(f"系{company}的负责人")
        else:
            if position:
                parts.append(f"系公司的{position}")
            else:
                parts.append("系公司负责人")

        # 构建结果
        base = f"答：我是{name}"
        if parts:
            base += "，" + "，".join(parts)

        base += "，负责公司的全面管理工作。"
        return base

    def collect_variables(self, role: str, case_type: int,
                          has_employer: bool, has_site: bool) -> Dict[str, Any]:
        """收集所有模板变量"""
        cache_key = f"{role}_{case_type}_{has_employer}_{has_site}"

        if cache_key in self.variables_cache:
            return self.variables_cache[cache_key]

        # 基础数据
        variables = self.data.to_template_dict()

        # 确保有 {{当前时期}} 字段
        from datetime import datetime
        current_date = variables.get('当前日期', datetime.now().strftime('%Y年%m月%d日'))
        current_time = variables.get('当前时间', datetime.now().strftime('%H时%M分'))
        variables['当前时期'] = f"{current_date}{current_time}"

        # 确保有本人年龄和性别
        if role == "本人":
            if '本人年龄' not in variables:
                age = self.data.basic_info.get('本人年龄', '')
                if not age:
                    age = self.data.basic_info.get('本人年龄', self.data.basic_info.get('年龄', ''))
                variables['本人年龄'] = age

            if '本人性别' not in variables:
                gender = self.data.basic_info.get('本人性别', '')
                if not gender:
                    gender = self.data.basic_info.get('本人性别', self.data.basic_info.get('性别', ''))
                variables['本人性别'] = gender

        # 案件类型描述
        case_descriptions = [
            "工伤保险条例第十四条第一款（普通案件）",
            "工伤保险条例第十四条第二款（工作前后案件）",
            "工伤保险条例第十四条第三款（暴力伤害案件）",
            "工伤保险条例第十四条第四款（患职业病案件）",
            "工伤保险条例第十四条第五款（因工外出案件）",
            "工伤保险条例第十四条第六款（上下班时案件）"
        ]
        variables['案件类型'] = case_descriptions[case_type]

        # 用工情况描述
        if has_employer:
            employer = self.data.company_info.get('用人单位', '')
            variables['用工情况'] = f"有用工单位：{employer}" if employer else "无用工单位"
        else:
            variables['用工情况'] = "无用工单位"

        # 工地情况描述
        if has_site:
            site = self.data.company_info.get('工地名称', '')
            variables['工地情况'] = f"在{site}工地" if site else "无特定工地"
        else:
            variables['工地情况'] = "无特定工地"

        # 缓存结果
        self.variables_cache[cache_key] = variables
        return variables


class TemplateService:
    """模板服务 - 负责所有模板相关操作"""

    def __init__(self, template_path: str):
        """
        初始化TemplateService
        :param template_path: 模板文件夹路径
        """
        from path_utils import path_utils

        self.template_path = template_path
        # 直接使用 path_utils 获取模板子目录
        self.template_base_path = str(path_utils.get_talk_template_path())
        self.document_template_path = str(path_utils.get_document_template_path())

    def get_template_path(self, role: str, case_type: int,
                          is_death_case: bool = False, is_personal: bool = False) -> str:
        """
        根据角色和案件类型获取模板路径
        :param role: 角色（本人/证人/法人）
        :param case_type: 案件类型索引（0-5）
        :param is_death_case: 是否为工亡案件
        :param is_personal: 是否为个人申请
        :return: 模板文件完整路径
        """
        if is_death_case:
            if is_personal:
                template_name = f"{role}谈话笔录（个人申请工亡案件）.docx"
            else:
                template_name = f"{role}谈话笔录（工亡案件）.docx"
        else:
            case_templates = {
                0: f"{role}谈话笔录（普通工伤案件）.docx",
                1: f"{role}谈话笔录（工作前后案件）.docx",
                2: f"{role}谈话笔录（暴力伤害案件）.docx",
                3: f"{role}谈话笔录（患职业病案件）.docx",
                4: f"{role}谈话笔录（因工外出案件）.docx",
                5: f"{role}谈话笔录（上下班时案件）.docx"
            }
            template_name = case_templates.get(case_type, f"{role}谈话笔录（普通工伤案件）.docx")

        template_full_path = os.path.join(self.template_base_path, template_name)

        if not os.path.exists(template_full_path):
            default_template = os.path.join(self.template_base_path, f"{role}谈话笔录（普通工伤案件）.docx")
            if os.path.exists(default_template):
                return default_template
            else:
                raise FileNotFoundError(f"模板文件不存在: {template_full_path}")

        return template_full_path

    def get_document_template_path(self, template_name: str) -> str:
        """
        获取文书模板路径
        :param template_name: 模板文件名
        :return: 完整模板路径
        """
        path = os.path.join(self.document_template_path, template_name)
        if not os.path.exists(path):
            raise FileNotFoundError(f"文书模板不存在: {path}")
        return path

    @staticmethod
    def _copy_basic_format(source_paragraph, target_paragraph):
        """复制基本段落格式"""
        try:
            target_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
            target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
            target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after

            if source_paragraph.runs and target_paragraph.runs:
                source_run = source_paragraph.runs[0]
                target_run = target_paragraph.runs[0]
                target_run.font.name = source_run.font.name
                target_run.font.size = source_run.font.size

        except Exception as e:
            print(f"[WARN] 格式复制失败: {e}")

    def _insert_introduction_to_doc(self, doc: Document, introduction_text: str):
        """向文档中插入自我介绍"""
        target_question = "问：请介绍一下你的姓名、住址、工作单位以及从事的工作？"

        for i, paragraph in enumerate(doc.paragraphs):
            if target_question in paragraph.text:
                if i + 1 < len(doc.paragraphs):
                    new_paragraph = doc.paragraphs[i + 1].insert_paragraph_before("")
                else:
                    new_paragraph = doc.add_paragraph("")

                run = new_paragraph.add_run(introduction_text)
                run.underline = True

                if i < len(doc.paragraphs):
                    self._copy_basic_format(doc.paragraphs[i], new_paragraph)
                break

    def _insert_questions_to_doc(self, doc: Document, questions: List[Dict[str, str]],
                                 variables: Dict[str, Any]):
        """向文档中插入问题"""
        for qa in questions:
            answer_hint = qa.get("answer_hint", "")
            if answer_hint:
                for key, value in variables.items():
                    if value:
                        placeholder = "{" + key + "}"
                        answer_hint = answer_hint.replace(placeholder, str(value))

                doc.add_paragraph(qa.get("question", ""))
                answer_para = doc.add_paragraph(answer_hint)

                for run in answer_para.runs:
                    run.underline = True

    def process_template(self, template_path: str, role_type: str, case_type: int,
                         variables: Dict[str, Any], need_special_questions: bool = False) -> str:
        """
        处理模板文件，插入动态内容
        :param template_path: 模板文件路径
        :param role_type: 角色类型
        :param case_type: 案件类型
        :param variables: 模板变量
        :param need_special_questions: 是否需要插入特殊问题
        :return: 处理后的临时模板文件路径
        """
        # 创建临时文件
        temp_path = os.path.join(
            tempfile.gettempdir(),
            f"temp_template_{os.path.basename(template_path)}"
        )

        # 复制模板到临时文件
        shutil.copy2(template_path, temp_path)

        # 打开文档并处理
        doc = Document(temp_path)

        # 插入自我介绍
        self_intro = variables.get('自我介绍内容', '')
        if self_intro:
            self._insert_introduction_to_doc(doc, self_intro)

        # 插入特殊问题（如果需要）
        if need_special_questions:
            questions = self.get_basic_question_template(case_type, role_type)
            if questions:
                self._insert_questions_to_doc(doc, questions, variables)

        # 保存修改
        doc.save(temp_path)
        return temp_path

    @staticmethod
    def get_basic_question_template(case_type: int, role: str = "本人") -> List[Dict[str, str]]:
        """
        根据案件类型和角色获取基本问题模板
        :param case_type: 案件类型索引
        :param role: 角色
        :return: 问题模板列表
        """
        templates = {
            "本人": {
                0: [
                    {"question": "问：你跟{{用人单位}}有没有签订劳动合同关系？有没有参加工伤保险？",
                     "answer_hint": "答：我与{{用人单位}}的劳动关系已经由永嘉县劳动人事争议仲裁委员会【浙永嘉劳人仲案（2025）XXX号】裁决书予以确认。"}
                ],
                1: [
                    {"question": "问：工作前进行了哪些准备工作？",
                     "answer_hint": "答：工作前我准备了{准备工作内容}。"}
                ],
                2: [
                    {"question": "问：请描述暴力伤害发生的具体经过？",
                     "answer_hint": "答：当时{施暴者姓名}突然{伤害行为}。"}
                ]
            },
            "证人": {
                0: [
                    {"question": "问：请介绍一下你的姓名、住址、工作单位以及从事的工作？",
                     "answer_hint": "答：我是{证人姓名}，住在{证人身份证地址}，为{用人单位}的职工，从事{证人岗位}工作。"}
                ]
            },
            "法人": {
                0: [
                    {"question": "问：请介绍一下你的姓名、职务以及公司基本情况？",
                     "answer_hint": "答：我是{法人姓名}，是{用人单位}的{法人职务}，负责公司{负责业务}。"}
                ]
            }
        }

        return templates.get(role, {}).get(case_type, [])