import os
import json
import datetime
import logging
from typing import Dict, List, Any, Tuple, Optional
from ctypes import windll, byref, create_string_buffer, c_int32, c_uint
import pandas as pd
from docx import Document
from docxtpl import DocxTemplate
from PyQt5.Qt import *
from PyQt5.QtCore import QThread, pyqtSignal, Qt
from PyQt5.QtWidgets import (
    QApplication, QWidget, QMessageBox, QDialog, QVBoxLayout,
    QLabel, QTextEdit, QPushButton, QHBoxLayout, QInputDialog,
    QLineEdit, QComboBox, QCompleter, QCheckBox, QProgressDialog,
    QTableWidget, QTableWidgetItem,
    QGroupBox
)
from PyQt5.QtGui import QFont

from ui_main_window import Ui_Form
from services import FileService, DataService, TemplateVariableManager
from ai_service import AIService
from config_service import ConfigService
from path_utils import path_utils
from main import UserManager, PasswordLineEdit

# 设置日志级别
logging.getLogger('config_service').setLevel(logging.WARNING)


# ============================================================================
# 工具函数
# ============================================================================

def _date_now() -> str:
    """当前日期，格式：2025年01月01日"""
    import datetime as _dt
    return _dt.datetime.now().strftime('%Y年%m月%d日')


def _time_now() -> str:
    """当前时间，格式：14时30分"""
    import datetime as _dt
    return _dt.datetime.now().strftime('%H时%M分')


def _timestamp_now() -> str:
    """时间戳，格式：20250101_143000"""
    import datetime as _dt
    return _dt.datetime.now().strftime('%Y%m%d_%H%M%S')


_CN_DIGITS = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']


def _witness_label(n: int) -> str:
    """把序号转成中文证人编号：1→证人一, 2→证人二, 10→证人十, 11→证人十一, 21→证人二十一"""
    if n <= 0:
        return f"证人{n}"

    if n <= 10:
        body = _CN_DIGITS[n] if n < 10 else "十"
    elif n < 20:
        body = "十" + (_CN_DIGITS[n % 10] if n % 10 else "")
    else:
        tens = n // 10
        ones = n % 10
        body = _CN_DIGITS[tens] + "十" + (_CN_DIGITS[ones] if ones else "")

    return f"证人{body}"


# ============================================================================
# 拟用条例 选项与格式互转
# ============================================================================

REGULATION_OPTIONS = [
    "第十四条第（一）项",
    "第十四条第（二）项",
    "第十四条第（三）项",
    "第十四条第（四）项",
    "第十四条第（五）项",
    "第十四条第（六）项",
    "第十四条第（七）项",
    "第十五条第（一）项",
    "第十五条第（二）项",
    "第十五条第（三）项",
]

# 拟用条例对应的法律要件（存入案件 JSON，供 AI 生成文书时突出关键证据要素）
REGULATION_ELEMENTS = {
    "第十四条第（一）项": ["工作时间", "工作场所", "因工作原因受到事故伤害"],
    "第十四条第（二）项": ["工作时间前后", "工作场所内", "从事与工作有关的预备性或收尾性工作", "受到事故伤害"],
    "第十四条第（三）项": ["工作时间", "工作场所内", "因履行工作职责", "受到暴力等意外伤害"],
    "第十四条第（四）项": ["患职业病（须符合国家职业病目录）"],
    "第十四条第（五）项": ["因工外出期间", "由于工作原因受到伤害 / 发生事故下落不明"],
    "第十四条第（六）项": ["上下班途中", "非本人主要责任", "交通事故（含轨道交通、客运轮渡、火车事故）"],
    "第十五条第（一）项": ["工作时间", "工作岗位", "突发疾病死亡 / 48小时内经抢救无效死亡"],
    "第十五条第（二）项": ["在抢险救灾等维护国家利益、公共利益活动中受到伤害"],
    "第十五条第（三）项": ["原在军队服役", "因战/因公负伤致残", "已取得革命伤残军人证", "到用人单位后旧伤复发"],
}


def _regulation_elements(short: str) -> list:
    """返回拟用条例对应的法律要件列表；未知条例（含第十四条第（七）项兜底）返回空列表"""
    if not short:
        return []
    return list(REGULATION_ELEMENTS.get(short.strip(), []))

_REG_CN_DIGITS = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}


def _cn_num_to_int(text: str):
    """中文数字转整数：一→1，六→6，十四→14，十五→15"""
    if not text:
        return None
    if text in _REG_CN_DIGITS:
        return _REG_CN_DIGITS[text]
    if text.startswith("十"):
        tail = text[1:]
        return 10 + (_REG_CN_DIGITS.get(tail, 0) if tail else 0)
    if "十" in text:
        tens, ones = text.split("十", 1)
        return _REG_CN_DIGITS.get(tens, 0) * 10 + (_REG_CN_DIGITS.get(ones, 0) if ones else 0)
    return None


_INT_TO_CN = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
              6: "六", 7: "七", 8: "八", 9: "九", 10: "十"}


def _int_to_cn_num(n) -> str:
    """整数转中文数字：1→一，6→六，14→十四，15→十五"""
    if not n:
        return ""
    if n <= 10:
        return _INT_TO_CN[n]
    if n < 20:
        tail = n - 10
        return "十" + (_INT_TO_CN[tail] if tail else "")
    tens = n // 10
    ones = n % 10
    return _INT_TO_CN[tens] + "十" + (_INT_TO_CN[ones] if ones else "")


def _regulation_full_to_short(text: str) -> str:
    """《工伤保险条例》第十四条第一款第一项 → 第十四条第（一）项"""
    import re
    if not text:
        return ""
    m = re.search(
        r"第([一二三四五六七八九十]+)条第[一二三四五六七八九十]+款第([一二三四五六七八九十]+)项",
        text,
    )
    if m:
        return f"第{m.group(1)}条第（{m.group(2)}）项"
    return text


def _regulation_short_to_full(short: str) -> str:
    """第十四条第（一）项 → 《工伤保险条例》第十四条第一款第一项"""
    import re
    if not short:
        return ""
    m = re.match(r"^第([一二三四五六七八九十]+)条第（([一二三四五六七八九十]+)）项$", short)
    if m:
        art = _cn_num_to_int(m.group(1))
        item = _cn_num_to_int(m.group(2))
        if art and item:
            return f"《工伤保险条例》第{_int_to_cn_num(art)}条第一款第{_int_to_cn_num(item)}项"
    return short


def _filter_provided(materials):
    """只保留「已提供（勾选）」的材料，转成 {name, notes} 格式"""
    if not isinstance(materials, list):
        return []
    return [
        {"name": m.get('name', ''), "notes": m.get('notes', '')}
        for m in materials
        if isinstance(m, dict) and m.get('provided') and m.get('name')
    ]


def _to_full_materials(provided_materials):
    """把 JSON 里的 {name, notes} 材料转回界面用的 {name, provided:True, notes}"""
    if not isinstance(provided_materials, list):
        return []
    return [
        {"name": m.get('name', ''), "provided": True, "notes": m.get('notes', '')}
        for m in provided_materials
        if isinstance(m, dict) and m.get('name')
    ]


# ============================================================================
# F2 测试数据预设（按 F2 轮换）
# ============================================================================

TEST_DATA_PRESETS = [
    # ═══════════════════════════════════════════════════════════════
    # 条例14条第(一)项 普通工伤 — 张三案（本人+证人+法人）
    # ═══════════════════════════════════════════════════════════════
    {
        "name": "1/18 张三案-本人(第一项)",
        "role": "本人",
        "name_pane": "张三",
        "idnumer_pane": "330324199003151234",
        "textEdit": "浙江省永嘉县瓯北街道XX路88号",
        "lineEdit_4": "13888880001",
        "lineEdit_5": "泥水工",
        "injured_worker": "张三",
        "comboBox": 0,
        "company_pane": "温州YY建筑劳务有限公司",
        "construction_company": "永嘉县XX建设工程有限公司",
        "construction_plant": "ZZ新城项目一期工地",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "我单位职工张三，男，1990年3月15日出生，身份证号330324199003151234，住永嘉县瓯北街道XX路88号。2026年7月20日下午16时20分许，张三在ZZ新城项目一期工地3号楼5层搬运水泥时，被滑落的水泥袋砸伤右脚，经永嘉县人民医院诊断为右足跖骨骨折。该事故属于在工作时间和工作场所内因工作原因受到的事故伤害，符合《工伤保险条例》第十四条第（一）项情形，现申请认定工伤。",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "医院诊断证明书", "provided": True, "notes": "永嘉县人民医院 右足跖骨骨折"},
            {"name": "劳动合同", "provided": True, "notes": ""},
            {"name": "工资银行流水", "provided": True, "notes": ""},
            {"name": "考勤记录", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "2/18 张三案-证人(李四)",
        "role": "证人",
        "name_pane": "李四",
        "idnumer_pane": "330324198508121235",
        "textEdit": "浙江省永嘉县桥头镇YY村123号",
        "lineEdit_4": "13966660002",
        "lineEdit_5": "钢筋工",
        "injured_worker": "张三",
        "comboBox": 0,
        "company_pane": "温州YY建筑劳务有限公司",
        "construction_company": "永嘉县XX建设工程有限公司",
        "construction_plant": "ZZ新城项目一期工地",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "证人证言", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "3/18 张三案-法人(王五)",
        "role": "法人",
        "name_pane": "王五",
        "idnumer_pane": "330324198212011234",
        "textEdit": "浙江省永嘉县上塘镇XX路66号",
        "lineEdit_4": "13777770003",
        "lineEdit_5": "法定代表人",
        "injured_worker": "张三",
        "comboBox": 0,
        "company_pane": "温州YY建筑劳务有限公司",
        "construction_company": "永嘉县XX建设工程有限公司",
        "construction_plant": "ZZ新城项目一期工地",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "公司营业执照副本", "provided": True, "notes": ""},
            {"name": "法定代表人身份证", "provided": True, "notes": ""},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },

    # ═══════════════════════════════════════════════════════════════
    # 条例14条第(二)项 工作前后预备性工作 — 刘十案（本人+证人+法人）
    # ═══════════════════════════════════════════════════════════════
    {
        "name": "4/18 刘十案-本人(第二项)",
        "role": "本人",
        "name_pane": "刘十",
        "idnumer_pane": "330324199207052345",
        "textEdit": "浙江省永嘉县大若岩镇XX村33号",
        "lineEdit_4": "13511112222",
        "lineEdit_5": "冲压工",
        "injured_worker": "刘十",
        "comboBox": 1,
        "company_pane": "温州AA金属制品有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "我单位职工刘十，系温州AA金属制品有限公司冲压工。2026年7月25日上午7时40分许（上班时间为8时），刘十按照惯例提前到岗对公司冲压设备进行例行检查和预热，在此过程中右手被机器夹伤，经温州附二医诊断为右手食指和中指挤压伤。该检查和预热工作系其职责范围内的预备性工作，符合《工伤保险条例》第十四条第（二）项情形，现申请认定工伤。",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "医院诊断证明", "provided": True, "notes": "右手食指和中指挤压伤"},
            {"name": "劳动合同", "provided": True, "notes": ""},
            {"name": "操作规程手册", "provided": True, "notes": "车间设备预热流程"},
            {"name": "考勤记录", "provided": True, "notes": "7月25日打卡7:35"},
        ],
    },
    {
        "name": "5/18 刘十案-证人(钱七)",
        "role": "证人",
        "name_pane": "钱七",
        "idnumer_pane": "330324199906081239",
        "textEdit": "浙江省永嘉县岩头镇ZZ村88号",
        "lineEdit_4": "13544440005",
        "lineEdit_5": "冲压工",
        "injured_worker": "刘十",
        "comboBox": 1,
        "company_pane": "温州AA金属制品有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "证人证言", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "6/18 刘十案-法人(王五)",
        "role": "法人",
        "name_pane": "王五",
        "idnumer_pane": "330324198212011234",
        "textEdit": "浙江省永嘉县上塘镇XX路66号",
        "lineEdit_4": "13777770003",
        "lineEdit_5": "法定代表人",
        "injured_worker": "刘十",
        "comboBox": 1,
        "company_pane": "温州AA金属制品有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "公司营业执照副本", "provided": True, "notes": ""},
            {"name": "法定代表人身份证", "provided": True, "notes": ""},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },

    # ═══════════════════════════════════════════════════════════════
    # 条例14条第(三)项 暴力伤害 — 周八案（本人+证人+法人）
    # ═══════════════════════════════════════════════════════════════
    {
        "name": "7/18 周八案-本人(第三项)",
        "role": "本人",
        "name_pane": "周八",
        "idnumer_pane": "330324199106011240",
        "textEdit": "浙江省永嘉县枫林镇XX村55号",
        "lineEdit_4": "13433330006",
        "lineEdit_5": "保安",
        "injured_worker": "周八",
        "comboBox": 2,
        "company_pane": "温州EE物业管理有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "我单位职工周八，系温州EE物业管理有限公司保安，派驻永嘉县XX商场担任安保工作。2026年8月1日晚21时许，周八在商场一楼巡逻时发现一名男子正在盗窃商户商品，上前制止时被对方用铁棍击打头部和右臂。商场其他保安闻讯赶来将嫌疑人控制并报警，周八被送至永嘉县人民医院治疗，诊断为脑震荡、右臂桡骨骨折。该伤害属于在履行工作职责过程中因履行工作职责受到暴力伤害，符合《工伤保险条例》第十四条第（三）项情形，现申请认定工伤。",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "医院诊断证明", "provided": True, "notes": "脑震荡、右臂桡骨骨折"},
            {"name": "公安局报案回执", "provided": True, "notes": "永嘉县公安局"},
            {"name": "商场监控录像", "provided": True, "notes": "XX商场一楼"},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "8/18 周八案-证人(李四)",
        "role": "证人",
        "name_pane": "李四",
        "idnumer_pane": "330324198508121235",
        "textEdit": "浙江省永嘉县桥头镇YY村123号",
        "lineEdit_4": "13966660002",
        "lineEdit_5": "保安",
        "injured_worker": "周八",
        "comboBox": 2,
        "company_pane": "温州EE物业管理有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "证人证言", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "9/18 周八案-法人(王五)",
        "role": "法人",
        "name_pane": "王五",
        "idnumer_pane": "330324198212011234",
        "textEdit": "浙江省永嘉县上塘镇XX路66号",
        "lineEdit_4": "13777770003",
        "lineEdit_5": "法定代表人",
        "injured_worker": "周八",
        "comboBox": 2,
        "company_pane": "温州EE物业管理有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "公司营业执照副本", "provided": True, "notes": ""},
            {"name": "法定代表人身份证", "provided": True, "notes": ""},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },

    # ═══════════════════════════════════════════════════════════════
    # 条例14条第(四)项 患职业病 — 孙八案（本人+证人+法人）
    # ═══════════════════════════════════════════════════════════════
    {
        "name": "10/18 孙八案-本人(第四项)",
        "role": "本人",
        "name_pane": "孙八",
        "idnumer_pane": "330324197506011245",
        "textEdit": "浙江省永嘉县巽宅镇XX村7号",
        "lineEdit_4": "13622223333",
        "lineEdit_5": "采掘工",
        "injured_worker": "孙八",
        "comboBox": 3,
        "company_pane": "温州DD矿业有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "我单位职工孙八，系温州DD矿业有限公司采掘工，自2015年起一直从事井下采掘作业，长期接触矽尘粉尘。2026年5月经温州市职业病防治院诊断，确认孙八患有职业性矽肺壹期。该疾病属于在职业活动中接触职业病危害因素所致，符合《工伤保险条例》第十四条第（四）项情形，现申请认定工伤。",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "职业病诊断证明书", "provided": True, "notes": "温州市职业病防治院 矽肺壹期"},
            {"name": "劳动合同", "provided": True, "notes": ""},
            {"name": "历年职业健康体检报告", "provided": True, "notes": ""},
            {"name": "工作场所粉尘检测报告", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "11/18 孙八案-证人(李四)",
        "role": "证人",
        "name_pane": "李四",
        "idnumer_pane": "330324198508121235",
        "textEdit": "浙江省永嘉县桥头镇YY村123号",
        "lineEdit_4": "13966660002",
        "lineEdit_5": "采掘工",
        "injured_worker": "孙八",
        "comboBox": 3,
        "company_pane": "温州DD矿业有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "证人证言", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "12/18 孙八案-法人(王五)",
        "role": "法人",
        "name_pane": "王五",
        "idnumer_pane": "330324198212011234",
        "textEdit": "浙江省永嘉县上塘镇XX路66号",
        "lineEdit_4": "13777770003",
        "lineEdit_5": "法定代表人",
        "injured_worker": "孙八",
        "comboBox": 3,
        "company_pane": "温州DD矿业有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "公司营业执照副本", "provided": True, "notes": ""},
            {"name": "法定代表人身份证", "provided": True, "notes": ""},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },

    # ═══════════════════════════════════════════════════════════════
    # 条例14条第(五)项 因工外出 — 钱七案（本人+证人+法人）
    # ═══════════════════════════════════════════════════════════════
    {
        "name": "13/18 钱七案-本人(第五项)",
        "role": "本人",
        "name_pane": "钱七",
        "idnumer_pane": "330324199906081239",
        "textEdit": "浙江省永嘉县岩头镇ZZ村88号",
        "lineEdit_4": "13544440005",
        "lineEdit_5": "技术员",
        "injured_worker": "钱七",
        "comboBox": 4,
        "company_pane": "永嘉县CC环保科技有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "我单位职工钱七，系永嘉县CC环保科技有限公司技术员。2026年7月28日，钱七受单位指派前往乐清市DD化工厂进行设备维护服务。上午9时30分许，其乘坐的公司车辆在乐清市柳市镇境内发生侧翻事故，造成钱七腰椎压缩性骨折，已送乐清市人民医院治疗。该事故属于因工外出期间由于工作原因受到的事故伤害，符合《工伤保险条例》第十四条第（五）项情形，现申请认定工伤。",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "公司派工单", "provided": True, "notes": "2026年7月27日签发"},
            {"name": "出差申请书", "provided": True, "notes": "经理王五审批"},
            {"name": "交通事故认定书", "provided": True, "notes": "单方事故 路面湿滑"},
            {"name": "医院诊断证明", "provided": True, "notes": "腰椎压缩性骨折"},
        ],
    },
    {
        "name": "14/18 钱七案-证人(李四)",
        "role": "证人",
        "name_pane": "李四",
        "idnumer_pane": "330324198508121235",
        "textEdit": "浙江省永嘉县桥头镇YY村123号",
        "lineEdit_4": "13966660002",
        "lineEdit_5": "技术员",
        "injured_worker": "钱七",
        "comboBox": 4,
        "company_pane": "永嘉县CC环保科技有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "证人证言", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "15/18 钱七案-法人(王五)",
        "role": "法人",
        "name_pane": "王五",
        "idnumer_pane": "330324198212011234",
        "textEdit": "浙江省永嘉县上塘镇XX路66号",
        "lineEdit_4": "13777770003",
        "lineEdit_5": "经理",
        "injured_worker": "钱七",
        "comboBox": 4,
        "company_pane": "永嘉县CC环保科技有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "公司营业执照副本", "provided": True, "notes": ""},
            {"name": "经理身份证", "provided": True, "notes": ""},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },

    # ═══════════════════════════════════════════════════════════════
    # 条例14条第(六)项 上下班途中 — 赵六案（本人+证人+法人）
    # ═══════════════════════════════════════════════════════════════
    {
        "name": "16/18 赵六案-本人(第六项)",
        "role": "本人",
        "name_pane": "赵六",
        "idnumer_pane": "330324199508031238",
        "textEdit": "浙江省永嘉县乌牛街道XX小区5栋301室",
        "lineEdit_4": "13655550004",
        "lineEdit_5": "装配工",
        "injured_worker": "赵六",
        "comboBox": 5,
        "company_pane": "温州BB电器有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "我单位职工赵六，系温州BB电器有限公司装配工，住永嘉县乌牛街道XX小区5栋301室。2026年8月5日下午17时40分许，赵六下班后骑电动车沿104国道从公司回乌牛街道家中，在104国道乌牛段被一辆小型轿车追尾，经永嘉县交警大队认定对方负全部责任。事故造成赵六左腿胫骨骨折，已送永嘉县人民医院治疗。该事故属于在上下班途中受到非本人主要责任的交通事故伤害，符合《工伤保险条例》第十四条第（六）项情形，现申请认定工伤。",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "道路交通事故认定书", "provided": True, "notes": "对方全责"},
            {"name": "医院诊断证明", "provided": True, "notes": "左腿胫骨骨折"},
            {"name": "劳动合同", "provided": True, "notes": ""},
            {"name": "路线示意图", "provided": True, "notes": "乌牛街道→104国道→公司"},
        ],
    },
    {
        "name": "17/18 赵六案-证人(李四)",
        "role": "证人",
        "name_pane": "李四",
        "idnumer_pane": "330324198508121235",
        "textEdit": "浙江省永嘉县桥头镇YY村123号",
        "lineEdit_4": "13966660002",
        "lineEdit_5": "装配工",
        "injured_worker": "赵六",
        "comboBox": 5,
        "company_pane": "温州BB电器有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "身份证复印件", "provided": True, "notes": ""},
            {"name": "证人证言", "provided": True, "notes": ""},
        ],
    },
    {
        "name": "18/18 赵六案-法人(王五)",
        "role": "法人",
        "name_pane": "王五",
        "idnumer_pane": "330324198212011234",
        "textEdit": "浙江省永嘉县上塘镇XX路66号",
        "lineEdit_4": "13777770003",
        "lineEdit_5": "法定代表人",
        "injured_worker": "赵六",
        "comboBox": 5,
        "company_pane": "温州BB电器有限公司",
        "construction_company": "",
        "construction_plant": "",
        "deathCaseCheckbox": False,
        "personalApplicationCheckbox": False,
        "statement_edit": "",
        "materials": [
            {"name": "公司营业执照副本", "provided": True, "notes": ""},
            {"name": "法定代表人身份证", "provided": True, "notes": ""},
            {"name": "劳动合同", "provided": True, "notes": ""},
        ],
    },
]


# ============================================================================
# AIWorker
# ============================================================================

class AIWorker(QThread):
    """AI工作线程"""
    finished = pyqtSignal(dict)  # 发送完成信号
    error = pyqtSignal(str)  # 发送错误信号
    progress = pyqtSignal(str, int)  # 发送进度信号 (消息, 进度百分比)

    def __init__(self, ai_service, file_path):
        super().__init__()
        self.ai_service = ai_service
        self.file_path = file_path

    def run(self):
        """线程运行的主函数"""
        try:
            # 第一步：提取文本
            self.progress.emit("正在提取文档文本...", 20)
            document_text = self.ai_service.extract_text_from_docx(self.file_path)

            # 第二步：AI分析
            self.progress.emit("正在调用DeepSeek API进行分析...", 50)
            result = self.ai_service.analyze_legal_document(document_text)

            # 第三步：完成
            self.progress.emit("分析完成，正在生成报告...", 90)
            self.finished.emit(result)

        except Exception as e:
            self.error.emit(str(e))

class RegulationAnalyzeWorker(QThread):
    """条例判断 + 证据分析 的 AI 工作线程"""
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, ai_service, case_obj):
        super().__init__()
        self.ai_service = ai_service
        self.case_obj = case_obj

    def run(self):
        try:
            result = self.ai_service.analyze_case_for_regulation(self.case_obj)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))

class TranscriptFromTemplateWorker(QThread):
    """把「发送给AI的模板」渲染全文发给 AI 生成询问笔录的线程"""
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, ai_service, full_text):
        super().__init__()
        self.ai_service = ai_service
        self.full_text = full_text

    def run(self):
        try:
            result = self.ai_service.generate_transcript_from_text(self.full_text)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))

class CaseDataModel:
    """案件数据模型 - 统一管理所有案件数据"""

    def __init__(self):
        self.basic_info: Dict[str, Any] = {}  # 基础个人信息
        self.company_info: Dict[str, Any] = {}  # 公司相关信息
        self.case_info: Dict[str, Any] = {}  # 案件信息
        self.investigation: Dict[str, Any] = {}  # 调查信息
        self.output_config: Dict[str, Any] = {}  # 输出配置
        self.witnesses: List[Dict[str, Any]] = []  # 多证人数据，每项含 序号/姓名/身份证号/身份证地址/手机号/岗位/性别/年龄
        self.current_witness_index: int = -1  # 当前正在编辑的证人下标，-1 表示无
        self._init_default_values()

    def _init_default_values(self):
        """初始化默认值"""
        self.case_info.update({
            '案件性质': '工伤案件',
            '申请类型': '单位申请'
        })
        self.output_config.update({
            '当前日期': _date_now(),
            '当前时间': _time_now()
        })

    def to_template_dict(self) -> Dict[str, Any]:
        """转换为模板渲染用的字典"""
        template_dict = {}

        # 确保日期时间是最新的
        self.output_config.update({
            '当前日期': _date_now(),
            '当前时间': _time_now()
        })

        # 按优先级合并
        template_dict.update(self.basic_info)
        template_dict.update(self.company_info)
        template_dict.update(self.case_info)
        template_dict.update(self.investigation)
        template_dict.update(self.output_config)

        return template_dict

    def update_basic_info(self, role: str, data: Dict[str, Any]):
        """更新基础信息"""
        prefixed_data = {}
        for key, value in data.items():
            if not key.startswith(role):
                new_key = f"{role}{key}" if key != "姓名" else f"{role}姓名"
            else:
                new_key = key
            prefixed_data[new_key] = value

        self.basic_info.update(prefixed_data)

    def update_company_info(self, company_name: str = "",
                            employer: str = "",
                            site: str = ""):
        """更新公司信息"""
        if company_name:
            self.company_info['用工单位'] = company_name
        if employer:
            self.company_info['用人单位'] = employer
        if site:
            self.company_info['工地名称'] = site

    def clear_role_data(self, role: str):
        """清除特定角色的数据"""
        role_prefix = role if role in ["本人", "证人", "法人"] else ""
        if not role_prefix:
            return

        keys_to_remove = [
            key for key in self.basic_info.keys()
            if key.startswith(role_prefix)
        ]

        for key in keys_to_remove:
            self.basic_info.pop(key, None)


# ============================================================================
# 关键证据定义（缺失时AI必须在笔录中追问）
# ============================================================================
KEY_EVIDENCE_NAMES = [
    "身份证",           # 身份证明
    "劳动合同",         # 劳动关系证明
    "医院诊断证明",     # 受伤事实证明
    "工资发放记录",     # 工资标准证明
    "考勤记录",         # 工作时间证明
    "证人证言",         # 事故见证
    "事故现场照片",     # 现场证据
    "监控录像",         # 现场证据
    "工伤认定书",       # 前置认定
    "劳动关系裁决书",   # 劳动关系确认
    "道路交通事故认定书",  # 上下班途中/因工外出
    "公安报案回执",     # 暴力伤害
    "死亡证明",         # 工亡
]


class MaterialListWidget(QWidget):
    """替代原有 QTextEdit 的材料管理组件。

    每行: [☑/☐ 复选框] [材料名称] [备注输入框]
    - 勾选 = 已提供
    - 未勾选 = 缺失，生成笔录时AI会追问
    - 备注 = 对该材料的补充说明
    """

    materials_changed = pyqtSignal()  # 材料变更信号

    def __init__(self, parent=None):
        super().__init__(parent)
        self._rows: List[Dict[str, Any]] = []  # [{name, provided, notes, widget_refs}]
        self._init_ui()

    def _init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(1)

        # 滚动区域
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #ccc;
                border-radius: 2px;
                background-color: #fafafa;
            }
        """)

        # 内容容器
        self._container = QWidget()
        self._container.setStyleSheet("background-color: #fafafa;")
        self._row_layout = QVBoxLayout(self._container)
        self._row_layout.setContentsMargins(4, 2, 4, 2)
        self._row_layout.setSpacing(2)
        self._row_layout.addStretch()  # 底部弹簧，把行推到顶部

        self.scroll.setWidget(self._container)
        layout.addWidget(self.scroll)

    def _make_row(self, name: str = "", provided: bool = False, notes: str = ""):
        """创建一行材料条目"""
        row = QWidget()
        row.setFixedHeight(23)
        h = QHBoxLayout(row)
        h.setContentsMargins(0, 0, 0, 0)
        h.setSpacing(3)

        # 复选框
        cb = QCheckBox()
        cb.setChecked(provided)
        cb.setFixedWidth(18)
        cb.setToolTip("勾选=已提供  |  不勾选=缺失")
        cb.toggled.connect(self._on_changed)

        # 材料名称输入框（可编辑）
        name_edit = QLineEdit(name if name else "")
        name_edit.setPlaceholderText("材料名称...")
        name_edit.setStyleSheet("""
            QLineEdit {
                font-size: 8pt;
                border: 1px solid #ddd;
                border-radius: 1px;
                padding: 1px 3px;
                background-color: #fff;
            }
            QLineEdit:focus {
                border-color: #3498db;
            }
        """)
        name_edit.setMinimumWidth(100)
        name_edit.textChanged.connect(self._on_changed)

        # 备注输入框
        note_edit = QLineEdit()
        note_edit.setText(notes)
        note_edit.setPlaceholderText("备注...")
        note_edit.setStyleSheet(name_edit.styleSheet())
        note_edit.textChanged.connect(self._on_changed)

        h.addWidget(cb)
        h.addWidget(name_edit, 1)   # stretch=1，自动填充剩余空间
        h.addWidget(note_edit, 2)   # stretch=2，备注更宽

        return row, cb, name_edit, note_edit

    def add_row(self, name: str = "", provided: bool = False, notes: str = ""):
        """在末尾添加一行"""
        row, cb, name_edit, note = self._make_row(name, provided, notes)
        # 在 stretch 之前插入
        self._row_layout.insertWidget(self._row_layout.count() - 1, row)

        item = {
            "name": name,
            "provided": provided,
            "notes": notes,
            "_cb": cb,
            "_name_edit": name_edit,
            "_note": note,
            "_row": row,
        }
        self._rows.append(item)

    def set_materials(self, data: List[Dict[str, Any]]):
        """批量设置材料列表"""
        self.clear()
        for item in data:
            self.add_row(
                name=item.get("name", ""),
                provided=item.get("provided", False),
                notes=item.get("notes", "")
            )

    def get_materials(self) -> List[Dict[str, Any]]:
        """获取所有材料数据（同步UI状态）"""
        result = []
        for row in self._rows:
            result.append({
                "name": row["_name_edit"].text(),
                "provided": row["_cb"].isChecked(),
                "notes": row["_note"].text(),
            })
        return result

    def get_provided(self) -> List[str]:
        """获取已提供的材料名称列表"""
        return [r["_name_edit"].text() for r in self._rows if r["_cb"].isChecked()]

    def get_missing(self) -> List[str]:
        """获取缺失的材料名称列表"""
        return [r["_name_edit"].text() for r in self._rows if not r["_cb"].isChecked()]

    def get_missing_key_evidence(self) -> List[str]:
        """获取缺失的关键证据列表"""
        missing = self.get_missing()
        return [m for m in missing if any(
            kw in m for kw in KEY_EVIDENCE_NAMES
        )]

    def clear(self):
        """清空所有行"""
        for row in self._rows:
            row["_row"].setParent(None)
        self._rows.clear()

    def get_summary_text(self) -> str:
        """生成可复制的文本摘要"""
        lines = []
        for i, r in enumerate(self.get_materials(), 1):
            status = "✓" if r["provided"] else "✗"
            line = f"{status} {i}. {r['name']}"
            if r["notes"]:
                line += f"（{r['notes']}）"
            lines.append(line)
        return "\n".join(lines)

    def copy_to_clipboard(self):
        """复制材料摘要到剪贴板"""
        text = self.get_summary_text()
        QApplication.clipboard().setText(text)

    def _on_changed(self):
        """复选框或备注变更时发出信号"""
        self.materials_changed.emit()


# ============================================================================
# CaseDataReviewDialog — 数据核对窗口（以 JSON 文本形式显示并可编辑）
# ============================================================================

class CaseDataReviewDialog(QDialog):
    """案件数据核对窗口。

    以 JSON 文本形式展示完整案件数据，用户可直接编辑；
    点击「保存并关闭」时解析 JSON（通过 get_case_obj() 读取），格式错误则提示且不关闭。
    """

    def __init__(self, case_obj: Dict[str, Any], parent=None):
        super().__init__(parent)
        self._case_obj: Optional[Dict[str, Any]] = None
        self._build_ui(case_obj)

    def _build_ui(self, case_obj):
        self.setWindowTitle("🔍 案件数据核对")
        self.resize(720, 800)
        self.setMinimumSize(640, 660)

        root = QVBoxLayout(self)

        title = QLabel("请核对并修改案件数据（JSON 格式），改完后点「保存并关闭」")
        title.setStyleSheet("font-size: 13px; font-weight: bold; padding: 4px;")
        root.addWidget(title)

        self.json_edit = QTextEdit()
        self.json_edit.setFont(QFont("Consolas", 10))
        self.json_edit.setPlainText(json.dumps(case_obj, ensure_ascii=False, indent=2))
        root.addWidget(self.json_edit, 1)

        btns = QHBoxLayout()
        btns.addStretch()

        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(self.reject)
        btns.addWidget(cancel_btn)

        save_btn = QPushButton("保存并关闭")
        save_btn.setStyleSheet(
            "QPushButton { background-color: #27ae60; color: white; font-weight: bold; "
            "padding: 6px 24px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #219150; }"
        )
        save_btn.clicked.connect(self._on_save)
        btns.addWidget(save_btn)

        root.addLayout(btns)

    def _on_save(self):
        text = self.json_edit.toPlainText().strip()
        try:
            obj = json.loads(text)
            if not isinstance(obj, dict):
                raise ValueError("JSON 顶层必须是对象 {…}")
            self._case_obj = obj
            self.accept()
        except Exception as e:
            QMessageBox.warning(self, "JSON 格式错误", f"无法解析 JSON：\n{str(e)}\n\n请修正后再保存。")

    def get_case_obj(self) -> Optional[Dict[str, Any]]:
        return self._case_obj


class ApprovalDecisionDialog(QDialog):
    """案件审批表 AI 分析结果对话框：认定工伤 / 不予认定工伤 / 保存"""

    def __init__(self, analysis: dict, parent=None):
        super().__init__(parent)
        self.choice = "保存"
        self._build_ui(analysis)

    def _build_ui(self, analysis: dict):
        self.setWindowTitle("🔍 AI 分析结果")
        self.resize(760, 560)
        self.setMinimumSize(640, 480)

        layout = QVBoxLayout(self)

        bias = analysis.get("偏向", "")
        bias_label = QLabel(f"AI 倾向：{bias}" if bias else "AI 倾向：未知")
        bias_label.setStyleSheet("font-size: 15px; font-weight: bold; padding: 4px;")
        layout.addWidget(bias_label)

        # 页签：综合分析 / 不予认定理由 / 诊断结论
        tab_widget = QTabWidget()

        tab1 = QWidget()
        t1 = QVBoxLayout(tab1)
        analysis_edit = QTextEdit()
        analysis_edit.setReadOnly(True)
        analysis_edit.setPlainText(analysis.get("分析", ""))
        t1.addWidget(analysis_edit)
        tab_widget.addTab(tab1, "综合分析")

        tab2 = QWidget()
        t2 = QVBoxLayout(tab2)
        reason_edit = QTextEdit()
        reason_edit.setReadOnly(True)
        reasons = analysis.get("关键理由", []) or []
        reason_edit.setPlainText("\n".join(f"• {r}" for r in reasons))
        t2.addWidget(reason_edit)
        tab_widget.addTab(tab2, "不予认定理由")

        tab3 = QWidget()
        t3 = QVBoxLayout(tab3)
        diag_edit = QTextEdit()
        diag_edit.setReadOnly(True)
        diag_edit.setPlainText(analysis.get("诊断结论", ""))
        t3.addWidget(diag_edit)
        tab_widget.addTab(tab3, "诊断结论")

        layout.addWidget(tab_widget, 1)

        # 三个按钮
        btn_row = QHBoxLayout()
        btn_row.addStretch()

        save_btn = QPushButton("保存")
        save_btn.clicked.connect(lambda: self._done("保存"))
        btn_row.addWidget(save_btn)

        no_btn = QPushButton("不予认定工伤")
        no_btn.setStyleSheet(
            "QPushButton { background-color: #e74c3c; color: white; font-weight: bold; "
            "padding: 6px 20px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #c0392b; }"
        )
        no_btn.clicked.connect(lambda: self._done("不予认定"))
        btn_row.addWidget(no_btn)

        yes_btn = QPushButton("认定工伤")
        yes_btn.setStyleSheet(
            "QPushButton { background-color: #27ae60; color: white; font-weight: bold; "
            "padding: 6px 20px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #219150; }"
        )
        yes_btn.clicked.connect(lambda: self._done("认定"))
        btn_row.addWidget(yes_btn)

        layout.addLayout(btn_row)

    def _done(self, choice: str):
        self.choice = choice
        self.accept()

    def get_choice(self) -> str:
        return self.choice


class MainWindow(QWidget, Ui_Form):

    def __init__(self, parent=None, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.setupUi(self)
        self._setup_radio_connections()

        self._test_data_index = -1  # F2 测试数据轮换索引
        self.current_case_id = ""  # 当前案件案本号（跨角色/跨步骤保持同案关联）

        # lineEdit_2 改为案本号显示
        self.label_10.setText("案本号：")
        self.lineEdit_2.setGeometry(80, 130, 180, 20)
        self.lineEdit_2.setPlaceholderText("输入本人姓名后自动生成")

        # 输入本人姓名后自动生成案本号
        self.name_pane.editingFinished.connect(self._on_name_pane_changed)

        # 创建用户管理器并整合API配置UI
        self.user_manager = UserManager()
        self._setup_api_config_ui()
        self._load_saved_user_config()
        self._setup_witness_ui()

        # 第一步：统一设置所有路径（必须在所有服务初始化之前）
        print("=" * 50)
        print("🚀 开始初始化 MainWindow")
        print("=" * 50)

        self._setup_paths()  # 统一使用 path_utils 设置路径

        # 第二步：初始化配置服务（但禁用其路径管理功能）
        self.config_service = ConfigService()
        # 禁用config_service的路径检查，避免干扰
        self.config_service.set('system.startup_check_disk_space', False)
        self.config_service.set('template.base_path', self.TEMPLATE_PATH)

        # 第三步：初始化组合框数据（必须在路径设置之后）
        self.init_combobox_data()

        # 工伤告知书按钮连接
        self.pushButton_7.clicked.connect(self.generate_injury_notice)

        # 第四步：初始化其他核心服务（使用正确的路径）
        self.file_service = FileService(self.BASE_PATH)
        self.data_service = DataService()

        # 第五步：更新所有服务的路径
        self._update_services_paths()

        # 第六步：初始化数据模型
        self.data_model = CaseDataModel()
        self.var_manager = TemplateVariableManager(self.data_model)
        self.data_model.output_config['用户名'] = self._get_current_username()

        # 简化日志系统
        self.setup_logging()

        # 保持向后兼容
        self._template_dict = self.data_model.to_template_dict()
        self.current_case_folder = None
        self.current_person_name = ""
        # self.case_versions = {}  # 注释掉，如果不再使用

        # 获取复选框控件
        self.death_case_checkbox = self.findChild(QCheckBox, "deathCaseCheckbox")
        self.personal_application_checkbox = self.findChild(QCheckBox, "personalApplicationCheckbox")

        # 连接信号
        self.death_case_checkbox.stateChanged.connect(self.on_case_type_changed)
        self.personal_application_checkbox.stateChanged.connect(self.on_case_type_changed)

        # 初始化数据
        case_config = self.config_service.get_case_config()
        self.set_data('案件性质', case_config.default_case_type, 'case')
        self.set_data('申请类型', case_config.default_application_type, 'case')

        # 连接公司相关信号
        self.company_pane.currentTextChanged.connect(self.company)
        self.construction_company.currentTextChanged.connect(self.sync_employer_to_dict)
        self.construction_plant.currentTextChanged.connect(self.c_plant)

        # 初始化案件类型下拉框
        self.comboBox.addItem("《工伤保险条例》第十四条第一款第一项")
        self.comboBox.addItem("《工伤保险条例》第十四条第一款第二项")
        self.comboBox.addItem("《工伤保险条例》第十四条第一款第三项")
        self.comboBox.addItem("《工伤保险条例》第十四条第一款第四项")
        self.comboBox.addItem("《工伤保险条例》第十四条第一款第五项")
        self.comboBox.addItem("《工伤保险条例》第十四条第一款第六项")

        # 初始化组合框（必须在 init_combobox_data 之后）
        self.init_comboboxes()

        # 应用UI设置
        self._apply_ui_settings()

        # 搜索按钮
        self.pushButton_6.clicked.connect(self.smart_search_cases)

        # 初始化AI服务（使用传入的api_config）
        self.ai_service = None  # 先初始化为None
        self.init_ai_service()

        # 连接AI审查按钮
        self.pushButton_ai_review.clicked.connect(self.ai_review_document)
        self._reconnect_approval_button()
        # 谈话通知书按钮连接
        self.pushButton_12.clicked.connect(self.on_pushButton_12_clicked)

        try:
            self.pushButton.clicked.disconnect()
        except TypeError:
            pass
        self.pushButton.clicked.connect(self.on_talk_button_clicked)

        # 初始状态设为可用
        self.pushButton.setEnabled(True)

        # 验证模板路径（使用已获取的路径）
        if os.path.exists(self.TEMPLATE_PATH):
            print(f"✅ 模板路径可访问: {self.TEMPLATE_PATH}")
            print(f"✅ 谈话模板路径: {self.TALK_TEMPLATE_PATH}")
            print(f"✅ 文书模板路径: {self.DOCUMENT_TEMPLATE_PATH}")
        else:
            print(f"❌ 模板路径不存在: {self.TEMPLATE_PATH}")

        print("=" * 50)
        print("🎉 MainWindow 初始化完成")
        print("=" * 50)

    def on_talk_button_clicked(self):
        """谈话笔录按钮点击事件处理 — 数据核对 + 按角色生成笔录"""
        try:
            # 证人：先同步表单证人数据到数据模型（open_data_review 会回填本人数据，避免覆盖证人）
            if self.get_current_role_type() == "证人":
                self._sync_form_to_current_witness()

            # ── 第一步：弹出数据核对窗口，逐项核对并允许修改 ──
            if not self.open_data_review():
                self._set_status('已取消', 'black')
                return

            # ── 第二步：按角色生成 ──
            role = self.get_current_role_type()
            if role == "证人":
                # 数据核对会把表单回填成本人数据，这里把表单切回当前证人显示
                self._sync_current_witness_to_form()
                # 证人：跳过条例分析，直接生成证人谈话笔录
                self._generate_witness_transcript()
            elif role == "法人":
                # 数据核对会把表单回填成本人数据，这里把表单切回法人显示
                self._sync_legal_to_form()
                # 法人：跳过条例分析，直接生成法人谈话笔录
                self._generate_legal_transcript()
            else:
                # 本人：AI 条例判断 + 证据分析
                self._analyze_case_with_ai(self.current_case_id)
        except Exception as e:
            print(f"谈话笔录按钮点击异常: {e}")
            import traceback
            traceback.print_exc()

    # ========================================================================
    # 数据核对（第一步）相关方法
    # ========================================================================

    def open_data_review(self) -> bool:
        """弹出案件数据核对窗口（JSON 文本形式）。

        点击时先自动生成案本号填入 case_id；用户可编辑 JSON；
        保存时解析、落盘并回写主界面。
        """
        data, materials, _, _ = self._collect_review_data()

        # ── 点击即自动生成案本号，填入 case_id ──
        case_id = str(data.get('case_id', '')).strip()
        if not case_id:
            case_id = self._auto_generate_case_number(
                data.get('name', ''), data.get('id_card', '')
            )
            data['case_id'] = case_id
        self.current_case_id = case_id

        case_obj = self._build_case_object(data, materials)

        dlg = CaseDataReviewDialog(case_obj, self)
        if dlg.exec_() != QDialog.Accepted:
            self._set_status('已取消数据核对', 'black')
            return False

        case_obj = dlg.get_case_obj()
        if not case_obj:
            return False

        # 以最终 JSON 里的 case_id 为准（用户可能在文本框里改过）
        case_id = str(case_obj.get('case_id', '')).strip() or case_id
        case_obj['case_id'] = case_id
        self.current_case_id = case_id

        # ── 回写主界面与数据模型 ──
        self._apply_case_object(case_obj)

        # ── 保存到 cases_data.json ──
        cases = self._load_cases_data()
        cases[case_id] = case_obj
        self._save_cases_data(cases)

        self._set_status(f'数据已核对并保存（案本号：{case_id}）', 'green')
        print(f"✅ 数据核对完成并保存到 cases_data.json（案本号：{case_id}）")
        return True

    def _collect_review_data(self):
        """从主界面与数据模型收集当前案件的全部字段（供核对窗口展示）"""
        # 本人字段以数据模型优先，避免证人/法人切换后 name_pane 串数据
        regulation_full = self.comboBox.currentText().strip()
        regulation_short = self.get_data('拟用条例', '') or _regulation_full_to_short(regulation_full)

        data = {
            'case_id': self.lineEdit_2.text().strip() or self.get_data('案本号', '') or self.current_case_id,
            'case_nature': '工亡案件' if self.death_case_checkbox.isChecked() else '工伤案件',
            'applicant_type': '个人申请' if self.personal_application_checkbox.isChecked() else '单位申请',
            'regulation': regulation_short,
            'apply_time': self._resolve_date_input(self.apply_time_edit.text()) if hasattr(self, 'apply_time_edit') else '',
            'accept_time': self._resolve_date_input(self.accept_time_edit.text()) if hasattr(self, 'accept_time_edit') else '',
            'visit_time': self._resolve_date_input(self.visit_time_edit.text()) if hasattr(self, 'visit_time_edit') else '',
            'name': self.get_data('本人姓名', '') or self.name_pane.text().strip(),
            'gender': self.get_data('本人性别', '') or self.lineEdit.text().strip(),
            'age': str(self.get_data('本人年龄', '') or self.age_pane.text().strip()),
            'id_card': self.get_data('本人身份证号', '') or self.idnumer_pane.text().strip(),
            'phone': self.get_data('本人手机号', '') or self.lineEdit_4.text().strip(),
            'address': self.get_data('本人身份证地址', '') or self.textEdit.toPlainText().strip(),
            'position': self.get_data('本人岗位', '') or self.lineEdit_5.text().strip(),
            'employer': self.get_data('用工单位', '') or self.construction_company.currentText().strip(),
            'labor_unit': self.get_data('用人单位', '') or self.company_pane.currentText().strip(),
            'site': self.get_data('工地名称', '') or self.construction_plant.currentText().strip(),
            'injury_desc': self.statement_edit.toPlainText().strip() if hasattr(self, 'statement_edit')
                           else self.get_data('受伤经过', ''),
        }
        materials = (self.material_list.get_materials() if hasattr(self, 'material_list') else []) \
            or self.data_model.investigation.get('本人材料', [])
        witnesses = list(self.data_model.witnesses)
        legal_reps = self._collect_legal_reps()
        return data, materials, witnesses, legal_reps

    def _collect_legal_reps(self) -> List[Dict[str, Any]]:
        """收集法人信息（当前系统法人是单条，通过 法人* 键存储，兼容未来多条）"""
        name = self.get_data('法人姓名', '')
        if not name:
            return []
        return [{
            'name': name,
            'position': self.get_data('法人职务', ''),
            'id_card': self.get_data('法人身份证号', ''),
            'address': self.get_data('法人身份证地址', ''),
            'phone': self.get_data('法人手机号', ''),
            'materials': self.data_model.investigation.get('法人材料', []),
        }]

    def _set_combo_or_type(self, combobox, text):
        """设置下拉框的值：存在则选中，否则输入（可编辑）或新增项（不可编辑）"""
        text = text or ""
        if not text:
            combobox.setCurrentIndex(-1)
            return
        idx = combobox.findText(text)
        if idx >= 0:
            combobox.setCurrentIndex(idx)
        elif combobox.isEditable():
            combobox.setEditText(text)
        else:
            combobox.addItem(text)
            combobox.setCurrentIndex(combobox.count() - 1)

    def _apply_regulation(self, short: str):
        """把核对后的「拟用条例」写回下拉框与数据模型"""
        short = (short or "").strip()
        self.set_data('拟用条例', short, 'case')
        full = _regulation_short_to_full(short)
        if full:
            self.set_data('引用条例', full, 'case')
            self._set_combo_or_type(self.comboBox, full)

    def _apply_case_object(self, case_obj: Dict[str, Any]):
        """把核对后的案件 JSON 对象回写到主界面控件与数据模型"""
        # 案件基本信息
        self.lineEdit_2.setText(str(case_obj.get('case_id', '')))
        self.set_data('案本号', case_obj.get('case_id', ''), 'case')

        self.death_case_checkbox.setChecked(case_obj.get('case_nature', '') == '工亡案件')
        self.personal_application_checkbox.setChecked(case_obj.get('applicant_type', '') == '个人申请')
        self.on_case_type_changed()

        self._apply_regulation(case_obj.get('proposed_article', ''))

        if hasattr(self, 'apply_time_edit'):
            self.apply_time_edit.setText(str(case_obj.get('apply_time', '')))
        if hasattr(self, 'accept_time_edit'):
            self.accept_time_edit.setText(str(case_obj.get('accept_time', '')))
        if hasattr(self, 'visit_time_edit'):
            self.visit_time_edit.setText(str(case_obj.get('visit_time', '')))
        self._save_date_inputs()

        # 本人信息
        self.name_pane.setText(str(case_obj.get('name', '')))
        self.set_data('本人姓名', case_obj.get('name', ''), 'basic')
        self.lineEdit.setText(str(case_obj.get('gender', '')))
        self.set_data('本人性别', case_obj.get('gender', ''), 'basic')
        self.age_pane.setText(str(case_obj.get('age', '')))
        self.set_data('本人年龄', case_obj.get('age', ''), 'basic')
        self.idnumer_pane.setText(str(case_obj.get('id_card', '')))
        self.set_data('本人身份证号', case_obj.get('id_card', ''), 'basic')
        self.lineEdit_4.setText(str(case_obj.get('phone', '')))
        self.set_data('本人手机号', case_obj.get('phone', ''), 'basic')
        self.lineEdit_5.setText(str(case_obj.get('position', '')))
        self.set_data('本人岗位', case_obj.get('position', ''), 'basic')
        self.textEdit.setPlainText(str(case_obj.get('address', '')))
        self.set_data('本人身份证地址', case_obj.get('address', ''), 'basic')

        # 单位信息（company_pane=用人单位，construction_company=用工单位）
        self._set_combo_or_type(self.company_pane, case_obj.get('labor_unit', ''))
        self.set_data('用人单位', case_obj.get('labor_unit', ''), 'company')
        self._set_combo_or_type(self.construction_company, case_obj.get('employer', ''))
        self.set_data('用工单位', case_obj.get('employer', ''), 'company')
        self._set_combo_or_type(self.construction_plant, case_obj.get('site', ''))
        self.set_data('工地名称', case_obj.get('site', ''), 'company')

        # 受伤经过
        if hasattr(self, 'statement_edit'):
            self.statement_edit.setPlainText(str(case_obj.get('injury_description', '')))
        self.set_data('受伤经过', case_obj.get('injury_description', ''), 'investigation')

        # 材料清单（本人）—— JSON 里只存了「已提供」，转回界面格式
        materials_full = _to_full_materials(case_obj.get('materials', []))
        if hasattr(self, 'material_list'):
            self.material_list.set_materials(materials_full)
        self.data_model.investigation['本人材料'] = materials_full

        # 法人信息回写
        legal_reps = case_obj.get('legal_reps', [])
        if legal_reps:
            lr = legal_reps[0]
            self.set_data('法人姓名', lr.get('name', ''), 'basic')
            self.set_data('法人职务', lr.get('position', ''), 'basic')
            self.set_data('法人身份证号', lr.get('id_card', ''), 'basic')
            self.set_data('法人身份证地址', lr.get('address', ''), 'basic')
            self.set_data('法人手机号', lr.get('phone', ''), 'basic')
            self.data_model.investigation['法人材料'] = lr.get('materials', [])

        # 刷新模板字典缓存
        self._template_dict = self.data_model.to_template_dict()
        print("✅ 核对数据已回写主界面与数据模型")

    # ========================================================================
    # 案件 JSON 持久化（案本号为键）
    # ========================================================================

    def _cases_data_path(self) -> str:
        """案件数据 JSON 文件路径"""
        return os.path.join(self.BASE_PATH, 'cases_data.json')

    def _load_cases_data(self) -> Dict[str, Any]:
        """加载全部案件数据，返回 {case_id: case_obj}"""
        path = self._cases_data_path()
        if not os.path.exists(path):
            return {}
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            if isinstance(data, dict) and isinstance(data.get('cases'), dict):
                return data['cases']
        except Exception as e:
            print(f"⚠️ 加载案件数据失败: {e}")
        return {}

    def _save_cases_data(self, cases: Dict[str, Any]) -> bool:
        """保存全部案件数据到 cases_data.json"""
        path = self._cases_data_path()
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                json.dump({"version": "2.0", "cases": cases}, f, ensure_ascii=False, indent=2)
            print(f"✅ 案件数据已保存: {path}（共 {len(cases)} 个案件）")
            return True
        except Exception as e:
            print(f"❌ 保存案件数据失败: {e}")
            return False

    def _update_case_field(self, case_number: str, **fields) -> bool:
        """把字段写回 cases_data.json 的指定案件"""
        try:
            cases = self._load_cases_data()
            case_obj = cases.get(case_number)
            if case_obj is None:
                return False
            case_obj.update(fields)
            self._save_cases_data(cases)
            return True
        except Exception as e:
            print(f"⚠️ 更新案件字段失败（非致命）: {e}")
            return False

    def _build_unified_template_data(self, case_obj: Dict[str, Any]) -> Dict[str, Any]:
        """构建统一的模板渲染字典。

        中文 key 为主（模板占位符统一用中文），同时附带英文 case_obj 字段名 key，
        模板里写中文或英文占位符都能被替换。
        """
        elements = case_obj.get('proposed_article_elements', []) or []
        materials = case_obj.get('materials', []) or []
        material_names = [m.get('name', '') for m in materials
                          if isinstance(m, dict) and m.get('name')]

        # 统一中文占位符
        zh = {
            '案本号': case_obj.get('case_id', ''),
            '案件性质': case_obj.get('case_nature', ''),
            '申请类型': case_obj.get('applicant_type', ''),
            '用工单位': case_obj.get('employer', ''),
            '用人单位': case_obj.get('labor_unit', ''),
            '工地名称': case_obj.get('site', ''),
            '申请时间': case_obj.get('apply_time', ''),
            '受理时间': case_obj.get('accept_time', ''),
            '拟用条例': case_obj.get('proposed_article', ''),
            '法律要件': ' + '.join(elements) if elements else '',
            '本人姓名': case_obj.get('name', ''),
            '本人性别': case_obj.get('gender', ''),
            '本人年龄': case_obj.get('age', ''),
            '本人身份证号': case_obj.get('id_card', ''),
            '本人手机号': case_obj.get('phone', ''),
            '本人身份证地址': case_obj.get('address', ''),
            '本人岗位': case_obj.get('position', ''),
            '受伤经过': case_obj.get('injury_description', ''),
            '已提供材料': '、'.join(material_names) if material_names else '',
            '记录人': case_obj.get('recorder', '') or self._get_current_username(),
            '申请人名称': case_obj.get('applicant_name', ''),
            '用户名': self._get_current_username(),
            '当前时期': self.get_data('当前时期', '') or (_date_now() + _time_now()),
        }

        # 英文 key（case_obj 字段名，兼容写法）
        en = {
            'case_id': case_obj.get('case_id', ''),
            'case_nature': case_obj.get('case_nature', ''),
            'applicant_type': case_obj.get('applicant_type', ''),
            'employer': case_obj.get('employer', ''),
            'labor_unit': case_obj.get('labor_unit', ''),
            'site': case_obj.get('site', ''),
            'apply_time': case_obj.get('apply_time', ''),
            'accept_time': case_obj.get('accept_time', ''),
            'proposed_article': case_obj.get('proposed_article', ''),
            'proposed_article_elements': ' + '.join(elements) if elements else '',
            'name': case_obj.get('name', ''),
            'gender': case_obj.get('gender', ''),
            'age': case_obj.get('age', ''),
            'id_card': case_obj.get('id_card', ''),
            'phone': case_obj.get('phone', ''),
            'address': case_obj.get('address', ''),
            'position': case_obj.get('position', ''),
            'injury_description': case_obj.get('injury_description', ''),
            'materials': '、'.join(material_names) if material_names else '',
            'recorder': case_obj.get('recorder', '') or self._get_current_username(),
            'applicant_name': case_obj.get('applicant_name', ''),
        }

        merged = dict(zh)
        merged.update(en)
        return merged

    def _build_case_object(self, data, materials) -> Dict[str, Any]:
        """构建单个案件对象（case_id 为第一字段）

        - 本人数据平铺在顶层，injury_description 仅本人
        - materials 只保留「已提供（勾选）」的证据
        - 含记录人、申请人名称、证人（并入 cases_data.json，单一数据源）
        """
        case = {
            "case_id": data.get('case_id', ''),
            "applicant_name": data.get('name', '') if data.get('applicant_type', '') == '个人申请' else data.get('labor_unit', ''),
            "case_nature": data.get('case_nature', ''),
            "applicant_type": data.get('applicant_type', ''),
            "employer": data.get('employer', ''),
            "labor_unit": data.get('labor_unit', ''),
            "site": data.get('site', ''),
            "apply_time": data.get('apply_time', ''),
            "accept_time": data.get('accept_time', ''),
            "visit_time": data.get('visit_time', ''),
            "proposed_article": data.get('regulation', ''),
            "proposed_article_elements": _regulation_elements(data.get('regulation', '')),
            # ── 本人（一套完整数据）──
            "name": data.get('name', ''),
            "gender": data.get('gender', ''),
            "age": data.get('age', ''),
            "id_card": data.get('id_card', ''),
            "phone": data.get('phone', ''),
            "address": data.get('address', ''),
            "position": data.get('position', ''),
            "injury_description": data.get('injury_desc', ''),
            "materials": _filter_provided(materials),
            # ── 记录人 / 证人 ──
            "recorder": self._get_current_username(),
            "witnesses": list(self.data_model.witnesses),
            "legal_reps": self._collect_legal_reps(),
        }
        return case

    # ========================================================================
    # AI 条例判断 + 证据分析
    # ========================================================================

    def _analyze_case_with_ai(self, case_id: str):
        """对已保存的案件进行 AI 条例与证据分析"""
        if not case_id:
            self._set_status('无案本号，无法分析', 'orange')
            return
        if not self.ai_service:
            self._set_status('未配置AI，无法分析', 'orange')
            QMessageBox.warning(self, "提示", "未配置API密钥，无法进行AI条例分析。\n请在顶部⚙配置中设置API密钥。")
            return
        case_obj = self._load_cases_data().get(case_id)
        if not case_obj:
            QMessageBox.warning(self, "提示", "未找到该案本号的案件数据")
            return

        self._set_status('正在AI分析条例与证据...', 'black')
        QApplication.processEvents()

        self.analysis_worker = RegulationAnalyzeWorker(self.ai_service, case_obj)
        self.analysis_worker.finished.connect(
            lambda result: self._on_analysis_finished(case_id, result)
        )
        self.analysis_worker.error.connect(self._on_analysis_error)
        self.analysis_worker.start()

    def _on_analysis_finished(self, case_id: str, result: dict):
        self._set_status('AI分析完成', 'green')
        if '错误' in result:
            QMessageBox.warning(self, "AI分析失败", result.get('错误', '未知错误'))
            return
        # AI 分析完成后，用案件数据渲染「发送给AI的模板」并打开给用户查看
        self._generate_ai_template_and_open(case_id)
        self._show_regulation_analysis(case_id, result)

    def _generate_ai_template_and_open(self, case_id: str):
        """AI 条例与证据分析完成后，内存渲染「发送给AI的模板」→ 把全文直接发给 AI 生成询问笔录（不再生成/打开 docx）"""
        try:
            case_obj = self._load_cases_data().get(case_id)
            if not case_obj:
                print(f"⚠️ 生成「发送给AI」全文失败：未找到案件数据（{case_id}）")
                return

            # ── 1. 构建模板数据（统一中文占位符，附带英文 key）──
            template_data = self._build_unified_template_data(case_obj)

            # ── 2. 获取模板路径 ──
            template_path = str(path_utils.get_document_template_path('发送给AI的模板.docx'))
            if not os.path.exists(template_path):
                print(f"⚠️ 模板文件不存在: {template_path}")
                return

            # ── 3. 渲染模板（内存中，不落盘）──
            from docxtpl import DocxTemplate
            word = DocxTemplate(template_path)
            word.render(template_data)

            # ── 4. 从内存提取渲染后的全文（不再生成/打开「发送给AI」docx）──
            import io as _io
            buf = _io.BytesIO()
            word.save(buf)
            buf.seek(0)
            from docx import Document
            rendered = Document(buf)
            full_text = "\n".join(p.text for p in rendered.paragraphs if p.text.strip())
            print(f"📝 提取模板全文 {len(full_text)} 字符，准备发送给 AI")

            # ── 5. 把全文发给 AI 生成询问笔录（后台线程）──
            self._start_transcript_generation(case_id, case_obj, full_text)

        except Exception as e:
            print(f"❌ 生成「发送给AI」全文异常: {e}")
            import traceback
            traceback.print_exc()

    def _start_transcript_generation(self, case_id: str, case_obj: dict, full_text: str):
        """启动后台线程，把模板全文发给 AI 生成询问笔录"""
        self._set_status('正在AI生成询问笔录...', 'black')
        QApplication.processEvents()
        self.transcript_worker = TranscriptFromTemplateWorker(self.ai_service, full_text)
        self.transcript_worker.finished.connect(
            lambda result, cid=case_id, co=case_obj: self._on_transcript_generated(cid, co, result)
        )
        self.transcript_worker.error.connect(self._on_transcript_error)
        self.transcript_worker.start()

    def _on_transcript_generated(self, case_id: str, case_obj: dict, result: dict):
        if result.get("状态") != "成功":
            err = result.get("错误信息", "未知错误")
            print(f"⚠️ 询问笔录生成失败: {err}")
            self._set_status(f'询问笔录生成失败: {err[:40]}', 'orange')
            return
        content = (result.get("内容", "") or "").strip()
        if not content:
            print("⚠️ 询问笔录生成失败：AI 返回内容为空")
            self._set_status('询问笔录生成失败：AI 返回内容为空', 'orange')
            return
        path = self._save_transcript_to_template(case_obj, content)
        if not path:
            return
        success, message = self.file_service.open_document(path)
        if success:
            self._set_status('已生成并打开本人谈话笔录', 'green')
        else:
            self._set_status(f'本人谈话笔录已生成，打开失败: {message}', 'orange')

    def _on_transcript_error(self, err: str):
        print(f"❌ 询问笔录生成出错: {err}")
        self._set_status('询问笔录生成出错', 'red')

    def _save_transcript_to_template(self, case_obj: dict, content: str) -> str:
        """渲染「本人谈话笔录（普通工伤案件）.docx」模板，把 AI 生成的问答内容（去标题）插入到告知程序之后，返回文件路径（失败返回空串）"""
        try:
            from docx.shared import Pt

            # ── 1. 渲染谈话模板（替换占位符）──
            template_path = str(path_utils.get_talk_template_path('本人谈话笔录（普通工伤案件）.docx'))
            if not os.path.exists(template_path):
                print(f"⚠️ 谈话模板不存在: {template_path}")
                return ""

            template_data = self._build_unified_template_data(case_obj)

            # ── 1.1 用 docxtpl 渲染占位符（保留占位符原有格式）──
            from docxtpl import DocxTemplate
            doc = DocxTemplate(template_path)
            doc.render(template_data)

            # ── 2. 定位「答：听清楚了，不申请回避。」锚点 ──
            anchor_index = None
            anchor_pf = None
            for i, p in enumerate(doc.paragraphs):
                if '答：听清楚了，不申请回避' in p.text:
                    anchor_index = i
                    anchor_pf = p.paragraph_format
                    break

            if anchor_index is None:
                print("⚠️ 未找到锚点「答：听清楚了，不申请回避」")
                return ""

            # ── 3. 删除锚点之后的模板样例问答（保留头部+告知，避免与AI问答重复）──
            body = doc.element.body
            anchor_elem = doc.paragraphs[anchor_index]._element
            after_anchor = False
            for child in list(body):
                if after_anchor:
                    body.remove(child)
                elif child is anchor_elem:
                    after_anchor = True

            # ── 4. 在锚点之后插入 AI 问答（每行下划线）──
            for line in content.splitlines():
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph()
                if anchor_pf is not None:
                    p.paragraph_format.alignment = anchor_pf.alignment
                    p.paragraph_format.first_line_indent = anchor_pf.first_line_indent
                    p.paragraph_format.space_before = anchor_pf.space_before
                    p.paragraph_format.space_after = anchor_pf.space_after
                run = p.add_run(line)
                run.font.size = Pt(15)
                run.underline = True

            # ── 5. 保存 ──
            if not self.current_case_folder or not os.path.exists(self.current_case_folder):
                folder_subject = str(case_obj.get('case_id', '') or case_obj.get('name', '') or '案件').strip()
                self.current_case_folder = os.path.join(self.BASE_PATH, folder_subject)
                os.makedirs(self.current_case_folder, exist_ok=True)

            subject = str(case_obj.get('name', '') or case_obj.get('case_id', '') or '案件').strip()
            file_name = f"{subject}本人谈话笔录.docx"
            target_path = os.path.join(self.current_case_folder, file_name)
            counter = 2
            while os.path.exists(target_path):
                file_name = f"{subject}本人谈话笔录({counter}).docx"
                target_path = os.path.join(self.current_case_folder, file_name)
                counter += 1

            doc.save(target_path)
            print(f"✅ 本人谈话笔录已生成: {target_path}")
            return target_path
        except Exception as e:
            print(f"❌ 生成本人谈话笔录失败: {e}")
            import traceback
            traceback.print_exc()
            self._set_status('生成本人谈话笔录失败', 'red')
            return ""

    # ========================================================================
    # 证人谈话笔录生成（仿本人流程，跳过条例分析）
    # ========================================================================

    def _build_witness_ai_prompt(self, case_obj: dict) -> str:
        """构建证人谈话笔录的 AI 提示词（案件数据 + 当前证人 替换占位符）"""
        from prompt_manager import load_prompt
        self._ensure_current_witness()
        # 注意：不调用 _sync_form_to_current_witness（open_data_review 已把表单回填成本人数据，会污染证人）
        w = self._current_witness() or {}

        prompt = load_prompt('witness_send_to_ai')
        replacements = {
            '案本号': case_obj.get('case_id', ''),
            '案件性质': case_obj.get('case_nature', ''),
            '申请类型': case_obj.get('applicant_type', ''),
            '本人姓名': case_obj.get('name', ''),
            '本人性别': case_obj.get('gender', ''),
            '本人身份证号': case_obj.get('id_card', ''),
            '用人单位': case_obj.get('labor_unit', ''),
            '用工单位': case_obj.get('employer', ''),
            '工地名称': case_obj.get('site', ''),
            '受伤经过': case_obj.get('injury_description', ''),
            '证人姓名': w.get('姓名', '') or self.get_data('证人姓名', ''),
            '证人身份证号': w.get('身份证号', '') or self.get_data('证人身份证号', ''),
            '证人岗位': w.get('岗位', '') or self.get_data('证人岗位', ''),
        }
        for key, val in replacements.items():
            prompt = prompt.replace('{{%s}}' % key, str(val))
        return prompt

    def _build_witness_template_data(self, case_obj: dict) -> dict:
        """构建证人谈话笔录模板的占位符数据"""
        self._ensure_current_witness()
        # 注意：不调用 _sync_form_to_current_witness（open_data_review 已把表单回填成本人数据，会污染证人）
        w = self._current_witness() or {}
        return {
            '当前时期': self.get_data('当前时期', '') or (_date_now() + _time_now()),
            '用户名': self._get_current_username(),
            '本人姓名': case_obj.get('name', ''),
            '证人姓名': w.get('姓名', '') or self.get_data('证人姓名', ''),
            '证人性别': w.get('性别', '') or self.get_data('证人性别', ''),
            '证人年龄': w.get('年龄', '') or self.get_data('证人年龄', ''),
            '证人身份证号': w.get('身份证号', '') or self.get_data('证人身份证号', ''),
            '证人身份证地址': w.get('身份证地址', '') or self.get_data('证人身份证地址', ''),
            '证人手机号': w.get('手机号', '') or self.get_data('证人手机号', ''),
            '证人岗位': w.get('岗位', '') or self.get_data('证人岗位', ''),
            '公司名称': case_obj.get('labor_unit', ''),  # 用人单位（签合同的单位）
        }

    def _generate_witness_transcript(self):
        """证人谈话笔录：构建提示词 → 发给AI → 渲染模板+插入（跳过条例分析）"""
        case_id = self.current_case_id or self.lineEdit_2.text().strip()
        if not case_id:
            self._set_status('无案本号，无法生成证人笔录', 'orange')
            return
        if not self.ai_service:
            self._set_status('未配置AI，无法生成证人笔录', 'orange')
            QMessageBox.warning(self, "提示", "未配置API密钥，无法生成证人笔录。\n请在顶部⚙配置中设置API密钥。")
            return
        case_obj = self._load_cases_data().get(case_id)
        if not case_obj:
            self._set_status('未找到该案本号的案件数据', 'orange')
            return

        prompt_text = self._build_witness_ai_prompt(case_obj)

        self._set_status('正在AI生成证人谈话笔录...', 'black')
        QApplication.processEvents()

        self.witness_transcript_worker = TranscriptFromTemplateWorker(self.ai_service, prompt_text)
        self.witness_transcript_worker.finished.connect(
            lambda result, cid=case_id, co=case_obj: self._on_witness_transcript_generated(cid, co, result)
        )
        self.witness_transcript_worker.error.connect(self._on_transcript_error)
        self.witness_transcript_worker.start()

    def _on_witness_transcript_generated(self, case_id: str, case_obj: dict, result: dict):
        if result.get("状态") != "成功":
            err = result.get("错误信息", "未知错误")
            print(f"⚠️ 证人谈话笔录生成失败: {err}")
            self._set_status(f'证人谈话笔录生成失败: {err[:40]}', 'orange')
            return
        content = (result.get("内容", "") or "").strip()
        if not content:
            print("⚠️ 证人谈话笔录生成失败：AI 返回内容为空")
            self._set_status('证人谈话笔录生成失败：AI 返回内容为空', 'orange')
            return
        path = self._save_witness_transcript_to_template(case_obj, content)
        if not path:
            return
        self._save_witnesses()  # 持久化证人数据
        success, message = self.file_service.open_document(path)
        if success:
            self._set_status('已生成并打开证人谈话笔录', 'green')
        else:
            self._set_status(f'证人谈话笔录已生成，打开失败: {message}', 'orange')

    def _save_witness_transcript_to_template(self, case_obj: dict, content: str) -> str:
        """渲染「证人谈话笔录（普通工伤案件）.docx」模板，删除样例问答后插入 AI 问答（下划线），返回路径"""
        try:
            from docx.shared import Pt

            # ── 1. 渲染证人谈话模板（替换占位符）──
            template_path = str(path_utils.get_talk_template_path('证人谈话笔录（普通工伤案件）.docx'))
            if not os.path.exists(template_path):
                print(f"⚠️ 证人谈话模板不存在: {template_path}")
                return ""

            template_data = self._build_witness_template_data(case_obj)

            from docxtpl import DocxTemplate
            doc = DocxTemplate(template_path)
            doc.render(template_data)

            # ── 2. 定位锚点「答：听清楚了，不申请回避」──
            anchor_index = None
            anchor_pf = None
            for i, p in enumerate(doc.paragraphs):
                if '答：听清楚了，不申请回避' in p.text:
                    anchor_index = i
                    anchor_pf = p.paragraph_format
                    break
            if anchor_index is None:
                print("⚠️ 未找到锚点「答：听清楚了，不申请回避」")
                return ""

            # ── 3. 删除锚点之后的模板样例问答（保留头部+告知，避免与AI问答重复）──
            body = doc.element.body
            anchor_elem = doc.paragraphs[anchor_index]._element
            after_anchor = False
            for child in list(body):
                if after_anchor:
                    body.remove(child)
                elif child is anchor_elem:
                    after_anchor = True

            # ── 4. 在锚点之后插入 AI 问答（每行下划线）──
            for line in content.splitlines():
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph()
                if anchor_pf is not None:
                    p.paragraph_format.alignment = anchor_pf.alignment
                    p.paragraph_format.first_line_indent = anchor_pf.first_line_indent
                    p.paragraph_format.space_before = anchor_pf.space_before
                    p.paragraph_format.space_after = anchor_pf.space_after
                run = p.add_run(line)
                run.font.size = Pt(15)
                run.underline = True

            # ── 5. 保存 ──
            if not self.current_case_folder or not os.path.exists(self.current_case_folder):
                folder_subject = str(case_obj.get('case_id', '') or case_obj.get('name', '') or '案件').strip()
                self.current_case_folder = os.path.join(self.BASE_PATH, folder_subject)
                os.makedirs(self.current_case_folder, exist_ok=True)

            subject = str(case_obj.get('name', '') or case_obj.get('case_id', '') or '案件').strip()
            file_name = f"{subject}证人谈话笔录.docx"
            target_path = os.path.join(self.current_case_folder, file_name)
            counter = 2
            while os.path.exists(target_path):
                file_name = f"{subject}证人谈话笔录({counter}).docx"
                target_path = os.path.join(self.current_case_folder, file_name)
                counter += 1

            doc.save(target_path)
            print(f"✅ 证人谈话笔录已生成: {target_path}")
            return target_path
        except Exception as e:
            print(f"❌ 生成证人谈话笔录失败: {e}")
            import traceback
            traceback.print_exc()
            self._set_status('生成证人谈话笔录失败', 'red')
            return ""

    # ========================================================================
    # 法人谈话笔录生成（仿证人流程，跳过条例分析）
    # ========================================================================

    def _sync_legal_to_form(self):
        """数据核对后把表单切回法人数据（open_data_review 会把表单回填成本人数据）"""
        if self.get_current_role_type() != "法人":
            return
        self.name_pane.setText(self.get_data('法人姓名', ''))
        self.idnumer_pane.setText(self.get_data('法人身份证号', ''))
        self.textEdit.setPlainText(self.get_data('法人身份证地址', ''))
        self.lineEdit_4.setText(self.get_data('法人手机号', ''))
        self.lineEdit_5.setText(self.get_data('法人职务', ''))
        self.lineEdit.setText(self.get_data('法人性别', ''))
        self.age_pane.setText(str(self.get_data('法人年龄', '')))

    def _build_legal_ai_prompt(self, case_obj: dict) -> str:
        """构建法人谈话笔录的 AI 提示词（案件数据 + 法人数据 替换占位符）"""
        from prompt_manager import load_prompt
        prompt = load_prompt('legal_send_to_ai')
        replacements = {
            '案本号': case_obj.get('case_id', ''),
            '案件性质': case_obj.get('case_nature', ''),
            '申请类型': case_obj.get('applicant_type', ''),
            '本人姓名': case_obj.get('name', ''),
            '本人性别': case_obj.get('gender', ''),
            '本人身份证号': case_obj.get('id_card', ''),
            '用人单位': case_obj.get('labor_unit', ''),
            '用工单位': case_obj.get('employer', ''),
            '工地名称': case_obj.get('site', ''),
            '受伤经过': case_obj.get('injury_description', ''),
            '法人姓名': self.get_data('法人姓名', ''),
            '法人职务': self.get_data('法人职务', ''),
            '法人身份证号': self.get_data('法人身份证号', ''),
        }
        for key, val in replacements.items():
            prompt = prompt.replace('{{%s}}' % key, str(val))
        return prompt

    def _build_legal_template_data(self, case_obj: dict) -> dict:
        """构建法人谈话笔录模板的占位符数据"""
        return {
            '当前时期': self.get_data('当前时期', '') or (_date_now() + _time_now()),
            '用户名': self._get_current_username(),
            '本人姓名': case_obj.get('name', ''),
            '法人姓名': self.get_data('法人姓名', ''),
            '法人性别': self.get_data('法人性别', ''),
            '法人年龄': self.get_data('法人年龄', ''),
            '法人身份证号': self.get_data('法人身份证号', ''),
            '法人身份证地址': self.get_data('法人身份证地址', ''),
            '法人手机号': self.get_data('法人手机号', ''),
            '法人职务': self.get_data('法人职务', ''),
            '公司名称': case_obj.get('labor_unit', ''),  # 用人单位（签合同的单位）
        }

    def _generate_legal_transcript(self):
        """法人谈话笔录：构建提示词 → 发给AI → 渲染模板+插入（跳过条例分析）"""
        case_id = self.current_case_id or self.lineEdit_2.text().strip()
        if not case_id:
            self._set_status('无案本号，无法生成法人笔录', 'orange')
            return
        if not self.ai_service:
            self._set_status('未配置AI，无法生成法人笔录', 'orange')
            QMessageBox.warning(self, "提示", "未配置API密钥，无法生成法人笔录。\n请在顶部⚙配置中设置API密钥。")
            return
        case_obj = self._load_cases_data().get(case_id)
        if not case_obj:
            self._set_status('未找到该案本号的案件数据', 'orange')
            return

        prompt_text = self._build_legal_ai_prompt(case_obj)

        self._set_status('正在AI生成法人谈话笔录...', 'black')
        QApplication.processEvents()

        self.legal_transcript_worker = TranscriptFromTemplateWorker(self.ai_service, prompt_text)
        self.legal_transcript_worker.finished.connect(
            lambda result, cid=case_id, co=case_obj: self._on_legal_transcript_generated(cid, co, result)
        )
        self.legal_transcript_worker.error.connect(self._on_transcript_error)
        self.legal_transcript_worker.start()

    def _on_legal_transcript_generated(self, case_id: str, case_obj: dict, result: dict):
        if result.get("状态") != "成功":
            err = result.get("错误信息", "未知错误")
            print(f"⚠️ 法人谈话笔录生成失败: {err}")
            self._set_status(f'法人谈话笔录生成失败: {err[:40]}', 'orange')
            return
        content = (result.get("内容", "") or "").strip()
        if not content:
            print("⚠️ 法人谈话笔录生成失败：AI 返回内容为空")
            self._set_status('法人谈话笔录生成失败：AI 返回内容为空', 'orange')
            return
        path = self._save_legal_transcript_to_template(case_obj, content)
        if not path:
            return
        success, message = self.file_service.open_document(path)
        if success:
            self._set_status('已生成并打开法人谈话笔录', 'green')
        else:
            self._set_status(f'法人谈话笔录已生成，打开失败: {message}', 'orange')

    def _save_legal_transcript_to_template(self, case_obj: dict, content: str) -> str:
        """渲染「法人谈话笔录（普通工伤案件）.docx」模板，删除样例问答后插入 AI 问答（下划线），返回路径"""
        try:
            from docx.shared import Pt

            # ── 1. 渲染法人谈话模板（替换占位符）──
            template_path = str(path_utils.get_talk_template_path('法人谈话笔录（普通工伤案件）.docx'))
            if not os.path.exists(template_path):
                print(f"⚠️ 法人谈话模板不存在: {template_path}")
                return ""

            template_data = self._build_legal_template_data(case_obj)

            from docxtpl import DocxTemplate
            doc = DocxTemplate(template_path)
            doc.render(template_data)

            # ── 2. 定位锚点「答：听清楚了，不申请回避」──
            anchor_index = None
            anchor_pf = None
            for i, p in enumerate(doc.paragraphs):
                if '答：听清楚了，不申请回避' in p.text:
                    anchor_index = i
                    anchor_pf = p.paragraph_format
                    break
            if anchor_index is None:
                print("⚠️ 未找到锚点「答：听清楚了，不申请回避」")
                return ""

            # ── 3. 删除锚点之后的模板样例问答（保留头部+告知，避免与AI问答重复）──
            body = doc.element.body
            anchor_elem = doc.paragraphs[anchor_index]._element
            after_anchor = False
            for child in list(body):
                if after_anchor:
                    body.remove(child)
                elif child is anchor_elem:
                    after_anchor = True

            # ── 4. 在锚点之后插入 AI 问答（每行下划线）──
            for line in content.splitlines():
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph()
                if anchor_pf is not None:
                    p.paragraph_format.alignment = anchor_pf.alignment
                    p.paragraph_format.first_line_indent = anchor_pf.first_line_indent
                    p.paragraph_format.space_before = anchor_pf.space_before
                    p.paragraph_format.space_after = anchor_pf.space_after
                run = p.add_run(line)
                run.font.size = Pt(15)
                run.underline = True

            # ── 5. 保存 ──
            if not self.current_case_folder or not os.path.exists(self.current_case_folder):
                folder_subject = str(case_obj.get('case_id', '') or case_obj.get('name', '') or '案件').strip()
                self.current_case_folder = os.path.join(self.BASE_PATH, folder_subject)
                os.makedirs(self.current_case_folder, exist_ok=True)

            subject = str(case_obj.get('name', '') or case_obj.get('case_id', '') or '案件').strip()
            file_name = f"{subject}法人谈话笔录.docx"
            target_path = os.path.join(self.current_case_folder, file_name)
            counter = 2
            while os.path.exists(target_path):
                file_name = f"{subject}法人谈话笔录({counter}).docx"
                target_path = os.path.join(self.current_case_folder, file_name)
                counter += 1

            doc.save(target_path)
            print(f"✅ 法人谈话笔录已生成: {target_path}")
            return target_path
        except Exception as e:
            print(f"❌ 生成法人谈话笔录失败: {e}")
            import traceback
            traceback.print_exc()
            self._set_status('生成法人谈话笔录失败', 'red')
            return ""

    def _on_analysis_error(self, err: str):
        self._set_status('AI分析出错', 'red')
        QMessageBox.critical(self, "AI分析错误", f"分析出错: {err}")

    def _show_regulation_analysis(self, case_id: str, result: dict):
        judged = result.get('judged_article', '')
        judged_reason = result.get('judged_article_reason', '')
        missing = result.get('missing_evidence', []) or []
        consistency = result.get('consistency', '')
        reason = result.get('reason', '')
        proposed = self._load_cases_data().get(case_id, {}).get('proposed_article', '')

        dlg = QDialog(self)
        dlg.setWindowTitle("AI 条例分析结果")
        dlg.resize(620, 600)
        layout = QVBoxLayout(dlg)

        title = QLabel("AI 条例与证据分析")
        title.setStyleSheet("font-size: 15px; font-weight: bold; padding: 4px;")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        text = QTextEdit()
        text.setReadOnly(True)
        text.setFont(QFont("微软雅黑", 10))
        lines = [
            "【判断应适用条例】",
            f"　{judged}",
        ]
        if judged_reason:
            lines.append(f"　依据：{judged_reason}")
        lines.append("")
        lines.append("【工作人员拟用条例】")
        lines.append(f"　{proposed or '未填写'}")
        lines.append("")
        lines.append("【一致/分歧判断】")
        lines.append(f"　{consistency}")
        if reason:
            lines.append(f"　理由：{reason}")
        lines.append("")
        lines.append("【尚缺关键证据清单】")
        if missing:
            for i, e in enumerate(missing, 1):
                lines.append(f"　{i}. {e}")
        else:
            lines.append("　（无）")
        text.setPlainText("\n".join(lines))
        layout.addWidget(text, 1)

        btn_layout = QHBoxLayout()
        if consistency == '分歧':
            q = QLabel("是否把拟用条例修改为 AI 判断的条例？")
            q.setStyleSheet("color:#c0392b; font-weight:bold;")
            btn_layout.addWidget(q)
            btn_layout.addStretch()
            no_btn = QPushButton("否")
            no_btn.clicked.connect(lambda: self._on_regulation_choice(False, dlg, case_id, judged, missing))
            btn_layout.addWidget(no_btn)
            yes_btn = QPushButton("是")
            yes_btn.setStyleSheet(
                "QPushButton{background-color:#27ae60;color:white;font-weight:bold;padding:5px 22px;border-radius:4px;}"
            )
            yes_btn.clicked.connect(lambda: self._on_regulation_choice(True, dlg, case_id, judged, missing))
            btn_layout.addWidget(yes_btn)
        else:
            btn_layout.addStretch()
            close_btn = QPushButton("关闭")
            close_btn.clicked.connect(dlg.accept)
            btn_layout.addWidget(close_btn)

        layout.addLayout(btn_layout)
        dlg.exec_()

    def _on_regulation_choice(self, yes: bool, dlg: QDialog, case_id: str, judged: str, missing: list):
        dlg.accept()
        if yes:
            self._apply_regulation_change(case_id, judged, missing)
            QMessageBox.information(self, "已修改", f"拟用条例已修改为：{judged}\n关键证据清单已相应更新。")
        else:
            self._set_status('已保留原拟用条例', 'black')

    def _apply_regulation_change(self, case_id: str, judged_article: str, missing_evidence: list):
        """把拟用条例修改为 AI 判断的条例，并同步界面 + JSON + 证据清单"""
        # 1. 主界面条例下拉框 + 数据模型
        self._apply_regulation(judged_article)
        # 2. JSON 里的 proposed_article
        cases = self._load_cases_data()
        case_obj = cases.get(case_id)
        if case_obj:
            case_obj['proposed_article'] = judged_article
            case_obj['proposed_article_elements'] = _regulation_elements(judged_article)
            self._save_cases_data(cases)
        # 3. 更新关键证据清单：把缺失证据作为未勾选项补进材料清单
        if missing_evidence and hasattr(self, 'material_list'):
            current = self.material_list.get_materials()
            existing = {m.get('name', '') for m in current}
            for ev in missing_evidence:
                if ev and ev not in existing:
                    current.append({"name": ev, "provided": False, "notes": ""})
                    existing.add(ev)
            self.material_list.set_materials(current)
            self.data_model.investigation['本人材料'] = current
        self._set_status(f'拟用条例已修改为：{judged_article}', 'green')

    # ========================================================================
    # F2 测试数据轮换
    # ========================================================================

    def keyPressEvent(self, event):
        """F2 键轮换测试数据"""
        if event.key() == Qt.Key_F2:
            self._cycle_test_data()
        else:
            super().keyPressEvent(event)

    def _cycle_test_data(self):
        """按 F2 切换到下一组测试数据"""
        self._test_data_index = (self._test_data_index + 1) % len(TEST_DATA_PRESETS)
        data = TEST_DATA_PRESETS[self._test_data_index]

        print(f"\n{'=' * 50}")
        print(f"F2 测试数据: {data['name']}")
        print(f"{'=' * 50}")

        # ── 角色单选按钮 ──
        role_map = {"本人": self.radioButton, "证人": self.radioButton_2, "法人": self.radioButton_3}
        for role_name, btn in role_map.items():
            btn.setChecked(role_name == data["role"])
        # 程序化 setChecked 不会触发 clicked，手动触发一次角色切换清理
        self.clear_role_fields()

        # ── 案例类型复选框 ──
        self.death_case_checkbox.setChecked(data["deathCaseCheckbox"])
        self.personal_application_checkbox.setChecked(data["personalApplicationCheckbox"])
        self.on_case_type_changed()

        # ── 基本信息输入框 ──
        self.name_pane.setText(data["name_pane"])
        self.idnumer_pane.setText(data["idnumer_pane"])
        self.textEdit.setPlainText(data["textEdit"])
        self.lineEdit_4.setText(data["lineEdit_4"])
        self.lineEdit_5.setText(data["lineEdit_5"])
        # 案本号：切到本人时清空（输入新姓名后自动生成）；切到证人/法人保留同一案本号
        if data["role"] == "本人":
            self.lineEdit_2.clear()

        # ── 条例下拉框 ──
        self.comboBox.setCurrentIndex(data["comboBox"])

        # ── 公司下拉框 ──
        self._set_combo_or_type(self.company_pane, data["company_pane"])
        self._set_combo_or_type(self.construction_company, data["construction_company"])
        self._set_combo_or_type(self.construction_plant, data["construction_plant"])

        # ── 自动计算年龄和性别 ──
        self.on_id_input_finished()

        # ── 同步角色数据到数据模型，供 JSON 保存（本人=键，证人是列表，法人是单条键） ──
        role = self.get_current_role_type()
        if role == "本人":
            self.set_data('本人姓名', data['name_pane'], 'basic')
            self.set_data('本人手机号', data['lineEdit_4'], 'basic')
            self.set_data('本人身份证地址', data['textEdit'], 'basic')
            self.set_data('本人岗位', data['lineEdit_5'], 'basic')
        elif role == "证人":
            name = data['name_pane']
            w = next((x for x in self.data_model.witnesses if x.get('姓名') == name), None)
            if w is None:
                w = {"序号": _witness_label(len(self.data_model.witnesses) + 1)}
                self.data_model.witnesses.append(w)
            w.update({
                "姓名": name,
                "身份证号": data['idnumer_pane'],
                "身份证地址": data['textEdit'],
                "手机号": data['lineEdit_4'],
                "岗位": data['lineEdit_5'],
                "性别": self.lineEdit.text().strip(),
                "年龄": self.age_pane.text().strip(),
            })
            # 指向该证人，避免索引无效导致生成/回填拿不到证人数据
            self.data_model.current_witness_index = self.data_model.witnesses.index(w)
        elif role == "法人":
            self.set_data('法人姓名', data['name_pane'], 'basic')
            self.set_data('法人身份证号', data['idnumer_pane'], 'basic')
            self.set_data('法人身份证地址', data['textEdit'], 'basic')
            self.set_data('法人手机号', data['lineEdit_4'], 'basic')
            self.set_data('法人职务', data['lineEdit_5'], 'basic')

        # ── 右侧面板（同案沿用） ──
        current_worker = data.get("injured_worker", "")
        prev_worker = getattr(self, '_prev_injured_worker', None)
        is_same_case = (prev_worker is not None and current_worker == prev_worker)
        if not is_same_case:
            self.current_case_id = ""  # 换了受伤职工 → 视为新案件，重置案本号关联
            self.set_data('案本号', '', 'case')  # 同时清掉数据模型里缓存的旧案本号

        if hasattr(self, 'statement_edit'):
            stmt = data.get("statement_edit", "")
            # 本人填案件陈述；证人/法人无陈述则清空
            self.statement_edit.setPlainText(stmt)
            if is_same_case:
                print(f"📋 同案沿用案件陈述（{current_worker}）")

        # 材料：只有本人的证据有效（证人/法人的证据不保存）
        if role == "本人":
            mats = data.get('materials', [])
            if hasattr(self, 'material_list'):
                self.material_list.set_materials(mats)
            self.data_model.investigation['本人材料'] = mats
        else:
            # 证人/法人：恢复显示本人证据，不显示证人/法人证据
            if hasattr(self, 'material_list'):
                self.material_list.set_materials(self.data_model.investigation.get('本人材料', []))

        self._prev_injured_worker = current_worker

        # ── 重置案件状态，允许重新生成笔录 ──
        self.pushButton.setEnabled(True)
        self.pushButton.setStyleSheet("")
        self.current_case_folder = None
        self.current_person_name = ""

        # ── 状态提示 ──
        label = (f"[测试 {self._test_data_index + 1}/{len(TEST_DATA_PRESETS)}] "
                 f"{data['name']}  |  F2=下一个")
        self._set_status(label, 'green')
        print(f"OK 测试数据已填充: {data['name']}")

    def on_role_changed(self):
        """当角色切换时调用"""
        print(f"🔄 角色切换: {self.get_current_role_type()}")

    def _setup_paths(self):
        """统一使用PathUtils设置所有路径"""
        print("=" * 50)
        print("🔄 使用统一的PathUtils设置所有路径...")

        # 使用PathUtils获取路径（path_utils.get_xxx() 方法已确保目录存在）
        self.BASE_PATH = str(path_utils.get_storage_path())
        self.TEMPLATE_PATH = str(path_utils.get_template_path())
        self.CONFIG_PATH = str(path_utils.get_config_path(""))
        self.DATA_PATH = str(path_utils.get_data_path(""))

        # 获取模板子目录（path_utils 已确保目录存在）
        self.TALK_TEMPLATE_PATH = str(path_utils.get_talk_template_path())
        self.DOCUMENT_TEMPLATE_PATH = str(path_utils.get_document_template_path())

    def _update_services_paths(self):
        """更新所有服务的路径"""
        print("🔄 更新服务路径...")

        # 更新FileService
        if hasattr(self, 'file_service'):
            self.file_service.BASE_PATH = self.BASE_PATH
            print(f"✅ 更新FileService路径: {self.BASE_PATH}")


        # 更新数据模型
        if hasattr(self, 'data_model'):
            self.data_model.company_info['存储路径'] = self.BASE_PATH
            self.data_model.company_info['模板路径'] = self.TEMPLATE_PATH

        print("✅ 服务路径更新完成")

    def select_all_questions(self, select_all: bool):
        """全选或全不选问题"""
        if not hasattr(self, 'question_list_widget'):
            return

        for i in range(self.question_list_widget.count()):
            item = self.question_list_widget.item(i)
            item.setCheckState(Qt.Checked if select_all else Qt.Unchecked)

    def insert_selected_questions(self, dialog):
        """将选中的问题插入到笔录文档"""
        try:
            # 获取选中的问题
            selected_questions = []
            for i in range(self.question_list_widget.count()):
                item = self.question_list_widget.item(i)
                if item.checkState() == Qt.Checked:
                    selected_questions.append(item.text())

            if not selected_questions:
                QMessageBox.warning(dialog, "提示", "请至少选择一个要插入的问题")
                return

            print(f"✅ 选择了 {len(selected_questions)} 个问题准备插入")

            # 检查案件文件夹
            if not self.current_case_folder:
                QMessageBox.warning(dialog, "提示", "请先保存案件信息")
                return

            # 查找本人笔录文件
            person_name = self.get_data("本人姓名", "") or self.name_pane.text().strip()
            if not person_name:
                QMessageBox.warning(dialog, "提示", "请先输入受伤职工姓名")
                return

            # 查找笔录文件
            person_files = []
            for file in os.listdir(self.current_case_folder):
                if "本人" in file and file.endswith('.docx') and "审批表" not in file:
                    person_files.append(file)

            if not person_files:
                QMessageBox.warning(dialog, "提示", "未找到本人笔录文件")
                return

            # 使用第一个找到的本人笔录
            file_path = os.path.join(self.current_case_folder, person_files[0])

            # 插入问题到文档
            success, message = self.insert_questions_to_document(file_path, selected_questions)

            if success:
                QMessageBox.information(dialog, "成功",
                                        f"已成功插入 {len(selected_questions)} 个问题到笔录中\n\n"
                                        f"文件：{os.path.basename(file_path)}")
                dialog.close()
            else:
                QMessageBox.critical(dialog, "失败", f"插入失败：{message}")

        except Exception as e:
            print(f"❌ 插入问题失败: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(dialog, "错误", f"插入过程中发生错误：{str(e)}")

    def insert_questions_to_document(self, file_path: str, questions: list) -> tuple:
        """
        将问题插入到Word文档中

        Args:
            file_path: Word文档路径
            questions: 问题列表

        Returns:
            (是否成功, 消息)
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor

            print(f"📄 正在处理文档: {file_path}")

            # 打开文档
            doc = Document(file_path)

            # 查找插入位置（文档末尾）
            # 我们可以找到最后一个段落，然后在其后插入

            # 添加分隔线
            separator = doc.add_paragraph("=" * 50)
            separator.alignment = 1  # 居中

            # 添加标题
            title = doc.add_paragraph("AI建议补充问题：")
            title.runs[0].bold = True
            title.alignment = 0  # 左对齐

            # 插入每个问题
            for i, question_text in enumerate(questions, 1):
                # 添加问题
                question_para = doc.add_paragraph()
                question_para.add_run(f"{i}. {question_text}")

                # 添加答案占位符（带下划线）
                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run("答：")
                answer_run.font.underline = True
                answer_run.font.color.rgb = RGBColor(0, 0, 0)

                # 添加空行
                doc.add_paragraph()

            # 保存文档
            doc.save(file_path)

            print(f"✅ 成功插入 {len(questions)} 个问题到文档")

            return True, "插入成功"

        except Exception as e:
            print(f"❌ 插入问题到文档失败: {e}")
            import traceback
            traceback.print_exc()
            return False, str(e)

    def parse_ai_result(self, ai_text: str) -> dict:
        """
        解析AI结果，提取审查结果和缺失问题

        Args:
            ai_text: AI返回的完整文本

        Returns:
            包含审查结果和缺失问题的字典
        """
        result = {
            "审查结果": "",
            "缺失问题": [],
            "原始文本": ai_text
        }

        try:
            # 分割审查结果和缺失问题
            if "【审查结果】" in ai_text and "【缺失问题列表】" in ai_text:
                # 提取审查结果部分
                start = ai_text.find("【审查结果】")
                end = ai_text.find("【缺失问题列表】")

                if start != -1 and end != -1:
                    review_text = ai_text[start:end]
                    # 清理标记
                    review_text = review_text.replace("【审查结果】", "").strip()
                    result["审查结果"] = review_text

                    # 提取缺失问题部分
                    questions_text = ai_text[end:]
                    # 按行分割
                    lines = questions_text.split('\n')

                    for line in lines:
                        line = line.strip()
                        # 查找带方框的问题行
                        if "□" in line and "问：" in line:
                            # 提取问题文本（去掉方框和序号）
                            # 示例：□ 1. 问：您与公司是否签订了书面劳动合同？
                            question = line
                            # 去掉方框标记
                            question = question.replace("□", "", 1).strip()
                            # 去掉序号（如"1. "）
                            if "." in question:
                                question = question.split(".", 1)[1].strip()

                            result["缺失问题"].append(question)

            # 如果格式不正确，尝试其他解析方式
            elif "审查结果" in ai_text and "缺失问题" in ai_text:
                # 尝试其他格式解析
                pass

            else:
                # 如果没有找到格式标记，整个文本作为审查结果
                result["审查结果"] = ai_text

        except Exception as e:
            print(f"解析AI结果失败: {e}")
            result["审查结果"] = ai_text

        print(f"✅ 解析结果: 审查结果长度={len(result['审查结果'])}, 问题数量={len(result['缺失问题'])}")
        return result

    def _reconnect_approval_button(self):
        """重新连接案件审批表按钮 (pushButton_11)"""
        try:
            # 断开所有现有连接
            self.pushButton_11.clicked.disconnect()
            print("🔌 已断开 pushButton_11 的旧连接")
        except:
            pass  # 如果没有连接，忽略

        # 连接到 approve() 方法
        self.pushButton_11.clicked.connect(self.approve)
        print("✅ 案件审批表按钮 (pushButton_11) 已连接到 approve()")

        # 可选：确认按钮文本
        print(f"📝 按钮文本: '{self.pushButton_11.text()}'")

    def on_pushButton_12_clicked(self):
        """谈话通知书按钮点击事件"""
        print("🔄 谈话通知书按钮被点击")

        # 调用谈话通知书生成函数
        self.generate_interview_notice_from_approval()

    def extract_data_from_approval_table(self, file_path):
        """
        从审批表Word文件中提取数据（简化版）
        """
        try:
            from docx import Document

            print(f"📄 开始提取审批表数据: {file_path}")

            # 搜索关键词映射
            search_mapping = {
                '用人单位': '用人单位',
                '职工姓名': '职工姓名',
                '身份证号': '职工身份证号',
                '申请时间': '申请时间',
                '受理时间': '受理时间',
                '受伤经过': '受伤经过',
                '医疗诊断': '医疗证明',  # 搜索"医疗诊断"
            }

            extracted_data = {}
            document = Document(file_path)

            # 遍历所有表格
            for table in document.tables:
                for row in table.rows:
                    # 遍历每个单元格
                    for i, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()

                        # 检查每个搜索词
                        for search_term, data_field in search_mapping.items():
                            # 如果已经找到了，跳过
                            if data_field in extracted_data:
                                continue

                            # 检查是否包含搜索词
                            if search_term in cell_text:
                                print(f"✅ 找到关键词 '{search_term}'")

                                # 尝试获取右边的单元格
                                if i + 1 < len(row.cells):
                                    right_cell = row.cells[i + 1]
                                    right_text = right_cell.text.strip()

                                    # 只有当右边单元格有内容且不是关键词时才使用
                                    if right_text and right_text != search_term:
                                        extracted_data[data_field] = right_text
                                        print(f"  提取 {data_field}: {right_text}")
                                    else:
                                        # 右边单元格是空的，标记为红色
                                        print(f"  ⚠️ {data_field}: 右边单元格为空")
                                else:
                                    print(f"  ⚠️ {data_field}: 没有右侧单元格")

            print("\n📋 提取结果:")
            required_fields = ['用人单位', '职工姓名', '职工身份证号', '申请时间', '受理时间', '受伤经过', '医疗证明']

            for field in required_fields:
                if field in extracted_data:
                    print(f"  ✅ {field}: {extracted_data[field]}")
                else:
                    print(f"  ❌ {field}: 未找到")

            # 填充缺失字段
            current_date = _date_now()

            if '申请时间' not in extracted_data:
                extracted_data['申请时间'] = current_date
                print(f"  🟥 使用默认值 申请时间: {current_date}")

            if '受理时间' not in extracted_data:
                extracted_data['受理时间'] = current_date
                print(f"  🟥 使用默认值 受理时间: {current_date}")

            if '医疗证明' not in extracted_data:
                extracted_data['医疗证明'] = '详见医疗诊断证明'
                print(f"  🟥 使用默认值 医疗证明: 详见医疗诊断证明")

            # 其他字段使用界面数据
            if '用人单位' not in extracted_data:
                extracted_data['用人单位'] = self.company_pane.currentText().strip() or '未知公司'

            if '职工姓名' not in extracted_data:
                extracted_data['职工姓名'] = self.get_data('本人姓名', '') or self.name_pane.text().strip() or '未知'

            if '职工身份证号' not in extracted_data:
                extracted_data['职工身份证号'] = self.get_data('本人身份证号', '')

            return extracted_data

        except Exception as e:
            print(f"❌ 提取审批表数据失败: {e}")
            import traceback
            traceback.print_exc()

            # 返回最小可用数据
            current_date = _date_now()
            return {
                '用人单位': self.company_pane.currentText().strip() or '未知公司',
                '职工姓名': self.get_data('本人姓名', '') or self.name_pane.text().strip() or '未知',
                '职工身份证号': self.get_data('本人身份证号', ''),
                '申请时间': current_date,
                '受理时间': current_date,
                '受伤经过': '详见谈话笔录',
                '医疗证明': '详见医疗诊断证明'
            }

    def generate_interview_notice_from_approval(self):
        """从审批表生成接受谈话通知书"""
        try:
            # 1. 检查本人姓名
            person_name = self.get_data("本人姓名", "")
            if not person_name:
                person_name = self.name_pane.text().strip()
            if not person_name:
                self._set_status('请先输入本人姓名', 'red')
                return

            # 2. 检查案件文件夹
            if not self.current_case_folder or not os.path.exists(self.current_case_folder):
                self._set_status('请先保存案件信息', 'red')
                return

            # 3. 查找本人案件审批表文件
            approval_file_name = f"{person_name}案件审批表.docx"
            approval_file_path = os.path.join(self.current_case_folder, approval_file_name)

            if not os.path.exists(approval_file_path):
                self._set_status(f'未找到审批表: {approval_file_name}', 'red')
                # 尝试查找其他可能的审批表文件
                all_files = os.listdir(self.current_case_folder)
                approval_files = [f for f in all_files if "审批表" in f and f.endswith('.docx')]

                if not approval_files:
                    self._set_status('案件文件夹中没有审批表文件', 'red')
                    return

                # 使用找到的第一个审批表文件
                approval_file_path = os.path.join(self.current_case_folder, approval_files[0])
                print(f"⚠️ 使用替代审批表: {approval_files[0]}")

            print(f"✅ 找到审批表文件: {approval_file_path}")

            # 4. 从审批表提取数据
            self._set_status('正在提取审批表数据...', 'black')
            QApplication.processEvents()

            extracted_data = self.extract_data_from_approval_table(approval_file_path)

            # 处理申请时间
            if '申请时间' in extracted_data:
                apply_time = extracted_data['申请时间']
                if isinstance(apply_time, str) and apply_time.isdigit() and len(apply_time) == 8:
                    extracted_data['申请时间'] = f"{apply_time[0:4]}年{apply_time[4:6]}月{apply_time[6:8]}日"

            # 处理受理时间
            if '受理时间' in extracted_data:
                accept_time = extracted_data['受理时间']
                if isinstance(accept_time, str) and accept_time.isdigit() and len(accept_time) == 8:
                    extracted_data['受理时间'] = f"{accept_time[0:4]}年{accept_time[4:6]}月{accept_time[6:8]}日"

            if not extracted_data:
                self._set_status('提取审批表数据失败', 'red')
                return

            # 检查必要字段
            required_fields = ['用人单位', '职工姓名', '职工身份证号', '申请时间', '受理时间', '受伤经过', '医疗证明']
            missing_required = []

            for field in required_fields:
                if field not in extracted_data or not extracted_data[field]:
                    missing_required.append(field)

            if missing_required:
                self._set_status(f'审批表缺少必要字段: {missing_required}', 'red')
                return

            self._set_status('审批表数据提取成功', 'green')

            # 5. 准备模板数据
            current_date = _date_now()

            # 使用docxtpl的RichText来设置红色
            from docxtpl import RichText

            template_data = {
                '用人单位': extracted_data.get('用人单位', self.company_pane.currentText().strip()),
                '本人姓名': extracted_data.get('职工姓名', self.get_data('本人姓名', '')),
                '职工性别': extracted_data.get('职工性别', self.get_data('本人性别', '')),
                '本人身份证号': extracted_data.get('职工身份证号', self.get_data('本人身份证号', '')),
                '受伤经过': extracted_data.get('受伤经过', '详见谈话笔录'),
                '当前时期': current_date,
                '当前日期': current_date,
                '当前时间': _time_now(),
                '案本号': self.get_data('案本号', '')
            }

            # 申请时间：如果有就用，没有就用红色的当前日期
            if '申请时间' in extracted_data and extracted_data['申请时间']:
                template_data['申请时间'] = extracted_data['申请时间']
            else:
                rt = RichText()
                rt.add(current_date, color='FF0000')  # 红色
                template_data['申请时间'] = rt

            # 受理时间：如果有就用，没有就用红色的当前日期
            if '受理时间' in extracted_data and extracted_data['受理时间']:
                template_data['受理时间'] = extracted_data['受理时间']
            else:
                rt = RichText()
                rt.add(current_date, color='FF0000')  # 红色
                template_data['受理时间'] = rt

            # 医疗证明：如果有就用，没有就用红色的"详见医疗诊断证明"
            if '医疗证明' in extracted_data and extracted_data['医疗证明']:
                template_data['医疗证明'] = extracted_data['医疗证明']
            else:
                rt = RichText()
                rt.add('详见医疗诊断证明', color='FF0000')  # 红色
                template_data['医疗证明'] = rt

            # ============ 关键修复：在这里检查模板文件并定义 template_path ============
            # 6. 检查模板文件
            template_path = str(path_utils.get_document_template_path('接受谈话通知书（样本）.docx'))

            # 打印哪些字段用了红色
            red_fields = []
            if '申请时间' not in extracted_data or not extracted_data['申请时间']:
                red_fields.append('申请时间')
            if '受理时间' not in extracted_data or not extracted_data['受理时间']:
                red_fields.append('受理时间')
            if '医疗证明' not in extracted_data or not extracted_data['医疗证明']:
                red_fields.append('医疗证明')

            if red_fields:
                print(f"🔴 以下字段使用红色: {red_fields}")

            # 6. 生成文档
            self._set_status('正在生成谈话通知书...', 'black')
            QApplication.processEvents()

            try:
                from docxtpl import DocxTemplate
                word = DocxTemplate(template_path)
                word.render(template_data)

                notice_file_name = f"{person_name}接受谈话通知书.docx"
                target_path = os.path.join(self.current_case_folder, notice_file_name)
                word.save(target_path)

                print(f"✅ 谈话通知书保存到: {target_path}")

                # 打开文件
                success, message = self.file_service.open_document(target_path)
                if success:
                    self._set_status('谈话通知书生成成功', 'green')
                else:
                    self._set_status(f'谈话通知书生成成功，但打开失败: {message}', 'orange')

            except Exception as e:
                print(f"❌ 生成谈话通知书失败: {e}")
                import traceback
                traceback.print_exc()
                self._set_status(f'生成谈话通知书失败: {str(e)}', 'red')

        except Exception as e:
            print(f"❌ 谈话通知书过程异常: {e}")
            import traceback
            traceback.print_exc()
            self._set_status(f'生成谈话通知书异常: {str(e)}', 'red')

    def _cleanup_resources(self):
        """统一清理所有临时资源（对话框、线程等）"""
        resources = ['wait_dialog', 'progress_dialog', 'ai_worker']

        for attr_name in resources:
            if hasattr(self, attr_name):
                try:
                    resource = getattr(self, attr_name)
                    if attr_name == 'ai_worker' and resource.isRunning():
                        resource.terminate()  # 改为terminate
                        resource.wait(5000)  # 等待5秒
                    elif hasattr(resource, 'close'):
                        resource.close()
                except Exception as e:
                    print(f"清理资源 {attr_name} 失败: {e}")
    def _set_status(self, text, color="black", label="label_14"):
        """设置状态标签"""
        label_widget = getattr(self, label, None)
        if not label_widget:
            return

        label_widget.setText(text)
        if color == "green":
            label_widget.setStyleSheet("QLabel{color:green;}")
        elif color == "red":
            label_widget.setStyleSheet("QLabel{color:red;}")
        elif color == "orange":
            label_widget.setStyleSheet("QLabel{color:orange;}")
        else:
            label_widget.setStyleSheet("QLabel{color:black;}")

    def _handle_ai_result(self, result=None, error=None, canceled=False):
        """
        统一处理AI操作结果
        """
        print(f"🔄 处理AI结果: result={result is not None}, error={error}, canceled={canceled}")

        # 添加简单的防重复
        if hasattr(self, '_is_handling_ai_result') and self._is_handling_ai_result:
            print("⚠️ 已经在处理AI结果，跳过重复调用")
            return

        self._is_handling_ai_result = True

        try:
            # 清理资源
            self._cleanup_resources()

            if canceled:
                print("⏹️ AI操作被用户取消")
                return

            if error:
                print(f"❌ AI操作出错: {error}")
                QMessageBox.critical(self, "AI审查错误", f"操作失败: {error}")
                return

            if result:
                print("✅ AI操作成功，显示结果")
                self.show_ai_review_result(result)
        except Exception as e:
            print(f"❌ 处理AI结果时出错: {e}")
            import traceback
            traceback.print_exc()
        finally:
            # 清理标志
            if hasattr(self, '_is_handling_ai_result'):
                delattr(self, '_is_handling_ai_result')

    def _setup_radio_connections(self):
        """设置单选按钮信号连接（简化版）"""
        # 断开现有连接（如果有）
        try:
            self.radioButton.clicked.disconnect()
            self.radioButton_2.clicked.disconnect()
            self.radioButton_3.clicked.disconnect()
        except:
            pass

        # 重新连接
        self.radioButton.clicked.connect(self.clear_role_fields)
        self.radioButton_2.clicked.connect(self.clear_role_fields)
        self.radioButton_3.clicked.connect(self.clear_role_fields)
        print("✅ 单选按钮信号重新连接")

    def _setup_api_config_ui(self):
        """在窗口顶部创建折叠配置栏，并下移所有现有控件；右侧增加辅助面板"""
        TOP_BAR_H = 28          # 顶部小横条高度
        RIGHT_PANEL_X = 478     # 右侧面板起始 x
        RIGHT_PANEL_W = 382     # 右侧面板宽度

        # --- 1. 调整窗口大小（左侧 470 + 右侧 400 = 870，加高以容纳日期组） ---
        WIN_W, WIN_H = 870, 850
        self.setMinimumSize(WIN_W, WIN_H)
        self.setMaximumSize(WIN_W, WIN_H)
        self.resize(WIN_W, WIN_H)

        # --- 2. 下移所有现有控件（只挪一个薄顶栏的空间） ---
        for child in self.children():
            if isinstance(child, QWidget) and child is not self:
                try:
                    geo = child.geometry()
                    child.setGeometry(geo.x(), geo.y() + TOP_BAR_H,
                                      geo.width(), geo.height())
                except Exception:
                    pass

        # ============================================================
        # 3. 顶部小横条（始终可见）：⚙ 配置 + 状态
        # ============================================================
        self.top_bar = QLabel(self)
        self.top_bar.setGeometry(0, 0, WIN_W, TOP_BAR_H)
        self.top_bar.setStyleSheet(
            "background-color: #e8e8e8; border-bottom: 1px solid #ccc;"
        )

        # 齿轮按钮：展开/收起用户配置
        self.config_toggle_btn = QPushButton("⚙", self)
        self.config_toggle_btn.setGeometry(4, 2, 28, 24)
        self.config_toggle_btn.setToolTip("显示/隐藏用户配置")
        self.config_toggle_btn.setStyleSheet(
            "QPushButton { background: transparent; border: none; font-size: 14px; }"
            "QPushButton:hover { background-color: #d0d0d0; border-radius: 3px; }"
        )
        self.config_toggle_btn.clicked.connect(self._toggle_config_panel)

        # 顶部状态文字
        self.top_status_label = QLabel("", self)
        self.top_status_label.setGeometry(37, 4, 430, 20)
        self.top_status_label.setStyleSheet("color: #888; background: transparent; border: none;")

        # ============================================================
        # 4. 可折叠的用户配置面板（默认隐藏）
        # ============================================================
        self.api_group = QGroupBox("用户配置", self)
        self.api_group.setGeometry(5, TOP_BAR_H + 2, 460, 84)
        self.api_group.setFont(QFont("微软雅黑", 9))
        self.api_group.hide()  # 默认隐藏

        # 第一行：用户名 + API密钥
        lbl_user = QLabel("用户:", self.api_group)
        lbl_user.setGeometry(10, 22, 35, 20)

        self.api_user_combo = QComboBox(self.api_group)
        self.api_user_combo.setEditable(True)
        self.api_user_combo.setGeometry(45, 20, 140, 22)
        self.api_user_combo.setPlaceholderText("输入用户名")
        self.api_user_combo.currentTextChanged.connect(self._on_user_combo_changed)

        lbl_key = QLabel("密钥:", self.api_group)
        lbl_key.setGeometry(195, 22, 35, 20)

        self.api_key_input = PasswordLineEdit(self.api_group)
        self.api_key_input.setGeometry(230, 20, 160, 22)
        self.api_key_input.setPlaceholderText("输入API密钥")

        # 第二行：记住我 + 保存 + 状态
        self.api_remember_cb = QCheckBox("记住我", self.api_group)
        self.api_remember_cb.setGeometry(10, 50, 70, 20)
        self.api_remember_cb.setChecked(True)

        self.api_save_btn = QPushButton("保存配置", self.api_group)
        self.api_save_btn.setGeometry(80, 48, 70, 23)
        self.api_save_btn.clicked.connect(self._on_save_api_config)

        self.api_status_label = QLabel("", self.api_group)
        self.api_status_label.setGeometry(160, 50, 290, 20)
        self.api_status_label.setStyleSheet("color: #888;")

        # ============================================================
        # 5. 右侧面板：案件申请陈述（上）
        # ============================================================
        STATEMENT_H = 350
        self.statement_group = QGroupBox("案件申请陈述", self)
        self.statement_group.setGeometry(RIGHT_PANEL_X, TOP_BAR_H + 3,
                                         RIGHT_PANEL_W, STATEMENT_H)
        self.statement_group.setFont(QFont("微软雅黑", 9))

        self.statement_edit = QTextEdit(self.statement_group)
        self.statement_edit.setGeometry(8, 18, RIGHT_PANEL_W - 16, STATEMENT_H - 50)
        self.statement_edit.setPlaceholderText("在此输入案件申请陈述...")
        self.statement_edit.setStyleSheet("""
            QTextEdit {
                border: 1px solid #ccc;
                border-radius: 2px;
                background-color: #fafafa;
                font-size: 9pt;
            }
            QTextEdit:focus {
                border-color: #3498db;
                background-color: #fff;
            }
        """)

        # 按钮：手动定位，不用 layout（避免和 QGroupBox 冲突）
        btn_y = STATEMENT_H - 28
        stmt_copy_btn = QPushButton("复制", self.statement_group)
        stmt_copy_btn.setGeometry(8, btn_y, 45, 23)
        stmt_copy_btn.clicked.connect(lambda: self._copy_statement())

        stmt_clear_btn = QPushButton("清空", self.statement_group)
        stmt_clear_btn.setGeometry(58, btn_y, 45, 23)
        stmt_clear_btn.clicked.connect(lambda: self.statement_edit.clear())

        # ============================================================
        # 6. 右侧面板：目前提供的材料分类（下）
        # ============================================================
        MATERIAL_H = 235
        material_top = TOP_BAR_H + STATEMENT_H + 10
        self.material_group = QGroupBox("目前提供的材料分类", self)
        self.material_group.setGeometry(RIGHT_PANEL_X, material_top,
                                        RIGHT_PANEL_W, MATERIAL_H)
        self.material_group.setFont(QFont("微软雅黑", 9))

        # 用 MaterialListWidget 替换原来的 QTextEdit
        self.material_list = MaterialListWidget(self.material_group)
        self.material_list.setGeometry(8, 18, RIGHT_PANEL_W - 16, MATERIAL_H - 50)

        # 按钮：手动定位
        mat_btn_y = MATERIAL_H - 28
        mat_copy_btn = QPushButton("复制", self.material_group)
        mat_copy_btn.setGeometry(8, mat_btn_y, 45, 23)
        mat_copy_btn.clicked.connect(lambda: self._copy_material())

        mat_clear_btn = QPushButton("清空", self.material_group)
        mat_clear_btn.setGeometry(58, mat_btn_y, 45, 23)
        mat_clear_btn.clicked.connect(lambda: self.material_list.clear())

        mat_add_btn = QPushButton("新增", self.material_group)
        mat_add_btn.setGeometry(108, mat_btn_y, 45, 23)
        mat_add_btn.clicked.connect(lambda: self.material_list.add_row())

        # ============================================================
        # 7. 左栏：申请 / 受理 / 就诊 时间
        # ============================================================
        DATE_H = 100
        self.date_group = QGroupBox("申请 / 受理 / 就诊时间", self)
        self.date_group.setGeometry(70, 655, 390, DATE_H)
        self.date_group.setFont(QFont("微软雅黑", 9))

        lbl_apply = QLabel("申请时间：", self.date_group)
        lbl_apply.setGeometry(8, 30, 60, 20)

        self.apply_time_edit = QLineEdit(self.date_group)
        self.apply_time_edit.setGeometry(68, 28, 300, 22)
        self.apply_time_edit.setPlaceholderText("留空=当前")
        self.apply_time_edit.setToolTip("输入8位日期如20260816；留空则使用系统当前日期")
        self.apply_time_edit.editingFinished.connect(self._save_date_inputs)

        lbl_accept = QLabel("受理时间：", self.date_group)
        lbl_accept.setGeometry(8, 56, 60, 20)

        self.accept_time_edit = QLineEdit(self.date_group)
        self.accept_time_edit.setGeometry(68, 54, 300, 22)
        self.accept_time_edit.setPlaceholderText("留空=当前")
        self.accept_time_edit.setToolTip("输入8位日期如20260816；留空则使用系统当前日期")
        self.accept_time_edit.editingFinished.connect(self._save_date_inputs)

        lbl_visit = QLabel("就诊时间：", self.date_group)
        lbl_visit.setGeometry(8, 82, 60, 20)

        self.visit_time_edit = QLineEdit(self.date_group)
        self.visit_time_edit.setGeometry(68, 80, 300, 22)
        self.visit_time_edit.setPlaceholderText("留空=当前")
        self.visit_time_edit.setToolTip("输入8位日期如20260816；留空则使用系统当前日期")
        self.visit_time_edit.editingFinished.connect(self._save_date_inputs)

        print("✅ API配置UI已创建")

    def _resolve_date_input(self, raw_value: str) -> str:
        """把输入框内容解析为日期字符串；为空时返回系统当前日期"""
        raw = (raw_value or "").strip()
        if not raw:
            return _date_now()
        # 8位纯数字 → YYYY年MM月DD日
        if raw.isdigit() and len(raw) == 8:
            return f"{raw[0:4]}年{raw[4:6]}月{raw[6:8]}日"
        return raw

    def _save_date_inputs(self):
        """把申请时间/受理时间/就诊时间保存到数据模型（空值用系统当前时间）"""
        if not hasattr(self, 'apply_time_edit') or not hasattr(self, 'accept_time_edit'):
            return
        self.set_data('申请时间', self._resolve_date_input(self.apply_time_edit.text()), 'case')
        self.set_data('受理时间', self._resolve_date_input(self.accept_time_edit.text()), 'case')
        if hasattr(self, 'visit_time_edit'):
            self.set_data('就诊时间', self._resolve_date_input(self.visit_time_edit.text()), 'case')

    def _load_saved_user_config(self):
        """加载已保存的用户配置到UI"""
        remembered = self.user_manager.get_remembered_user()
        if remembered:
            username = remembered.get("username", "")
            api_key = remembered.get("api_key", "")
            # 填充用户下拉列表
            user_list = self.user_manager.get_user_list()
            self.api_user_combo.addItems(user_list)
            if username:
                idx = self.api_user_combo.findText(username)
                if idx >= 0:
                    self.api_user_combo.setCurrentIndex(idx)
                else:
                    self.api_user_combo.setCurrentText(username)
            if api_key:
                self.api_key_input.setText(api_key)
            print(f"✅ 已加载记住的用户: {username}")
        else:
            # 至少填充用户列表
            user_list = self.user_manager.get_user_list()
            self.api_user_combo.addItems(user_list)
            print("ℹ️ 没有记住的用户")

    def _get_current_username(self) -> str:
        """获取当前用户名"""
        if hasattr(self, 'api_user_combo'):
            return self.api_user_combo.currentText().strip()
        return "未登录用户"

    def _toggle_config_panel(self):
        """展开/收起用户配置面板"""
        visible = self.api_group.isVisible()
        self.api_group.setVisible(not visible)
        arrow = "▼" if not visible else "⚙"
        self.config_toggle_btn.setText(arrow)

    def _update_api_status(self):
        """更新API状态标签"""
        if not hasattr(self, 'api_status_label'):
            return
        if self.ai_service:
            self.api_status_label.setText("✅ AI已就绪")
            self.api_status_label.setStyleSheet("color: green; font-weight: bold;")
            self.top_status_label.setText("✅ AI已就绪")
            self.top_status_label.setStyleSheet("color: green; background: transparent; border: none;")
        else:
            username = self._get_current_username()
            if not username or username == "未登录用户":
                msg = "⚠️ 请配置用户名和API密钥"
                self.api_status_label.setText(msg)
                self.api_status_label.setStyleSheet("color: orange;")
                self.top_status_label.setText(msg)
                self.top_status_label.setStyleSheet("color: orange; background: transparent; border: none;")
            else:
                msg = "⚠️ API密钥未配置，AI功能不可用"
                self.api_status_label.setText(msg)
                self.api_status_label.setStyleSheet("color: orange;")
                self.top_status_label.setText(msg)
                self.top_status_label.setStyleSheet("color: orange; background: transparent; border: none;")

    def _on_save_api_config(self):
        """保存API配置"""
        username = self.api_user_combo.currentText().strip()
        api_key = self.api_key_input.text().strip()
        remember = self.api_remember_cb.isChecked()

        if not username:
            QMessageBox.warning(self, "提示", "请输入用户名")
            return

        # 保存到UserManager
        api_url = "https://api.deepseek.com"
        success = self.user_manager.save_user_config(
            username=username,
            api_url=api_url,
            api_key=api_key,
            remember_me=remember,
            service="DeepSeek"
        )

        if success:
            # 更新下拉列表
            if self.api_user_combo.findText(username) < 0:
                self.api_user_combo.addItem(username)
                self.api_user_combo.setCurrentText(username)

            # 更新数据模型
            self.data_model.output_config['用户名'] = username

            # 重新初始化AI服务
            self.init_ai_service()

            self._set_status(f"配置已保存 - 用户: {username}", "green")
            print(f"✅ API配置已保存: 用户={username}, 密钥={'已设置' if api_key else '未设置'}")
        else:
            QMessageBox.critical(self, "错误", "保存配置失败，请重试")

    def _on_user_combo_changed(self, text):
        """用户名下拉框变化时自动加载对应的API密钥"""
        if not text or not text.strip():
            return
        username = text.strip()
        user_data = self.user_manager.users_data.get('users', {}).get(username, {})
        if user_data:
            api_key = user_data.get('api_key', '')
            remember = user_data.get('remember_me', True)
            self.api_key_input.setText(api_key)
            self.api_remember_cb.setChecked(remember)
            if api_key:
                print(f"✅ 已加载用户 '{username}' 的API配置")
            else:
                print(f"ℹ️ 用户 '{username}' 未配置API密钥")

    def init_ai_service(self):
        """初始化AI服务 - 从UserManager读取配置"""
        try:
            # 从UserManager获取当前用户的API配置
            username = self._get_current_username()
            if not username:
                print("⚠️ 未找到用户配置，AI功能将不可用")
                self.ai_service = None
                self._update_api_status()
                return

            user_data = self.user_manager.users_data.get('users', {}).get(username, {})
            api_key = user_data.get('api_key', '')
            api_url = user_data.get('api_url', 'https://api.deepseek.com')

            # 检查配置是否完整
            if not api_key or not api_url:
                print("⚠️ API配置不完整，AI功能将不可用")
                print(f"  API地址: {api_url if api_url else '未设置'}")
                print(f"  API密钥: {'已设置' if api_key else '未设置'}")
                self.ai_service = None
                self._update_api_status()
                return

            print(f"✅ 使用用户 '{username}' 的API配置初始化AI服务")
            print(f"  API地址: {api_url}")
            print(f"  API密钥前8位: {api_key[:8]}...")

            # 创建AI服务实例
            self.ai_service = AIService(api_key, api_url)
            print("✅ AI服务初始化成功")
            self._update_api_status()

        except Exception as e:
            print(f"❌ AI服务初始化失败: {e}")
            import traceback
            traceback.print_exc()
            self.ai_service = None
            self._update_api_status()

    def ai_review_document(self):
        """AI审查文档"""
        try:
            print("=" * 50)
            print("🔄 AI审查开始执行...")

            # 检查AI服务
            if not self.ai_service:
                print("❌ AI服务未初始化")
                QMessageBox.warning(self, "AI审查", "请先配置API密钥。")
                self.init_ai_service()
                return

            print("✅ AI服务检查通过")

            # 检查案件文件夹
            if not self.current_case_folder:
                print("❌ 当前案件文件夹为空")
                QMessageBox.warning(self, "AI审查", "请先保存案件信息。")
                return

            print(f"📁 案件文件夹: {self.current_case_folder}")

            # 查找本人笔录文件
            person_files = []
            for file in os.listdir(self.current_case_folder):
                print(f"📄 检查文件: {file}")
                if "本人" in file and file.endswith('.docx') and "审批表" not in file:
                    person_files.append(file)
                    print(f"✅ 找到本人笔录: {file}")

            if not person_files:
                print("❌ 未找到本人笔录文件")
                QMessageBox.warning(self, "AI审查", "未找到本人笔录文件。")
                return

            print(f"✅ 找到{len(person_files)}个本人笔录文件")

            # 使用第一个找到的本人笔录
            file_path = os.path.join(self.current_case_folder, person_files[0])
            print(f"📄 使用文件路径: {file_path}")

            # 检查文件是否存在
            if not os.path.exists(file_path):
                print(f"❌ 文件不存在")
                QMessageBox.warning(self, "AI审查", f"文件不存在: {file_path}")
                return

            print("✅ 文件存在")

            # ======================
            # 在这里创建进度对话框和AI工作线程
            # ======================

            # 创建进度对话框
            self.progress_dialog = QProgressDialog("正在分析文档...", "取消", 0, 100, self)
            self.progress_dialog.setWindowTitle("AI审查")
            self.progress_dialog.setWindowModality(Qt.WindowModal)

            # 创建AI工作线程（现在file_path已经定义）
            self.ai_worker = AIWorker(self.ai_service, file_path)

            # 连接信号到统一处理方法
            self.ai_worker.finished.connect(
                lambda result: self._handle_ai_result(result=result)
            )
            self.ai_worker.error.connect(
                lambda error: self._handle_ai_result(error=error)
            )
            self.ai_worker.progress.connect(
                lambda msg, value: self.progress_dialog.setLabelText(f"{msg} ({value}%)")
            )

            # 进度对话框取消
            self.progress_dialog.canceled.connect(
                lambda: self._handle_ai_result(canceled=True)
            )

            # 显示进度对话框并启动线程
            self.progress_dialog.show()
            self.ai_worker.start()

        except Exception as e:
            print(f"🔥 整体审查过程异常: {str(e)}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "AI审查错误", f"审查失败: {str(e)}")

    def show_ai_review_result(self, review_result):
        """显示AI审查结果 - 带问题选择功能"""
        print(f"🖥️ 显示审查结果，结果类型: {type(review_result)}")

        # 处理不同的结果格式
        if isinstance(review_result, str):
            result_text = review_result
        elif isinstance(review_result, dict):
            if "结果" in review_result:
                result_text = review_result["结果"]
            elif "错误信息" in review_result:
                result_text = f"错误: {review_result['错误信息']}"
            elif "原始回复" in review_result:
                result_text = review_result["原始回复"]
            else:
                result_text = str(review_result)
        else:
            result_text = str(review_result)

        print(f"📝 要解析的文本长度: {len(result_text)}")

        # 解析AI结果
        parsed_result = self.parse_ai_result(result_text)

        # 创建对话框
        dialog = QDialog(self)
        dialog.setWindowTitle("AI法律审查结果")
        dialog.resize(700, 600)

        layout = QVBoxLayout()

        # 标题
        title = QLabel("AI法律审查报告")
        title.setStyleSheet("font-size: 16px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # 创建标签页
        tab_widget = QTabWidget()

        # 标签1：审查结果
        review_tab = QWidget()
        review_layout = QVBoxLayout()

        review_label = QLabel("审查结果分析：")
        review_label.setStyleSheet("font-weight: bold;")
        review_layout.addWidget(review_label)

        # 审查结果显示区域
        review_text_edit = QTextEdit()
        review_text_edit.setReadOnly(True)
        review_text_edit.setPlainText(parsed_result["审查结果"])
        review_text_edit.setMinimumHeight(300)
        review_layout.addWidget(review_text_edit)

        review_tab.setLayout(review_layout)
        tab_widget.addTab(review_tab, "审查结果")

        # 标签2：缺失问题（如果有）
        if parsed_result["缺失问题"]:
            questions_tab = QWidget()
            questions_layout = QVBoxLayout()

            questions_label = QLabel(f"发现 {len(parsed_result['缺失问题'])} 个缺失问题，请勾选需要添加到笔录的问题：")
            questions_label.setStyleSheet("font-weight: bold; color: #e74c3c;")
            questions_layout.addWidget(questions_label)

            # 创建问题列表（带复选框）
            self.question_list_widget = QListWidget()

            for question in parsed_result["缺失问题"]:
                item = QListWidgetItem(question)
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                item.setCheckState(Qt.Unchecked)  # 默认未选中
                self.question_list_widget.addItem(item)

            questions_layout.addWidget(self.question_list_widget)

            # 全选/全不选按钮
            select_buttons_layout = QHBoxLayout()

            btn_select_all = QPushButton("全选")
            btn_select_all.clicked.connect(lambda: self.select_all_questions(True))

            btn_select_none = QPushButton("全不选")
            btn_select_none.clicked.connect(lambda: self.select_all_questions(False))

            select_buttons_layout.addWidget(btn_select_all)
            select_buttons_layout.addWidget(btn_select_none)
            select_buttons_layout.addStretch()

            questions_layout.addLayout(select_buttons_layout)

            questions_tab.setLayout(questions_layout)
            tab_widget.addTab(questions_tab, f"缺失问题 ({len(parsed_result['缺失问题'])})")

        layout.addWidget(tab_widget)

        # 按钮区域
        button_layout = QHBoxLayout()

        # 插入到笔录按钮（只在有问题时显示）
        if parsed_result.get("缺失问题"):
            btn_insert = QPushButton("插入选中问题到笔录")
            btn_insert.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold;")
            btn_insert.clicked.connect(lambda: self.insert_selected_questions(dialog))
            button_layout.addWidget(btn_insert)

        btn_copy = QPushButton("复制结果")
        btn_copy.clicked.connect(lambda: self.copy_to_clipboard(parsed_result["审查结果"]))

        btn_save = QPushButton("保存报告")
        btn_save.clicked.connect(lambda: self.save_ai_report(parsed_result))

        btn_close = QPushButton("关闭")
        btn_close.clicked.connect(dialog.close)

        button_layout.addWidget(btn_copy)
        button_layout.addWidget(btn_save)
        button_layout.addWidget(btn_close)

        layout.addLayout(button_layout)

        dialog.setLayout(layout)

        # 保存解析结果到对话框对象，以便后续使用
        dialog.parsed_result = parsed_result

        dialog.exec_()

    def copy_to_clipboard(self, text, parent=None):
        """复制文本到剪贴板"""
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(parent or self, "复制成功", "已复制到剪贴板")

    def _copy_statement(self):
        """复制案件申请陈述"""
        text = self.statement_edit.toPlainText().strip()
        if text:
            self.copy_to_clipboard(text)
        else:
            QMessageBox.information(self, "提示", "案件申请陈述为空")

    def _copy_material(self):
        """复制材料分类"""
        if hasattr(self, 'material_list'):
            self.material_list.copy_to_clipboard()
        QMessageBox.information(self, "提示", "材料列表已复制到剪贴板")

    def save_ai_report(self, parsed_result):
        """保存AI审查报告"""
        if not self.current_case_folder:
            QMessageBox.warning(self, "提示", "请先保存案件信息")
            return

        person_name = self.get_data('本人姓名', '未知')
        timestamp = _timestamp_now()
        filename = f"{person_name}_AI审查报告_{timestamp}.txt"
        filepath = os.path.join(self.current_case_folder, filename)

        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write("AI法律审查报告\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"审查时间：{datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
                f.write(f"审查对象：{person_name}\n\n")

                f.write("【审查结果】\n")
                f.write(parsed_result.get("审查结果", "无审查结果") + "\n\n")

                if parsed_result.get("缺失问题"):
                    f.write("【缺失问题列表】\n")
                    for i, question in enumerate(parsed_result["缺失问题"], 1):
                        f.write(f"{i}. {question}\n")

            QMessageBox.information(self, "保存成功", f"报告已保存为：\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "保存失败", f"保存失败：{str(e)}")

    def generate_injury_notice(self):
        """生成工伤告知书：按案本号找目录 → 检查审批表 → 按JSON结论选模板渲染"""
        try:
            # 1. 读案本号
            case_number = self.lineEdit_2.text().strip()
            if not case_number:
                self._set_status('未找到案本号', 'red')
                QMessageBox.warning(self, "提示", "请先输入或生成案本号")
                return

            # 2. 用案本号找案件文件夹
            case_folder = os.path.join(self.BASE_PATH, case_number)
            if not os.path.exists(case_folder):
                self._set_status('未找到案件目录', 'red')
                QMessageBox.warning(self, "提示", f"未找到案本号对应的案件目录：\n{case_folder}")
                return

            # 3. 检查审批表文件
            approval_files = [f for f in os.listdir(case_folder)
                              if f.endswith('.docx') and '审批表' in f]
            if not approval_files:
                self._set_status('目录下没有审批表文件，请先生成案件审批表', 'red')
                QMessageBox.warning(self, "提示", "目录下没有审批表文件，请先生成案件审批表。")
                return

            # 多个审批表：弹窗选择
            selected_file = approval_files[0]
            if len(approval_files) > 1:
                choice, ok = QInputDialog.getItem(self, "选择审批表",
                                                  "目录下有多个审批表，请选择：",
                                                  approval_files, 0, False)
                if not ok or not choice:
                    self._set_status('已取消', 'black')
                    return
                selected_file = choice
            print(f"✅ 使用审批表: {selected_file}")

            # 4. 用 JSON 数据生成字典
            case_obj = self._load_cases_data().get(case_number)
            if not case_obj:
                self._set_status('未找到该案本号的案件数据', 'red')
                QMessageBox.warning(self, "提示", f"未在数据中找到案本号：{case_number}")
                return
            template_data = self._build_notice_template_data(case_obj)

            # 5. 按 JSON 结论选模板
            conclusion = case_obj.get('conclusion', '')
            if conclusion == "不予认定":
                template_name = '不予工伤认定告知书（样本）.docx'
            else:
                template_name = '工伤认定告知书（样本）.docx'
            template_path = str(path_utils.get_document_template_path(template_name))
            if not os.path.exists(template_path):
                self._set_status(f'模板文件不存在: {template_name}', 'red')
                QMessageBox.critical(self, "错误", f"找不到模板文件:\n{template_path}")
                return

            # 6. 渲染
            self._set_status('正在生成工伤告知书...', 'black')
            QApplication.processEvents()
            from docxtpl import DocxTemplate
            word = DocxTemplate(template_path)
            word.render(template_data)

            # 7. 保存 + 打开
            name = case_obj.get('name', '') or '职工'
            if conclusion == "不予认定":
                notice_file_name = f"{name}不予工伤认定告知书.docx"
            else:
                notice_file_name = f"{name}工伤认定告知书.docx"
            target_path = os.path.join(case_folder, notice_file_name)
            word.save(target_path)
            print(f"✅ 工伤告知书保存到: {target_path}")

            success, message = self.file_service.open_document(target_path)
            if success:
                self._set_status('工伤告知书生成成功', 'green')
                QMessageBox.information(self, "成功", f"工伤告知书已生成:\n{notice_file_name}")
            else:
                self._set_status(f'工伤告知书生成成功，但打开失败: {message}', 'orange')
                QMessageBox.information(self, "成功",
                                        f"工伤告知书已生成:\n{notice_file_name}\n\n但打开失败: {message}")

        except Exception as e:
            print(f"❌ 生成工伤告知书失败: {e}")
            import traceback
            traceback.print_exc()
            self._set_status(f'生成工伤告知书失败: {str(e)}', 'red')
            QMessageBox.critical(self, "错误", f"生成工伤告知书失败:\n{str(e)}")

    def _build_notice_template_data(self, case_obj: dict) -> dict:
        """构建工伤告知书模板字典（从案件 JSON 数据）"""
        current_date = _date_now()
        return {
            '本人姓名': case_obj.get('name', ''),
            '本人身份证号': case_obj.get('id_card', ''),
            '用人单位': case_obj.get('labor_unit', ''),  # 用人单位（签合同单位）
            '受伤经过': case_obj.get('injury_process', case_obj.get('injury_description', '详见谈话笔录')),
            '医疗证明': case_obj.get('medical_conclusion', ''),
            '申请时间': case_obj.get('apply_time', current_date),
            '受理时间': case_obj.get('accept_time', current_date),
            '当前时期': current_date + _time_now(),
            '告知日期': current_date,
            '受理编号': case_obj.get('case_id', ''),
            '案本号': case_obj.get('case_id', ''),
        }

    def _apply_ui_settings(self):
        """应用UI设置"""
        try:
            ui_settings = self.config_service.get_ui_settings()

            # 设置字体
            font = QFont(ui_settings.font_family, ui_settings.font_size)
            self.setFont(font)

            # 注意：窗口大小由 _setup_api_config_ui 固定为 870x740，此处不再覆盖
        except Exception as e:
            print(f"⚠️ 应用UI设置失败: {e}")

    def clear_role_fields(self):
        """
        清空当前角色的字段
        """
        role = self.get_current_role_type()

        print(f"🧹 清空{role}字段")

        # 清空输入控件
        self.clear_fields()

        # 清空该角色的数据模型与模板缓存，避免多名证人/人员数据串用
        self._clear_role_data(role)

        # 当角色切换时，更新按钮状态
        self.on_role_changed()

        # 切换回本人时清空案本号，输入新姓名后自动生成
        if role == "本人":
            self.lineEdit_2.clear()

        # 多证人：切换到证人时显示下拉框并加载；切走时保存当前证人并隐藏
        if role == "证人":
            self._show_witness_ui()
        else:
            self._sync_form_to_current_witness()
            self._hide_witness_ui()

    def _clear_role_data(self, role: str):
        """清除指定角色在数据模型与模板字典中的所有数据，确保多人数据一一对应"""
        if role not in ("本人", "证人", "法人"):
            return
        self.data_model.clear_role_data(role)
        for key in [k for k in list(self._template_dict.keys()) if k.startswith(role)]:
            self._template_dict.pop(key, None)
        self.var_manager.clear_cache()

    # ========================================================================
    # 多证人管理
    # ========================================================================

    def _setup_witness_ui(self):
        """创建证人编号下拉框与「添加证人」按钮（代码创建，放在左栏底部空位）"""
        self.witness_label = QLabel("证人编号：", self)
        self.witness_label.setObjectName("witness_label")
        self.witness_label.setGeometry(70, 765, 70, 20)

        self.witness_combo = QComboBox(self)
        self.witness_combo.setObjectName("witness_combo")
        self.witness_combo.setGeometry(140, 763, 180, 24)
        self.witness_combo.currentIndexChanged.connect(self._on_witness_selected)

        self.add_witness_btn = QPushButton("添加证人", self)
        self.add_witness_btn.setObjectName("add_witness_btn")
        self.add_witness_btn.setGeometry(330, 763, 80, 24)
        self.add_witness_btn.clicked.connect(self._add_witness)

        self.witness_label.hide()
        self.witness_combo.hide()
        self.add_witness_btn.hide()

    def _show_witness_ui(self):
        """切换到证人角色时显示下拉框，并加载当前/首位证人"""
        self.witness_label.show()
        self.witness_combo.show()
        self.add_witness_btn.show()

        if not self.data_model.witnesses:
            if self.witness_combo.count() == 0:
                self.witness_combo.blockSignals(True)
                self.witness_combo.addItem("（暂无证人，请添加）")
                self.witness_combo.blockSignals(False)
            return

        if self.data_model.current_witness_index < 0:
            self.data_model.current_witness_index = 0
        self._refresh_witness_combo()
        self._sync_current_witness_to_form()

    def _hide_witness_ui(self):
        self.witness_label.hide()
        self.witness_combo.hide()
        self.add_witness_btn.hide()

    def _current_witness(self) -> Optional[Dict[str, Any]]:
        idx = self.data_model.current_witness_index
        if 0 <= idx < len(self.data_model.witnesses):
            return self.data_model.witnesses[idx]
        return None

    def _refresh_witness_combo(self):
        """重建下拉框内容（阻塞信号，避免触发切换逻辑）"""
        self.witness_combo.blockSignals(True)
        self.witness_combo.clear()
        for w in self.data_model.witnesses:
            name = w.get("姓名") or "未命名"
            self.witness_combo.addItem(f"{w.get('序号', '')} · {name}")
        if 0 <= self.data_model.current_witness_index < self.witness_combo.count():
            self.witness_combo.setCurrentIndex(self.data_model.current_witness_index)
        self.witness_combo.blockSignals(False)

    def _mirror_witness_to_flat(self, w: Dict[str, Any]):
        """把证人 dict 同步到扁平 证人* 键（供模板/AI 使用）"""
        mapping = [
            ("姓名", "证人姓名"),
            ("身份证号", "证人身份证号"),
            ("身份证地址", "证人身份证地址"),
            ("手机号", "证人手机号"),
            ("岗位", "证人岗位"),
            ("性别", "证人性别"),
            ("年龄", "证人年龄"),
        ]
        for src, dst in mapping:
            val = w.get(src, "")
            if val:
                self.set_data(dst, val, 'basic')

    def _sync_form_to_current_witness(self):
        """把表单内容写回当前证人，并同步扁平 证人* 键"""
        if self.get_current_role_type() != "证人":
            return
        w = self._current_witness()
        if w is None:
            return

        w["姓名"] = self.name_pane.text().strip()
        w["身份证号"] = self.idnumer_pane.text().strip()
        w["身份证地址"] = self.textEdit.toPlainText().strip()
        w["手机号"] = self.lineEdit_4.text().strip()
        w["岗位"] = self.lineEdit_5.text().strip()
        w["性别"] = self.lineEdit.text().strip()
        w["年龄"] = self.age_pane.text().strip()

        self._mirror_witness_to_flat(w)

        # 更新下拉框当前项的显示（不重建，避免选中状态被打乱）
        idx = self.data_model.current_witness_index
        if 0 <= idx < self.witness_combo.count():
            self.witness_combo.blockSignals(True)
            self.witness_combo.setItemText(idx, f"{w.get('序号', '')} · {w.get('姓名') or '未命名'}")
            self.witness_combo.blockSignals(False)

        self.var_manager.clear_cache()

    def _sync_current_witness_to_form(self):
        """把当前证人数据填入表单 + 扁平 证人* 键"""
        w = self._current_witness()
        if w is None:
            return
        self.name_pane.setText(w.get("姓名", ""))
        self.idnumer_pane.setText(w.get("身份证号", ""))
        self.textEdit.setPlainText(w.get("身份证地址", ""))
        self.lineEdit_4.setText(w.get("手机号", ""))
        self.lineEdit_5.setText(w.get("岗位", ""))
        self.lineEdit.setText(w.get("性别", ""))
        self.age_pane.setText(str(w.get("年龄", "")) if w.get("年龄") else "")

        self._mirror_witness_to_flat(w)
        self.var_manager.clear_cache()

    def _ensure_current_witness(self):
        """生成证人笔录前调用：确保 current_witness_index 指向一个有效证人"""
        if self.get_current_role_type() != "证人":
            return
        if 0 <= self.data_model.current_witness_index < len(self.data_model.witnesses):
            return
        if self.data_model.witnesses:
            # 已有证人但索引无效（如 F2 填充后未设索引）→ 指向第一个
            self.data_model.current_witness_index = 0
            self._refresh_witness_combo()
            return
        new_index = len(self.data_model.witnesses)
        new_witness = {
            "序号": _witness_label(new_index + 1),
            "姓名": "", "身份证号": "", "身份证地址": "",
            "手机号": "", "岗位": "", "性别": "", "年龄": "",
        }
        self.data_model.witnesses.append(new_witness)
        self.data_model.current_witness_index = new_index
        self._refresh_witness_combo()

    def _add_witness(self):
        """添加一个新证人，自动编号为 证人一/证人二/…"""
        self._sync_form_to_current_witness()  # 先保存当前证人的编辑

        new_index = len(self.data_model.witnesses)
        new_witness = {
            "序号": _witness_label(new_index + 1),
            "姓名": "", "身份证号": "", "身份证地址": "",
            "手机号": "", "岗位": "", "性别": "", "年龄": "",
        }
        self.data_model.witnesses.append(new_witness)
        self.data_model.current_witness_index = new_index

        self._refresh_witness_combo()
        self._sync_current_witness_to_form()  # 清空表单，准备录入新证人
        self._save_witnesses()

    def _on_witness_selected(self, index):
        """切换选中的证人"""
        if index < 0 or index >= len(self.data_model.witnesses):
            return
        if self.data_model.current_witness_index != index:
            self._sync_form_to_current_witness()  # 保存上一个证人
        self.data_model.current_witness_index = index
        self._sync_current_witness_to_form()

    def _save_witnesses(self):
        """把证人数据写回 cases_data.json 的当前案件（单一数据源，不再单独存证人信息.json）"""
        case_id = (self.current_case_id or '').strip() or self.lineEdit_2.text().strip()
        if not case_id:
            return
        try:
            cases = self._load_cases_data()
            case_obj = cases.get(case_id)
            if case_obj is None:
                return
            case_obj['witnesses'] = list(self.data_model.witnesses)
            self._save_cases_data(cases)
            print(f"✅ 证人信息已保存到案件数据: {len(self.data_model.witnesses)} 位证人")
        except Exception as e:
            print(f"保存证人信息失败: {e}")

    def _load_witnesses(self):
        """从 cases_data.json 的当前案件读取证人数据"""
        case_id = (self.current_case_id or '').strip() or self.lineEdit_2.text().strip()
        if not case_id:
            return
        try:
            case_obj = self._load_cases_data().get(case_id)
            if case_obj is None:
                return
            self.data_model.witnesses = list(case_obj.get('witnesses', []))
            self.data_model.current_witness_index = -1
            self._refresh_witness_combo()
            print(f"✅ 加载证人信息: {len(self.data_model.witnesses)} 位证人")
        except Exception as e:
            print(f"加载证人信息失败: {e}")

    def setup_logging(self):
        """简化日志系统"""
        self.log_warning = lambda msg: print(f"警告: {msg}")
        self.log_error = lambda msg: print(f"错误: {msg}")
    def on_case_type_changed(self):
        """当案件类型选择改变时调用"""
        is_death_case = self.death_case_checkbox.isChecked()
        is_personal = self.personal_application_checkbox.isChecked()

        self.set_data('案件性质', "工亡案件" if is_death_case else "工伤案件", 'case')
        self.set_data('申请类型', "个人申请" if is_personal else "单位申请", 'case')
        self.update_case_type_hint()

    def update_case_type_hint(self):
        """更新案件类型提示信息"""
        is_death_case = self.death_case_checkbox.isChecked()
        is_personal = self.personal_application_checkbox.isChecked()

        hint_text = "当前案件类型: "
        if is_death_case and is_personal:
            hint_text += "个人申请的工亡案件"
        elif is_death_case:
            hint_text += "单位申请的工亡案件"
        elif is_personal:
            hint_text += "个人申请的工伤案件"
        else:
            hint_text += "单位申请的工伤案件"

        hint_label = self.findChild(QLabel, "labelCaseHint")
        if hint_label:
            hint_label.setText(hint_text)

    def init_combobox_data(self):
        """初始化组合框数据（简化版）"""
        try:
            # 使用 path_utils 的数据路径
            from path_utils import path_utils

            doc_template_dir = path_utils.get_document_template_path("")
            print(f"🔍 文书模板目录: {doc_template_dir}")

            # 用工单位 - 使用文书模板目录
            company_file = str(path_utils.get_document_template_path('用工单位汇总.xlsx'))
            print(f"🔍 用工单位文件: {company_file}")

            if os.path.exists(company_file):
                try:
                    file = pd.read_excel(company_file)
                    self.items_list = file['用工单位汇总'].tolist()
                    print(f"✅ 加载用工单位: {len(self.items_list)}个")
                    if self.items_list:
                        print(f"   示例: {self.items_list[:3]}")
                except Exception as e:
                    print(f"❌ 读取用工单位文件失败: {e}")
                    self.items_list = ['公司A', '公司B', '公司C']  # 默认数据
            else:
                print("⚠️ 用工单位文件不存在，创建默认文件")
                self.items_list = ['公司A', '公司B', '公司C']
                # 创建默认文件
                try:
                    df = pd.DataFrame(self.items_list, columns=['用工单位汇总'])
                    df.to_excel(company_file, index=False)
                    print(f"✅ 创建默认用工单位文件")
                except Exception as e:
                    print(f"❌ 创建用工单位文件失败: {e}")

            # 用人单位 - 使用文书模板目录
            employer_file = str(path_utils.get_document_template_path('用人单位汇总.xlsx'))
            print(f"🔍 用人单位文件: {employer_file}")

            if os.path.exists(employer_file):
                try:
                    file1 = pd.read_excel(employer_file)
                    self.items_list1 = file1['用人单位汇总'].tolist()
                    print(f"✅ 加载用人单位: {len(self.items_list1)}个")
                except Exception as e:
                    print(f"❌ 读取用人单位文件失败: {e}")
                    self.items_list1 = ['用人单位A', '用人单位B']
            else:
                print("⚠️ 用人单位文件不存在，创建默认文件")
                self.items_list1 = ['用人单位A', '用人单位B']
                try:
                    df = pd.DataFrame(self.items_list1, columns=['用人单位汇总'])
                    df.to_excel(employer_file, index=False)
                    print(f"✅ 创建默认用人单位文件")
                except Exception as e:
                    print(f"❌ 创建用人单位文件失败: {e}")

            # 工地名称 - 使用文书模板目录
            site_file = str(path_utils.get_document_template_path('工地名称汇总.xlsx'))
            print(f"🔍 工地名称文件: {site_file}")

            if os.path.exists(site_file):
                try:
                    file2 = pd.read_excel(site_file)
                    self.items_list2 = file2['工地名称汇总'].tolist()
                    print(f"✅ 加载工地名称: {len(self.items_list2)}个")
                except Exception as e:
                    print(f"❌ 读取工地名称文件失败: {e}")
                    self.items_list2 = ['工地A', '工地B']
            else:
                print("⚠️ 工地名称文件不存在，创建默认文件")
                self.items_list2 = ['工地A', '工地B']
                try:
                    df = pd.DataFrame(self.items_list2, columns=['工地名称汇总'])
                    df.to_excel(site_file, index=False)
                    print(f"✅ 创建默认工地名称文件")
                except Exception as e:
                    print(f"❌ 创建工地名称文件失败: {e}")

        except Exception as e:
            print(f"❌ 初始化组合框数据失败: {e}")
            import traceback
            traceback.print_exc()

            # 设置默认数据
            self.items_list = ['公司A', '公司B', '公司C']
            self.items_list1 = ['用人单位A', '用人单位B']
            self.items_list2 = ['工地A', '工地B']
            print("✅ 使用默认数据")

    def on_id_input_finished(self):
        """当身份证输入框完成编辑时自动计算年龄和性别"""
        try:
            role = self.get_current_role_type()
            idcard = self.idnumer_pane.text().strip()

            if not idcard or len(idcard) not in (15, 18):
                return

            self.set_data(f"{role}身份证号", idcard, 'basic')
            self.process_id_info(role)
            self.calculate_age_from_id(role)

        except Exception as e:
            import traceback
            traceback.print_exc()

    def save_company(self):
        """保存用人单位到Excel（company_pane 现为用人单位）"""
        new_item = self.company_pane.currentText().strip()
        if new_item and new_item not in self.items_list1:
            self.items_list1 = self.file_service.save_to_excel(
                "",
                '用人单位汇总.xlsx',
                '用人单位汇总',
                new_item,
                self.items_list1
            )
            self.init_combobox(self.company_pane, self.items_list1)
            print(f"💾 保存用人单位: {new_item}")

    def save_construction_company(self):
        """保存用工单位到Excel（construction_company 现为用工单位）"""
        new_item = self.construction_company.currentText().strip()
        if new_item and new_item not in self.items_list:
            self.items_list = self.file_service.save_to_excel(
                "",
                '用工单位汇总.xlsx',
                '用工单位汇总',
                new_item,
                self.items_list
            )
            self.init_combobox(self.construction_company, self.items_list)
            print(f"💾 保存用工单位: {new_item}")

    def save_construction_plant(self):
        """保存工地名称到Excel"""
        new_item = self.construction_plant.currentText().strip()
        if new_item and new_item not in self.items_list2:
            self.items_list2 = self.file_service.save_to_excel(
                "",  # 空字符串
                '工地名称汇总.xlsx',
                '工地名称汇总',
                new_item,
                self.items_list2
            )
            self.init_combobox(self.construction_plant, self.items_list2)
            print(f"💾 保存工地名称: {new_item}")

    def init_combobox(self, combobox, items):
        """初始化组合框"""
        print(f"🔍 初始化 {combobox.objectName()}，数据长度: {len(items)}")

        combobox.clear()

        if items:
            for item in items:
                combobox.addItem(str(item))
            print(f"✅ 添加了 {len(items)} 个选项")
        else:
            print("⚠️ 没有数据可添加")
            combobox.addItem("暂无数据")

        combobox.setCurrentIndex(-1)  # 清空选择

        # 设置自动完成
        completer = QCompleter(items)
        completer.setFilterMode(Qt.MatchContains)
        completer.setCompletionMode(QCompleter.PopupCompletion)
        combobox.setCompleter(completer)

        print(f"✅ {combobox.objectName()} 初始化完成，当前项数: {combobox.count()}")

    def calculate_age_from_id(self, role):
        """根据身份证号计算年龄（使用DataService）"""
        try:
            idcard = self.get_data(f"{role}身份证号", "")
            if not idcard:
                return

            # 使用DataService计算年龄
            age = self.data_service.calculate_age_from_idcard(idcard)
            if age is None:
                self._set_status('身份证号格式错误', 'red', 'label_12')
                return

            # 设置年龄数据
            self.set_data(f"{role}年龄", age, 'basic')
            self.age_pane.setText(str(age))

            # 超龄检查（仅对本人）
            if role == "本人":
                gender = self.get_data(f"{role}性别", "")
                if (gender == "男" and age > 60) or (gender == "女" and age > 50):
                    self._set_status('此人已经超龄', 'red', 'label_12')
                else:
                    self._set_status('信息提示', 'black', 'label_12')

        except Exception as e:
            print(f"[calculate_age_from_id] 错误: {e}")
            self._set_status('年龄计算错误', 'red', 'label_12')

    def get_data(self, key: str, default: Any = None) -> Any:
        """统一的数据访问方法"""
        model_value = self._get_from_data_model(key)
        if model_value is not None:
            if key not in self._template_dict or self._template_dict[key] != model_value:
                self._template_dict[key] = model_value
            return model_value

        dict_value = self._template_dict.get(key)
        if dict_value is not None:
            return dict_value

        return default

    def set_data(self, key: str, value: Any, category: str = "auto") -> None:
        """统一的数据设置方法"""
        if value is None:
            value = ""

        if category == "auto":
            category = self._detect_data_category(key)

        self._store_to_data_model(key, value, category)
        self._template_dict[key] = value
        self._handle_special_keys(key, value)

    def _handle_special_keys(self, key: str, value: Any):
        """处理特殊键的同步逻辑"""
        if key == "当前时期":
            if not value:
                current_date = _date_now()
                current_time = _time_now()
                self.set_data(key, f"{current_date}{current_time}", 'output')

    def _detect_data_category(self, key: str) -> str:
        """自动检测数据类别"""
        role_prefixes = ['本人', '证人', '法人']
        for prefix in role_prefixes:
            if key.startswith(prefix):
                base_key = key[len(prefix):]
                return self._detect_base_category(base_key)
        return self._detect_base_category(key)

    def _detect_base_category(self, key: str) -> str:
        """检测基础键名的类别"""
        if key in ['姓名', '年龄', '性别', '身份证号', '身份证地址', '手机号', '岗位', '职务']:
            return 'basic'
        elif key in ['用工单位', '用人单位', '工地名称']:
            return 'company'
        elif key in ['案件性质', '申请类型', '案本号', '案件类型']:
            return 'case'
        elif key in ['当前日期', '当前时间', '当前时期']:
            return 'output'
        else:
            return 'investigation'

    def _get_from_data_model(self, key: str) -> Any:
        """从数据模型的各个部分获取数据

        注意：_store_to_data_model 存储时保留完整 key（含角色前缀），
        此处 first check 即可命中，不需要再去前缀查找。
        """
        if key in self.data_model.basic_info:
            return self.data_model.basic_info[key]

        for category in (
            self.data_model.company_info,
            self.data_model.case_info,
            self.data_model.investigation,
            self.data_model.output_config,
        ):
            if key in category:
                return category[key]

        return None

    def _store_to_data_model(self, key: str, value: Any, category: str) -> None:
        """存储数据到数据模型"""
        role_prefixes = ['本人', '证人', '法人']
        for prefix in role_prefixes:
            if key.startswith(prefix):
                self.data_model.basic_info[key] = value
                return

        category_map = {
            'basic': self.data_model.basic_info,
            'company': self.data_model.company_info,
            'case': self.data_model.case_info,
            'investigation': self.data_model.investigation,
            'output': self.data_model.output_config,
        }

        if category in category_map:
            category_map[category][key] = value
        else:
            self.data_model.basic_info[key] = value

    def clear_fields(self):
        """清空所有输入控件"""
        fields = [
            self.name_pane, self.age_pane, self.lineEdit,
            self.idnumer_pane, self.textEdit, self.lineEdit_4, self.lineEdit_5
        ]
        for field in fields:
            if isinstance(field, QLineEdit):
                field.clear()
            elif isinstance(field, QTextEdit):
                field.clear()

        # 清空右侧辅助面板
        if hasattr(self, 'statement_edit'):
            self.statement_edit.clear()
        if hasattr(self, 'material_list'):
            self.material_list.clear()

        self._set_status('信息提示', 'black', 'label_14')
        self._set_status('信息提示', 'black', 'label_12')
    def _read_all_transcripts(self, case_folder: str) -> str:
        """读取案件目录下所有谈话笔录（本人/证人/法人）全文，按文件名分隔"""
        try:
            if not case_folder or not os.path.exists(case_folder):
                return ""
            exclude_kw = ("审批表", "告知书", "通知书", "发送给AI")
            parts = []
            for fname in sorted(os.listdir(case_folder)):
                if not fname.endswith('.docx'):
                    continue
                if any(k in fname for k in exclude_kw):
                    continue
                fpath = os.path.join(case_folder, fname)
                try:
                    doc = Document(fpath)
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    if text.strip():
                        parts.append(f"=== {fname} ===\n{text}")
                except Exception:
                    continue
            if not parts:
                print("⚠️ 目录下没有可用的谈话笔录")
                return ""
            print(f"📚 读取全部笔录: {len(parts)} 份")
            return "\n\n".join(parts)
        except Exception as e:
            print(f"⚠️ 读取全部笔录失败: {e}")
            return ""

    def approve(self):
        """生成案件审批表 — 读取案本号 → JSON查数据 → 渲染模板"""
        try:
            # ── 1. 读取主界面上的案本号 ──
            case_number = self.lineEdit_2.text().strip()
            if not case_number:
                self._set_status('未找到案本号', 'red')
                QMessageBox.warning(self, "提示", "请先输入或生成案本号")
                return

            # ── 2. 在 cases_data.json 里查找对应数据 ──
            case_obj = self._load_cases_data().get(case_number)
            if not case_obj:
                self._set_status('JSON中未找到该案本号', 'red')
                QMessageBox.warning(self, "提示", f"未在数据中找到案本号：{case_number}")
                return
            # 用案件数据构建 case_data（字段名兼容原有索引字段，后面逻辑不变）
            case_data = {
                'case_id': case_obj.get('case_id', ''),
                'person_name': case_obj.get('name', ''),
                'company_name': case_obj.get('labor_unit', ''),
                'applicant_name': case_obj.get('applicant_name', ''),
                'regulation': _regulation_short_to_full(case_obj.get('proposed_article', '')),
                'folder_name': case_obj.get('folder_name', ''),
                'person_gender': case_obj.get('gender', ''),
                'id_card': case_obj.get('id_card', ''),
            }

            person_name = case_data.get('person_name', '')
            company_name = case_data.get('company_name', '')

            # ── 3. 申请人名称（从JSON读取，保存时已按申请类型算好）──
            applicant_name = case_data.get('applicant_name', '')
            self.set_data('申请人名称', applicant_name, 'case')

            # ── 4. 确定案件文件夹（笔录都保存在 BASE_PATH/<案本号>/ 下）──
            case_folder = os.path.join(self.BASE_PATH, case_number)
            if not os.path.exists(case_folder):
                # 兼容旧数据：JSON folder_name 或当前案件文件夹或兜底
                folder_name = case_data.get('folder_name', '')
                if folder_name:
                    case_folder = os.path.join(self.BASE_PATH, folder_name)
                elif self.current_case_folder and os.path.exists(self.current_case_folder):
                    case_folder = self.current_case_folder
                else:
                    case_folder = str(path_utils.get_storage_path(
                        f"{person_name}-工伤案件" if person_name else "未命名案件"
                    ))

            # ── 4.0 检查目录下是否已有审批表文件 ──
            existing_approval = None
            if os.path.isdir(case_folder):
                for fname in sorted(os.listdir(case_folder)):
                    if fname.endswith('.docx') and '审批表' in fname:
                        existing_approval = os.path.join(case_folder, fname)
                        break
            if existing_approval:
                msg = QMessageBox(self)
                msg.setWindowTitle("审批表已存在")
                msg.setText("该案件目录下已有审批表文件，如何处理？")
                view_btn = msg.addButton("查看已有审批表", QMessageBox.AcceptRole)
                new_btn = msg.addButton("新建一个审批表", QMessageBox.ActionRole)
                msg.exec_()
                if msg.clickedButton() == view_btn:
                    success, _ = self.file_service.open_document(existing_approval)
                    if success:
                        self._set_status('已打开已有审批表', 'green')
                    else:
                        self._set_status('打开已有审批表失败', 'orange')
                    return
                # 选择新建：继续走新建流程（文件名会自动加一）

            # ── 4.1 医疗结论 / 受伤经过（AI 分析全部笔录 → 认定/不予认定决策）──
            medical_conclusion = ''
            injury_process = ''
            conclusion = "予以认定"

            if not self.ai_service:
                # 无 AI：提醒并停止（不生成审批表）
                QMessageBox.warning(self, "提示", "未配置AI，无法分析认定/不予认定，不能生成案件审批表。")
                return

            # 解析适用条款情形，供 AI 突出关键证据要素
            from case_classifier import CaseClassifier
            regulation = case_data.get('regulation', '')
            reg_desc = ""
            reg_elements = []
            for _idx, _reg in CaseClassifier.REGULATIONS.items():
                if _reg.get("text") == regulation:
                    reg_desc = _reg.get("desc", "")
                    reg_elements = _reg.get("elements", [])
                    break

            # 读取该案全部谈话笔录（本人/证人/法人）
            all_text = self._read_all_transcripts(case_folder)
            if not all_text:
                QMessageBox.warning(self, "提示", "该案件目录下没有找到谈话笔录，无法分析。\n请先生成本人/证人/法人谈话笔录。")
                return

            # AI 分析 → 认定/不予认定 偏向
            self._set_status('正在AI分析笔录...', 'black')
            QApplication.processEvents()
            analysis = self.ai_service.analyze_approval_transcripts(
                all_text,
                case_id=case_number,
                regulation_text=regulation,
                regulation_desc=reg_desc,
                regulation_elements=reg_elements,
            )
            if not analysis:
                QMessageBox.warning(self, "提示", "AI 分析失败，请重试。")
                return

            # 展示对话框（认定工伤 / 不予认定工伤 / 保存）
            dlg = ApprovalDecisionDialog(analysis, self)
            dlg.exec_()
            choice = dlg.get_choice()

            if choice == "保存":
                # 保存分析结果到 JSON，不生成审批表
                self._update_case_field(case_number,
                                        analysis_result=analysis.get("分析", ""),
                                        conclusion_bias=analysis.get("偏向", ""))
                self._set_status('已保存 AI 分析结果', 'green')
                return

            # 生成「调查核实情况」段落（受伤经过）
            self._set_status('正在AI生成调查核实情况...', 'black')
            QApplication.processEvents()
            gen_result = self.ai_service.generate_injury_and_conclusion(
                all_text,
                regulation_text=regulation,
                regulation_desc=reg_desc,
                regulation_elements=reg_elements,
            )
            if gen_result:
                injury_process = gen_result.get("受伤经过", "") or ""
                medical_conclusion = gen_result.get("诊断结论", "") or ""
            if not injury_process:
                QMessageBox.warning(self, "提示", "生成调查核实情况失败，请重试。")
                return

            if choice == "不予认定":
                # 不予认定：在调查核实后面追加关键理由
                reasons = analysis.get("关键理由", []) or []
                if reasons:
                    injury_process += "\n不予认定的关键理由：\n" + "\n".join(
                        f"{i + 1}. {r}" for i, r in enumerate(reasons))
                conclusion = "不予认定"
            else:
                conclusion = "予以认定"

            # 回写 JSON：受伤经过 / 诊断结论 / 结论
            self.set_data('本人受伤经过', injury_process, 'investigation')
            self.set_data('认定结论', conclusion, 'case')
            self._update_case_field(case_number,
                                    injury_process=injury_process,
                                    medical_conclusion=medical_conclusion,
                                    conclusion=conclusion)

            # 医疗结论写入数据模型
            if medical_conclusion:
                self.data_model.investigation['医院诊断'] = medical_conclusion
                self.data_model.investigation['医疗结论'] = medical_conclusion

            # ── 4.2 用JSON数据构建模板变量（{{受伤经过}}先不替换）──
            template_data = {
                '用人单位': company_name,
                '申请人名称': applicant_name,
                '本人姓名': person_name,
                '本人性别': case_data.get('person_gender', ''),
                '本人身份证号': case_data.get('id_card', ''),
                '受伤经过': injury_process,
                '医疗结论': medical_conclusion,
                '引用条例': case_data.get('regulation', ''),
                '申请时间': self._resolve_date_input(self.apply_time_edit.text()),
                '受理时间': self._resolve_date_input(self.accept_time_edit.text()),
            }

            # ── 5. 获取模板路径 ──
            template_path = str(path_utils.get_document_template_path(
                '工伤案件审批表（模板）.docx'
            ))
            if not os.path.exists(template_path):
                self._set_status('模板文件不存在', 'red')
                QMessageBox.critical(self, "错误",
                    f"找不到模板文件:\n{template_path}")
                return

            # ── 5. 预处理模板：根据结论在「认定工伤/不予认定工伤」方框打勾 ──
            import tempfile
            import shutil
            temp_template = os.path.join(tempfile.gettempdir(), '_temp_approval_table.docx')
            shutil.copy2(template_path, temp_template)

            from docx import Document as DocxEditor
            doc_edit = DocxEditor(temp_template)
            check_confirm = (conclusion == "予以认定")
            for table in doc_edit.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if check_confirm and "□认定工伤" in p.text:
                                full = p.text.replace("□认定工伤", "☑认定工伤", 1)
                            elif (not check_confirm) and "□不予认定工伤" in p.text:
                                full = p.text.replace("□不予认定工伤", "☑不予认定工伤", 1)
                            else:
                                continue
                            if p.runs:
                                p.runs[0].text = full
                                for r in p.runs[1:]:
                                    r.text = ""
                            break
            doc_edit.save(temp_template)

            # ── 6. 渲染模板（仅替换带 {{}} 外框的占位符）──
            word = DocxTemplate(temp_template)
            word.render(template_data)

            # ── 6. 保存 ──
            os.makedirs(case_folder, exist_ok=True)

            file_name = f"{person_name}案件审批表.docx" if person_name else "案件审批表.docx"
            target_path = os.path.join(case_folder, file_name)
            counter = 2
            while os.path.exists(target_path):
                file_name = f"{person_name}案件审批表({counter}).docx"
                target_path = os.path.join(case_folder, file_name)
                counter += 1

            word.save(target_path)
            print(f"✅ 案件审批表已保存: {target_path}")

            # ── 清理临时模板 ──
            try:
                os.remove(temp_template)
            except Exception:
                pass

            # ── 9. 打开 ──
            success, message = self.file_service.open_document(target_path)
            if success:
                self._set_status('案件审批表生成成功', 'green')
            else:
                self._set_status(f'审批表已生成，打开失败: {message}', 'orange')

        except Exception as e:
            print(f"❌ 生成审批表异常: {e}")
            import traceback
            traceback.print_exc()
            self._set_status(f'生成审批表失败: {str(e)}', 'red')

    def _on_name_pane_changed(self):
        """当 name_pane 输入完成时，如果是本人角色，自动生成案本号"""
        try:
            role = self.get_current_role_type()
            if role != "本人":
                return
            name = self.name_pane.text().strip()
            if not name:
                return
            # 只有当前没有案本号时才自动生成（避免覆盖用户手动编辑）
            current = self.lineEdit_2.text().strip()
            if not current:
                id_card = self.idnumer_pane.text().strip()
                case_num = self._auto_generate_case_number(name, id_card)
                self.lineEdit_2.setText(case_num)
                self.set_data('案本号', case_num, 'case')
        except Exception as e:
            print(f"自动生成案本号失败: {e}")

    def _auto_generate_case_number(self, person_name: str, id_card: str = "") -> str:
        """根据本人姓名和身份证后四位自动生成案本号"""
        is_death = self.death_case_checkbox.isChecked()
        prefix = "工亡" if is_death else "案本"
        date = datetime.datetime.now().strftime("%Y%m%d")
        id_last4 = id_card[-4:] if id_card and len(id_card) >= 4 else "xxxx"
        return f"{person_name}-{prefix}{date}{id_last4}"
    def get_company_info(self) -> Dict[str, str]:
        """获取公司相关信息（company_pane=用人单位，construction_company=用工单位）"""
        employer = self.construction_company.currentText().strip()
        labor_unit = self.company_pane.currentText().strip()
        site = self.construction_plant.currentText().strip()

        if not employer:
            employer = self.get_data('用工单位', '')
        if not labor_unit:
            labor_unit = self.get_data('用人单位', '')
        if not site:
            site = self.get_data('工地名称', '')

        return {
            '用工单位': employer,
            '用人单位': labor_unit,
            '工地名称': site
        }
    def init_comboboxes(self):
        """初始化所有组合框"""
        self.init_combobox(self.company_pane, self.items_list1)
        self.init_combobox(self.construction_company, self.items_list)
        self.init_combobox(self.construction_plant, self.items_list2)

        self.company_pane.setCurrentIndex(-1)
        self.construction_company.setCurrentIndex(-1)
        self.construction_plant.setCurrentIndex(-1)

    def company(self):
        """更新用人单位信息（company_pane 现为用人单位）"""
        employer_name = self.company_pane.currentText().strip()
        self.set_data('用人单位', employer_name, 'company')

    def sync_employer_to_dict(self):
        """更新用工单位信息（construction_company 现为用工单位）"""
        company_name = self.construction_company.currentText().strip()
        self.set_data('用工单位', company_name, 'company')

    def c_plant(self):
        """更新工地名称信息"""
        site_name = self.construction_plant.currentText().strip()
        self.set_data('工地名称', site_name, 'company')

    def id_clicked(self):
        """读取身份证信息"""
        try:
            dll = windll.LoadLibrary("./sdtapi.dll")
            port = c_int32(1001)
            ifopen = c_int32(1)
            pucManaInfo = create_string_buffer(4)
            pucManaMsg = create_string_buffer(8)
            dll.SDT_StartFindIDCard(port, pucManaInfo, ifopen)
            dll.SDT_SelectIDCard(port, pucManaMsg, ifopen)
            pucCHMsg = create_unicode_buffer(256)
            pucPHMsg = create_string_buffer(1024)
            puiCHMsgLen = c_uint(0)
            puiPHMsgLen = c_uint(0)
            ret = dll.SDT_ReadBaseMsg(port, pucCHMsg, byref(puiCHMsgLen), pucPHMsg,
                                      byref(puiPHMsgLen), ifopen)
            if ret == 65:
                return
            dll.SDT_ClosePort(port)
            self.set_data('当前时期', _date_now() + _time_now(), 'output')

            role = self.get_current_role_type()
            self.process_id(pucCHMsg, role)

        except Exception as e:
            print(f"读取身份证信息时出错: {str(e)}")

    def process_id(self, pucCHMsg, role):
        """处理身份证信息"""
        try:
            name = pucCHMsg.value[0:15].strip()
            self.data_model.update_basic_info(role, {'姓名': name})
            self.set_data(f"{role}姓名", name, 'basic')
            self.name_pane.setText(name)

            if len(pucCHMsg.value) >= 79:
                id_number = pucCHMsg.value[61:79].strip()
                self.data_model.update_basic_info(role, {'身份证号': id_number})
                self.set_data(f"{role}身份证号", id_number, 'basic')
                self.idnumer_pane.setText(id_number)

            if len(pucCHMsg.value) >= 61:
                address = pucCHMsg.value[26:61].strip()
                self.data_model.update_basic_info(role, {'身份证地址': address})
                self.set_data(f"{role}身份证地址", address, 'basic')
                self.textEdit.setText(address)

            self.process_id_info(role)
            self.calculate_age_from_id(role)

        except Exception as e:
            import traceback
            traceback.print_exc()

    def update_role_info(self, role):
        """更新角色信息"""
        try:
            self._save_date_inputs()
            name = self.name_pane.text().strip()
            idcard = self.idnumer_pane.text().strip()
            address = self.textEdit.toPlainText().strip()
            phone = self.lineEdit_4.text().strip()
            position = self.lineEdit_5.text().strip()

            data_updates = {
                f'{role}姓名': name,
                f'{role}身份证号': idcard,
                f'{role}身份证地址': address,
                f'{role}手机号': phone,
            }

            if role == "法人":
                data_updates[f'{role}职务'] = position
            else:
                data_updates[f'{role}岗位'] = position

            for key, value in data_updates.items():
                if value:
                    self.set_data(key, value, 'basic')

            if role == "本人":
                # 自动生成案本号
                current_case = self.lineEdit_2.text().strip()
                if not current_case:
                    id_card = self.idnumer_pane.text().strip()
                    case_num = self._auto_generate_case_number(name, id_card)
                    self.lineEdit_2.setText(case_num)
                    self.set_data('案本号', case_num, 'case')

            if role == "法人":
                company_name = self.get_data('用工单位', '')
                if not company_name:
                    company_name = self.construction_company.currentText().strip()
                    if company_name:
                        self.set_data('用工单位', company_name, 'company')

        except Exception as e:
            import traceback
            traceback.print_exc()
    def _save_death_case_info(self):
        """工亡案件：只保存本人信息 + 生成案本号 + 创建文件夹，不生成笔录"""
        try:
            # 1. 收集本人信息
            self.update_role_info("本人")

            # 2. 获取受伤职工姓名
            person_name = self.get_data("本人姓名", "")
            if not person_name:
                person_name = self.name_pane.text().strip()
                if person_name:
                    self.set_data("本人姓名", person_name, "basic")
                else:
                    self._set_status("请输入姓名", "red")
                    return

            # 3. 更新公司信息
            company_info = self.get_company_info()
            for key, value in company_info.items():
                if value:
                    self.set_data(key, value, "company")

            # 4. 使用或生成案本号
            case_number = self.lineEdit_2.text().strip()
            if not case_number:
                id_card = self.idnumer_pane.text().strip()
                case_number = self._auto_generate_case_number(person_name, id_card)
                self.lineEdit_2.setText(case_number)
            self.set_data("案本号", case_number, "case")

            # 5. 创建案件文件夹
            from datetime import datetime
            case_info = {
                "case_number": case_number,
                "person_name": person_name.strip(),
                "id_card": self.get_data("本人身份证号", ""),
            }
            self.current_case_folder = self.file_service.create_enhanced_case_folder(
                self.BASE_PATH, case_info
            )
            self.current_person_name = person_name

            # 6. 保存本人信息到文件夹（TXT）
            info_file = os.path.join(self.current_case_folder, f"{person_name}基本信息.txt")
            with open(info_file, "w", encoding="utf-8") as f:
                f.write(f"案本号：{case_number}\n")
                f.write(f"案件性质：工亡\n")
                f.write(f"申请类型：{'个人申请' if self.personal_application_checkbox.isChecked() else '单位申请'}\n")
                f.write(f"姓名：{person_name}\n")
                f.write(f"性别：{self.lineEdit.text().strip()}\n")
                f.write(f"身份证号：{self.idnumer_pane.text().strip()}\n")
                f.write(f"年龄：{self.age_pane.text().strip()}\n")
                f.write(f"身份证地址：{self.textEdit.toPlainText().strip()}\n")
                f.write(f"电话：{self.lineEdit_4.text().strip()}\n")
                f.write(f"岗位：{self.lineEdit_5.text().strip()}\n")
                f.write(f"公司：{company_info.get('用工单位', '')}\n")
                f.write(f"用人单位：{company_info.get('用人单位', '')}\n")
                f.write(f"工地名称：{company_info.get('工地名称', '')}\n")
                f.write(f"案件陈述：{self.statement_edit.toPlainText().strip() if hasattr(self, 'statement_edit') else ''}\n")
                f.write(f"保存时间：{_date_now()}{_time_now()}\n")
            print(f"📄 基本信息已保存: {info_file}")

            # 7. 更新案件数据
            self._update_case_in_data(
                case_number=case_number,
                person_name=person_name,
                folder_path=self.current_case_folder,
                transcript_file=f"{person_name}基本信息.txt"
            )

            # 8. 状态提示
            self._set_status(
                f"工亡案件信息已保存 案本号:{case_number} | 可搜索'{person_name}'关联证人和家属笔录",
                "green"
            )
            print(f"✅ 工亡案件已保存: {case_number} 文件夹: {self.current_case_folder}")

        except Exception as e:
            print(f"❌ 保存工亡信息失败: {e}")
            import traceback
            traceback.print_exc()
            self._set_status(f"保存失败: {e}", "red")
    def process_id_info(self, role):
        """处理身份证信息并更新性别显示（使用DataService）"""
        try:
            idcard = self.get_data(f"{role}身份证号", "")
            if not idcard:
                return

            # 使用DataService提取性别
            gender = self.data_service.extract_gender_from_idcard(idcard)
            if gender is None:
                print(f"[process_id_info] 无法从身份证提取性别: {idcard}")
                return

            # 设置性别数据
            self.set_data(f"{role}性别", gender, 'basic')
            self.lineEdit.setText(gender)

        except Exception as e:
            print(f"[process_id_info] 错误: {e}")
            import traceback
            traceback.print_exc()

    def get_current_role_type(self) -> str:
        """获取当前选中的角色类型"""
        if self.radioButton.isChecked():
            return "本人"
        elif self.radioButton_2.isChecked():
            return "证人"
        elif self.radioButton_3.isChecked():
            return "法人"
        return "本人"
    def smart_search_cases(self):
        """按案本号在 cases_data.json 中模糊搜索，并回填主界面"""
        try:
            keyword = self.lineEdit_2.text().strip()
            if not keyword:
                keyword = self.get_data("本人姓名", "")
            if not keyword:
                keyword = self.name_pane.text().strip()
            if not keyword:
                QMessageBox.warning(self, "提示", "请输入案本号进行搜索")
                return

            cases = self._load_cases_data()  # {case_id: case_obj}
            if not cases:
                QMessageBox.information(self, "提示", "还没有保存任何案件数据")
                return

            # 模糊匹配：案本号包含关键字
            kw = keyword.lower()
            matched = [(cid, obj) for cid, obj in cases.items()
                       if kw in str(cid).lower()]

            if not matched:
                QMessageBox.information(self, "提示", f"未找到案本号包含「{keyword}」的案件")
                return

            if len(matched) == 1:
                self._load_case_to_form(matched[0][1])
            else:
                self._show_case_search_dialog(matched)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"搜索失败: {str(e)}")
            import traceback
            traceback.print_exc()

    def _load_case_to_form(self, case_obj: Dict[str, Any]):
        """把单个案件 JSON 回填主界面与数据模型"""
        # 切回本人角色再回填
        self.radioButton.setChecked(True)
        self.clear_role_fields()
        self._apply_case_object(case_obj)
        self.current_case_id = str(case_obj.get('case_id', ''))
        self._set_status(f"已加载案件：{self.current_case_id}", 'green')
        QMessageBox.information(
            self, "加载成功",
            f"已加载案件数据：\n案本号：{self.current_case_id}\n姓名：{case_obj.get('name', '')}"
        )

    def _show_case_search_dialog(self, matched):
        """弹出窗口列出匹配案件供选择"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"选择案件（{len(matched)} 条）")
        dialog.resize(720, 460)

        layout = QVBoxLayout()
        title = QLabel(f"找到 {len(matched)} 条匹配案件，请选择：")
        layout.addWidget(title)

        table = QTableWidget()
        table.setColumnCount(4)
        table.setHorizontalHeaderLabels(['案本号', '姓名', '用人单位', '拟用条例'])
        table.setRowCount(len(matched))
        table.setSelectionBehavior(QTableWidget.SelectRows)
        table.setSelectionMode(QTableWidget.SingleSelection)
        table.setEditTriggers(QTableWidget.NoEditTriggers)
        table.verticalHeader().setVisible(False)

        for i, (case_id, case_obj) in enumerate(matched):
            table.setItem(i, 0, QTableWidgetItem(str(case_id)))
            table.setItem(i, 1, QTableWidgetItem(str(case_obj.get('name', ''))))
            table.setItem(i, 2, QTableWidgetItem(str(case_obj.get('labor_unit', ''))))
            table.setItem(i, 3, QTableWidgetItem(str(case_obj.get('proposed_article', ''))))

        table.resizeColumnsToContents()
        table.horizontalHeader().setStretchLastSection(True)
        if table.rowCount() > 0:
            table.selectRow(0)
        layout.addWidget(table)

        # 双击直接选择
        table.cellDoubleClicked.connect(
            lambda r, c: self._on_case_search_selected(matched, r, dialog)
        )

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(cancel_btn)
        select_btn = QPushButton("选择")
        select_btn.clicked.connect(
            lambda: self._on_case_search_selected(matched, table.currentRow(), dialog)
        )
        btn_layout.addWidget(select_btn)
        layout.addLayout(btn_layout)

        dialog.setLayout(layout)
        dialog.exec_()

    def _on_case_search_selected(self, matched, row: int, dialog: QDialog):
        """用户选中某条案件后的处理"""
        if row < 0 or row >= len(matched):
            QMessageBox.warning(dialog, "提示", "请先选择一行")
            return
        case_obj = matched[row][1]
        dialog.accept()
        self._load_case_to_form(case_obj)



if __name__ == "__main__":
    import sys

    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())