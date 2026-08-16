# template_service.py
import os
import tempfile
import shutil
from typing import Dict, List, Any, Optional
from docx import Document


class TemplateVariableManager:
    """模板变量管理器 - 从原文件迁移过来"""

    def __init__(self, data_model):
        self.data = data_model
        self.variables_cache: Dict[str, Any] = {}
        self.introduction_cache: Dict[str, str] = {}

    def clear_cache(self):
        """清空所有缓存（切换证人/角色等数据变化时调用，避免命中旧数据）"""
        self.variables_cache.clear()
        self.introduction_cache.clear()

    def get_self_introduction(self, role: str,
                              company: str, employer: str, site: str,
                              name: str, position: str, id_address: str = "") -> str:
        """生成自我介绍（带缓存）"""
        cache_key = f"{role}_{name}_{company}_{employer}_{site}"

        if cache_key in self.introduction_cache:
            return self.introduction_cache[cache_key]

        # 根据角色生成自我介绍
        if role == "证人":
            intro = self._generate_witness_intro(name, position, id_address, company, employer, site)
        elif role == "法人":
            intro = self._generate_legal_intro(name, position, id_address, company)
        else:  # 本人
            intro = self._generate_person_intro(name, position, id_address, company, employer, site)

        self.introduction_cache[cache_key] = intro
        return intro

    def _generate_person_intro(self, name: str, position: str, id_address: str,
                               company: str, employer: str, site: str) -> str:
        """生成本人自我介绍"""
        parts = []

        if id_address:
            parts.append(f"身份证地址为{id_address}")

        if company:
            parts.append(f"系{company}的职工")
            if employer:
                parts.append(f"被指派到{employer}")
                if site:
                    parts.append(f"承建的{site}工地")
            elif site:
                parts.append(f"被指派到{site}工地")
        else:
            parts.append("")

        # 构建结果
        base = f"答：我是{name}"
        if parts:
            base += "，" + "，".join(filter(None, parts))

        if position:
            base += f"，从事{position}工作。"
        else:
            base += "。"

        return base

    def _generate_witness_intro(self, name: str, position: str, id_address: str,
                                company: str, employer: str, site: str) -> str:
        """生成证人自我介绍"""
        parts = []

        if id_address:
            parts.append(f"身份证地址为{id_address}")

        if company:
            parts.append(f"系{company}的职工")
            if employer:
                parts.append(f"被指派到{employer}")
                if site:
                    parts.append(f"承建的{site}工地")
            elif site:
                parts.append(f"被指派到{site}工地")
        else:
            parts.append("")

        # 构建结果
        base = f"答：我是{name}"
        if parts:
            base += "，" + "，".join(filter(None, parts))

        if position:
            base += f"，从事{position}工作。"
        else:
            base += "。"

        return base

    def _generate_legal_intro(self, name: str, position: str,
                              id_address: str, company: str) -> str:
        """生成法人自我介绍"""
        parts = []

        if id_address:
            parts.append(f"身份证地址为{id_address}")

        if company:
            if position:
                parts.append(f"系{company}的{position}")
            else:
                parts.append(f"系{company}的负责人")
        else:
            if position:
                parts.append(f"系公司的{position}")
            else:
                parts.append("系公司负责人")

        # 构建结果
        base = f"答：我是{name}"
        if parts:
            base += "，" + "，".join(parts)

        base += "，负责公司的全面管理工作。"
        return base

    def collect_variables(self, role: str, case_type: int,
                          has_employer: bool, has_site: bool) -> Dict[str, Any]:
        """收集所有模板变量"""
        cache_key = f"{role}_{case_type}_{has_employer}_{has_site}"

        if cache_key in self.variables_cache:
            return self.variables_cache[cache_key]

        # 基础数据
        variables = self.data.to_template_dict()

        # 确保有 {{当前时期}} 字段
        from datetime import datetime
        current_date = variables.get('当前日期', datetime.now().strftime('%Y年%m月%d日'))
        current_time = variables.get('当前时间', datetime.now().strftime('%H时%M分'))
        variables['当前时期'] = f"{current_date}{current_time}"

        # 确保有本人年龄和性别
        if role == "本人":
            if '本人年龄' not in variables:
                age = self.data.basic_info.get('本人年龄', '')
                if not age:
                    age = self.data.basic_info.get('本人年龄', self.data.basic_info.get('年龄', ''))
                variables['本人年龄'] = age

            if '本人性别' not in variables:
                gender = self.data.basic_info.get('本人性别', '')
                if not gender:
                    gender = self.data.basic_info.get('本人性别', self.data.basic_info.get('性别', ''))
                variables['本人性别'] = gender

        # 案件类型描述
        case_descriptions = [
            "工伤保险条例第十四条第一款（普通案件）",
            "工伤保险条例第十四条第二款（工作前后案件）",
            "工伤保险条例第十四条第三款（暴力伤害案件）",
            "工伤保险条例第十四条第四款（患职业病案件）",
            "工伤保险条例第十四条第五款（因工外出案件）",
            "工伤保险条例第十四条第六款（上下班时案件）"
        ]
        variables['案件类型'] = case_descriptions[case_type]

        # 用工情况描述
        if has_employer:
            employer = self.data.company_info.get('用人单位', '')
            variables['用工情况'] = f"有用工单位：{employer}" if employer else "无用工单位"
        else:
            variables['用工情况'] = "无用工单位"

        # 工地情况描述
        if has_site:
            site = self.data.company_info.get('工地名称', '')
            variables['工地情况'] = f"在{site}工地" if site else "无特定工地"
        else:
            variables['工地情况'] = "无特定工地"

        # 缓存结果
        self.variables_cache[cache_key] = variables
        return variables


class TemplateService:
    """模板服务 - 负责所有模板相关操作"""

    def __init__(self, template_path: str):
        """
        初始化TemplateService
        :param template_path: 模板文件夹路径
        """
        from path_utils import path_utils

        self.template_path = template_path
        # 直接使用 path_utils 获取模板子目录
        self.template_base_path = str(path_utils.get_talk_template_path())
        self.document_template_path = str(path_utils.get_document_template_path())

    def get_template_path(self, role: str, case_type: int,
                          is_death_case: bool = False, is_personal: bool = False) -> str:
        """
        根据角色和案件类型获取模板路径
        :param role: 角色（本人/证人/法人）
        :param case_type: 案件类型索引（0-5）
        :param is_death_case: 是否为工亡案件
        :param is_personal: 是否为个人申请
        :return: 模板文件完整路径
        """
        if is_death_case:
            if is_personal:
                template_name = f"{role}谈话笔录（个人申请工亡案件）.docx"
            else:
                template_name = f"{role}谈话笔录（工亡案件）.docx"
        else:
            case_templates = {
                0: f"{role}谈话笔录（普通工伤案件）.docx",
                1: f"{role}谈话笔录（工作前后案件）.docx",
                2: f"{role}谈话笔录（暴力伤害案件）.docx",
                3: f"{role}谈话笔录（患职业病案件）.docx",
                4: f"{role}谈话笔录（因工外出案件）.docx",
                5: f"{role}谈话笔录（上下班时案件）.docx"
            }
            template_name = case_templates.get(case_type, f"{role}谈话笔录（普通工伤案件）.docx")

        template_full_path = os.path.join(self.template_base_path, template_name)

        if not os.path.exists(template_full_path):
            default_template = os.path.join(self.template_base_path, f"{role}谈话笔录（普通工伤案件）.docx")
            if os.path.exists(default_template):
                return default_template
            else:
                raise FileNotFoundError(f"模板文件不存在: {template_full_path}")

        return template_full_path

    def get_document_template_path(self, template_name: str) -> str:
        """
        获取文书模板路径
        :param template_name: 模板文件名
        :return: 完整模板路径
        """
        path = os.path.join(self.document_template_path, template_name)
        if not os.path.exists(path):
            raise FileNotFoundError(f"文书模板不存在: {path}")
        return path

    @staticmethod
    def _copy_basic_format(source_paragraph, target_paragraph):
        """复制基本段落格式"""
        try:
            target_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
            target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
            target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after

            if source_paragraph.runs and target_paragraph.runs:
                source_run = source_paragraph.runs[0]
                target_run = target_paragraph.runs[0]
                target_run.font.name = source_run.font.name
                target_run.font.size = source_run.font.size

        except Exception as e:
            print(f"[WARN] 格式复制失败: {e}")

    def _insert_introduction_to_doc(self, doc: Document, introduction_text: str):
        """向文档中插入自我介绍"""
        target_question = "问：请介绍一下你的姓名、住址、工作单位以及从事的工作？"

        for i, paragraph in enumerate(doc.paragraphs):
            if target_question in paragraph.text:
                if i + 1 < len(doc.paragraphs):
                    new_paragraph = doc.paragraphs[i + 1].insert_paragraph_before("")
                else:
                    new_paragraph = doc.add_paragraph("")

                run = new_paragraph.add_run(introduction_text)
                run.underline = True

                if i < len(doc.paragraphs):
                    self._copy_basic_format(doc.paragraphs[i], new_paragraph)
                break

    def _insert_questions_to_doc(self, doc: Document, questions: List[Dict[str, str]],
                                 variables: Dict[str, Any]):
        """向文档中插入问题"""
        for qa in questions:
            answer_hint = qa.get("answer_hint", "")
            if answer_hint:
                for key, value in variables.items():
                    if value:
                        placeholder = "{" + key + "}"
                        answer_hint = answer_hint.replace(placeholder, str(value))

                doc.add_paragraph(qa.get("question", ""))
                answer_para = doc.add_paragraph(answer_hint)

                for run in answer_para.runs:
                    run.underline = True

    def process_template(self, template_path: str, role_type: str, case_type: int,
                         variables: Dict[str, Any], need_special_questions: bool = False) -> str:
        """
        处理模板文件，插入动态内容
        :param template_path: 模板文件路径
        :param role_type: 角色类型
        :param case_type: 案件类型
        :param variables: 模板变量
        :param need_special_questions: 是否需要插入特殊问题
        :return: 处理后的临时模板文件路径
        """
        # 创建临时文件
        temp_path = os.path.join(
            tempfile.gettempdir(),
            f"temp_template_{os.path.basename(template_path)}"
        )

        # 复制模板到临时文件
        shutil.copy2(template_path, temp_path)

        # 打开文档并处理
        doc = Document(temp_path)

        # 插入自我介绍
        self_intro = variables.get('自我介绍内容', '')
        if self_intro:
            self._insert_introduction_to_doc(doc, self_intro)

        # 插入特殊问题（如果需要）
        if need_special_questions:
            questions = self.get_basic_question_template(case_type, role_type)
            if questions:
                self._insert_questions_to_doc(doc, questions, variables)

        # 保存修改
        doc.save(temp_path)
        return temp_path

    @staticmethod
    def get_basic_question_template(case_type: int, role: str = "本人") -> List[Dict[str, str]]:
        """
        根据案件类型和角色获取基本问题模板
        :param case_type: 案件类型索引
        :param role: 角色
        :return: 问题模板列表
        """
        templates = {
            "本人": {
                0: [
                    {"question": "问：你跟{{公司名称}}有没有签订劳动合同关系？有没有参加工伤保险？",
                     "answer_hint": "答：我与{{公司名称}}的劳动关系已经由永嘉县劳动人事争议仲裁委员会【浙永嘉劳人仲案（2025）XXX号】裁决书予以确认。"}
                ],
                1: [
                    {"question": "问：工作前进行了哪些准备工作？",
                     "answer_hint": "答：工作前我准备了{准备工作内容}。"}
                ],
                2: [
                    {"question": "问：请描述暴力伤害发生的具体经过？",
                     "answer_hint": "答：当时{施暴者姓名}突然{伤害行为}。"}
                ]
            },
            "证人": {
                0: [
                    {"question": "问：请介绍一下你的姓名、住址、工作单位以及从事的工作？",
                     "answer_hint": "答：我是{证人姓名}，住在{证人身份证地址}，为{公司名称}的职工，从事{证人岗位}工作。"}
                ]
            },
            "法人": {
                0: [
                    {"question": "问：请介绍一下你的姓名、职务以及公司基本情况？",
                     "answer_hint": "答：我是{法人姓名}，是{公司名称}的{法人职务}，负责公司{负责业务}。"}
                ]
            }
        }

        return templates.get(role, {}).get(case_type, [])